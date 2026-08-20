from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any, Iterable, Mapping
from zoneinfo import ZoneInfo

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from tenable_reports.config.profile import ClientProfile


TEMPLATE_VERSION = "base-docx-v0.1"
REPORT_TITLE = "RELATÓRIO DE VULNERABILIDADES TENABLE"

BLUE = "2E59FC"
NAVY = "0B1F4A"
LIGHT_BLUE = "EAF0FF"
MID_GRAY = "68728A"
LIGHT_GRAY = "F3F5F9"
WHITE = "FFFFFF"
CRITICAL = "FF0000"
HIGH = "F28C00"
MEDIUM = "FFF200"
LOW = "00B050"

ASSET_HEADERS = (
    "IP Address",
    "Asset Name",
    "Crítica",
    "Alta",
    "Média",
    "Baixa",
    "Total",
    "Exploitable",
)
ASSET_PLACEHOLDERS = (
    "{{ASSET_IP}}",
    "{{ASSET_NAME}}",
    "{{CRITICAL}}",
    "{{HIGH}}",
    "{{MEDIUM}}",
    "{{LOW}}",
    "{{TOTAL}}",
    "{{EXPLOITABLE}}",
)
ASSET_WIDTHS_TWIPS = (1450, 2150, 930, 930, 930, 930, 1100, 1420)

MONTHS_PT = (
    "",
    "JANEIRO",
    "FEVEREIRO",
    "MARÇO",
    "ABRIL",
    "MAIO",
    "JUNHO",
    "JULHO",
    "AGOSTO",
    "SETEMBRO",
    "OUTUBRO",
    "NOVEMBRO",
    "DEZEMBRO",
)


@dataclass(frozen=True, slots=True)
class BaseReportRenderResult:
    output_path: Path
    template_version: str
    client_id: str
    period_id: str
    top_asset_rows: int
    masked_sensitive_fields: bool


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell: Any, width_twips: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def _set_table_fixed(table: Any, width_twips: int) -> None:
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def _set_paragraph_bottom_border(paragraph: Any, color: str, size: int = 18) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _set_run_font(run: Any, *, size: float | None = None, color: str | None = None,
                  bold: bool | None = None, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _set_language(run: Any, language: str = "pt-BR") -> None:
    properties = run._element.get_or_add_rPr()
    lang = properties.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        properties.append(lang)
    lang.set(qn("w:val"), language)


def _add_picture(paragraph: Any, image_path: Path, *, width: Any, alt_text: str) -> Any:
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=width)
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    return shape


def _add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, result, end):
        run._r.append(element)
    _set_run_font(run, size=8, color=MID_GRAY)


def _configure_styles(document: DocxDocument) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(BLUE)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True

def _configure_page(section: Any) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    section.different_first_page_header_footer = True


def _configure_headers_and_footers(document: DocxDocument, assets_dir: Path) -> None:
    section = document.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.4))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_fixed(table, 9860)
    _set_cell_width(table.cell(0, 0), 3200)
    _set_cell_width(table.cell(0, 1), 6660)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    left_paragraph = left.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_picture(
        left_paragraph,
        assets_dir / "itprotect-logo.png",
        width=Inches(1.55),
        alt_text="Logotipo ITProtect",
    )
    right_paragraph = right.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title = right_paragraph.add_run("RELATÓRIO DE VULNERABILIDADES\n")
    _set_run_font(title, size=7.5, color=BLUE, bold=True)
    client = right_paragraph.add_run("{{CLIENT_NAME}}")
    _set_run_font(client, size=7.5, color=NAVY, bold=True)
    _set_paragraph_bottom_border(right_paragraph, BLUE, size=8)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Cm(17.4))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    _set_table_fixed(footer_table, 9860)
    _set_cell_width(footer_table.cell(0, 0), 7900)
    _set_cell_width(footer_table.cell(0, 1), 1960)
    left_footer = footer_table.cell(0, 0).paragraphs[0]
    left_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left_footer.add_run(
        "Relatório Tenable  |  Uso restrito  |  {{TEMPLATE_VERSION}}"
    )
    _set_run_font(run, size=7.5, color=MID_GRAY)
    _add_page_number(footer_table.cell(0, 1).paragraphs[0])


def _add_spacer(document: DocxDocument, points: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing = Pt(1)


def _build_cover(document: DocxDocument, assets_dir: Path) -> None:
    pattern = document.add_paragraph()
    pattern.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pattern.paragraph_format.space_after = Pt(42)
    _add_picture(
        pattern,
        assets_dir / "brand-pattern.png",
        width=Cm(17.4),
        alt_text="Grafismo geométrico azul da identidade visual",
    )

    client_paragraph = document.add_paragraph()
    client_paragraph.paragraph_format.space_after = Pt(9)
    client_run = client_paragraph.add_run("{{CLIENT_NAME}}")
    _set_run_font(client_run, size=16, color=NAVY, bold=True)

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(18)
    title_run = title_paragraph.add_run("RELATÓRIO DE\nVULNERABILIDADES\nTENABLE")
    _set_run_font(title_run, size=28, color=BLUE, bold=True)

    period_paragraph = document.add_paragraph()
    period_paragraph.paragraph_format.space_after = Pt(5)
    period_run = period_paragraph.add_run("{{PERIOD_LABEL}}")
    _set_run_font(period_run, size=15, color=NAVY, bold=True)

    range_paragraph = document.add_paragraph()
    range_run = range_paragraph.add_run("{{PERIOD_RANGE}}")
    _set_run_font(range_run, size=9, color=MID_GRAY)

    _add_spacer(document, 78)
    logos = document.add_table(rows=1, cols=2)
    logos.alignment = WD_TABLE_ALIGNMENT.CENTER
    logos.autofit = False
    _set_table_fixed(logos, 9860)
    for cell, width in zip(logos.rows[0].cells, (4930, 4930), strict=True):
        _set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tenable = logos.cell(0, 0).paragraphs[0]
    tenable.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_picture(
        tenable,
        assets_dir / "tenable-logo.png",
        width=Inches(1.65),
        alt_text="Logotipo Tenable",
    )
    itprotect = logos.cell(0, 1).paragraphs[0]
    itprotect.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_picture(
        itprotect,
        assets_dir / "itprotect-logo.png",
        width=Inches(2.0),
        alt_text="Logotipo ITProtect",
    )

    document.add_page_break()


def _format_card_value(cell: Any, placeholder: str) -> None:
    _set_cell_shading(cell, LIGHT_BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(placeholder)
    _set_run_font(run, size=18, color=BLUE, bold=True)


def _format_card_label(cell: Any, label: str) -> None:
    _set_cell_shading(cell, LIGHT_BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(label)
    _set_run_font(run, size=7.5, color=NAVY, bold=True)


def _build_summary(document: DocxDocument) -> None:
    heading = document.add_paragraph("1. VISÃO GERAL DO PERÍODO", style="Heading 1")
    _set_paragraph_bottom_border(heading, BLUE, size=8)

    intro = document.add_paragraph()
    intro.add_run(
        "Esta prova editorial é gerada exclusivamente a partir do dataset mensal "
        "materializado, sem nova consulta à API durante a renderização. "
    )
    grain = intro.add_run("Grão das contagens: instância de finding.")
    grain.bold = True

    period = document.add_paragraph()
    period.add_run("Período: ").bold = True
    period.add_run("{{PERIOD_RANGE}}  |  ")
    period.add_run("Execução: ").bold = True
    period.add_run("{{EXECUTION_LABEL}}")

    cards = document.add_table(rows=2, cols=4)
    cards.alignment = WD_TABLE_ALIGNMENT.CENTER
    cards.autofit = False
    _set_table_fixed(cards, 9860)
    for row in cards.rows:
        for cell in row.cells:
            _set_cell_width(cell, 2465)
            _set_cell_shading(cell, LIGHT_BLUE)
    values = (
        "{{NM}}",
        "{{VA}}",
        "{{EX}}",
        "{{TA}}",
    )
    labels = (
        "FINDINGS NÃO MITIGADOS",
        "ATIVOS VULNERÁVEIS",
        "FINDINGS EXPLOITABLE",
        "ATIVOS PRIORIZADOS",
    )
    for index, value in enumerate(values):
        _format_card_value(cards.cell(0, index), value)
        _format_card_label(cards.cell(1, index), labels[index])
    _prevent_row_split(cards.rows[0])
    _prevent_row_split(cards.rows[1])

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_before = Pt(5)
    run = metadata.add_run(
        "Definição métrica {{METRIC_DEFINITION_VERSION}}  |  "
        "run_id {{RUN_ID}}  |  coleta {{COLLECTION_STATUS}}"
    )
    _set_run_font(run, size=7.5, color=MID_GRAY)


def _format_asset_table(table: Any) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_fixed(table, sum(ASSET_WIDTHS_TWIPS))
    header_fills = (BLUE, BLUE, CRITICAL, HIGH, MEDIUM, LOW, BLUE, NAVY)
    header_font_colors = (WHITE, WHITE, "000000", "000000", "000000", "000000", WHITE, WHITE)
    for index, cell in enumerate(table.rows[0].cells):
        _set_cell_width(cell, ASSET_WIDTHS_TWIPS[index])
        _set_cell_shading(cell, header_fills[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            _set_run_font(run, size=7.2, color=header_font_colors[index], bold=True)
    _set_repeat_table_header(table.rows[0])
    _prevent_row_split(table.rows[0])

    for row_index, row in enumerate(table.rows[1:], start=1):
        _prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            _set_cell_width(cell, ASSET_WIDTHS_TWIPS[column_index])
            _set_cell_shading(cell, WHITE if row_index % 2 else LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if column_index in (0, 1)
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                _set_run_font(run, size=7.2, color=NAVY)


def _build_top_assets(document: DocxDocument) -> None:
    heading = document.add_paragraph("2. PRINCIPAIS ATIVOS VULNERÁVEIS", style="Heading 1")
    _set_paragraph_bottom_border(heading, BLUE, size=8)
    note = document.add_paragraph()
    note.add_run(
        "Ranking dos ativos pelo total de findings não mitigados. "
        "Exploitable é um subconjunto de Total e não integra a soma das severidades. "
    )
    sensitive = note.add_run(
        "IP Address e Asset Name seguem a política de exposição escolhida na execução."
    )
    sensitive.italic = True

    table = document.add_table(rows=2, cols=len(ASSET_HEADERS))
    table.style = "Table Grid"
    for index, header in enumerate(ASSET_HEADERS):
        table.cell(0, index).text = header
        table.cell(1, index).text = ASSET_PLACEHOLDERS[index]
    _format_asset_table(table)

    methodology = document.add_paragraph()
    methodology.paragraph_format.space_before = Pt(5)
    run = methodology.add_run(
        "Fonte: report-dataset.json. Severidade informativa não participa deste perfil. "
        "Valores ausentes não são convertidos em zero."
    )
    _set_run_font(run, size=7.5, color=MID_GRAY)


def _enable_field_updates(document: DocxDocument) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def create_base_template(
    output_path: str | Path,
    *,
    assets_dir: str | Path,
) -> Path:
    """Create the controlled, sanitised Phase 5 base template."""

    output = Path(output_path)
    assets = Path(assets_dir)
    required_assets = ("tenable-logo.png", "itprotect-logo.png", "brand-pattern.png")
    missing = [name for name in required_assets if not (assets / name).is_file()]
    if missing:
        raise ValueError(f"Assets ausentes para o template: {', '.join(missing)}")

    document = Document()
    section = document.sections[0]
    _configure_page(section)
    _configure_styles(document)
    _configure_headers_and_footers(document, assets)
    _build_cover(document, assets)
    _build_summary(document)
    _build_top_assets(document)
    _enable_field_updates(document)

    document.core_properties.title = "Template controlado do relatório-base Tenable"
    document.core_properties.subject = "Fase 5 - capa e seção de prova"
    document.core_properties.author = "ITProtect"
    document.core_properties.keywords = f"Tenable, relatório, {TEMPLATE_VERSION}"
    document.core_properties.comments = (
        "Template sanitizado: não contém nomes de pessoas, e-mails, IPs ou hostnames reais."
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def _iter_paragraphs(document: DocxDocument) -> Iterable[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.footer,
            section.first_page_footer,
        ):
            yield from container.paragraphs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _replace_tokens(document: DocxDocument, replacements: Mapping[str, str]) -> None:
    for paragraph in _iter_paragraphs(document):
        for run in paragraph.runs:
            original = run.text
            text = original
            for token, value in replacements.items():
                text = text.replace(token, value)
            # Assigning .text to an image-only run removes its drawing XML.
            if text != original:
                run.text = text


def _parse_utc(value: str) -> datetime:
    parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=UTC)
    return parsed.astimezone(UTC)


def _period_labels(period: Mapping[str, Any]) -> tuple[str, str]:
    timezone_name = str(period.get("timezone") or "UTC")
    timezone = ZoneInfo(timezone_name)
    start = _parse_utc(str(period["start_at"])).astimezone(timezone)
    end_exclusive = _parse_utc(str(period["end_at"])).astimezone(timezone)
    end_inclusive = end_exclusive - timedelta(microseconds=1)
    if start.year == end_inclusive.year and start.month == end_inclusive.month:
        label = f"{MONTHS_PT[start.month]}/{start.year}"
    else:
        label = f"{start:%d/%m/%Y} A {end_inclusive:%d/%m/%Y}"
    date_range = f"{start:%d/%m/%Y} a {end_inclusive:%d/%m/%Y}"
    return label, date_range


def _integer(value: Any, field_name: str) -> int:
    if isinstance(value, bool):
        raise ValueError(f"{field_name} deve ser inteiro.")
    try:
        result = int(value)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{field_name} deve ser inteiro.") from exc
    if result < 0:
        raise ValueError(f"{field_name} não pode ser negativo.")
    return result


def _validate_dataset(dataset: Mapping[str, Any], profile: ClientProfile) -> None:
    if dataset.get("schema_version") != 1:
        raise ValueError("A Fase 5 exige report-dataset schema_version=1.")
    if dataset.get("client_id") != profile.client_id:
        raise ValueError("client_id do dataset difere do perfil selecionado.")
    if not isinstance(dataset.get("period"), Mapping):
        raise ValueError("Dataset sem período válido.")
    if not isinstance(dataset.get("metrics"), Mapping):
        raise ValueError("Dataset sem métricas válidas.")
    top_assets = dataset.get("top_assets")
    if not isinstance(top_assets, list):
        raise ValueError("Dataset sem top_assets válido.")
    for index, row in enumerate(top_assets):
        if not isinstance(row, Mapping):
            raise ValueError(f"top_assets[{index}] deve ser um objeto.")
        total = _integer(row.get("total"), f"top_assets[{index}].total")
        exploitable = _integer(
            row.get("exploitable"), f"top_assets[{index}].exploitable"
        )
        if exploitable > total:
            raise ValueError(
                f"top_assets[{index}].exploitable não pode exceder total."
            )


def _find_asset_table(document: DocxDocument) -> Any:
    for table in document.tables:
        if table.rows and tuple(cell.text.strip() for cell in table.rows[0].cells) == ASSET_HEADERS:
            return table
    raise ValueError("O template não contém a tabela contratual de ativos.")


def _fill_asset_table(
    document: DocxDocument,
    rows: list[Mapping[str, Any]],
    *,
    mask_sensitive: bool,
) -> None:
    table = _find_asset_table(document)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)

    if not rows:
        empty = table.add_row()
        merged = empty.cells[0]
        for cell in empty.cells[1:]:
            merged = merged.merge(cell)
        merged.text = "Nenhum ativo disponível para o período e os filtros selecionados."
        merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        for item in rows:
            row = table.add_row()
            values = (
                "" if mask_sensitive else str(item.get("ip_address") or ""),
                "" if mask_sensitive else str(item.get("asset_name") or ""),
                str(_integer(item.get("critical"), "critical")),
                str(_integer(item.get("high"), "high")),
                str(_integer(item.get("medium"), "medium")),
                str(_integer(item.get("low"), "low")),
                str(_integer(item.get("total"), "total")),
                str(_integer(item.get("exploitable"), "exploitable")),
            )
            for index, value in enumerate(values):
                row.cells[index].text = value
    _format_asset_table(table)


def _load_dataset(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except OSError as exc:
        raise ValueError(f"Não foi possível ler o dataset: {path}") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"Dataset JSON inválido na linha {exc.lineno}.") from exc
    if not isinstance(value, dict):
        raise ValueError("O dataset deve conter um objeto JSON na raiz.")
    return value


def generate_base_report(
    *,
    template_path: str | Path,
    dataset_path: str | Path,
    profile: ClientProfile,
    output_path: str | Path,
    mask_sensitive: bool = False,
) -> BaseReportRenderResult:
    """Render a DOCX only from a profile and a materialised report dataset."""

    template = Path(template_path)
    dataset_file = Path(dataset_path)
    output = Path(output_path)
    if not template.is_file():
        raise ValueError(f"Template Word não encontrado: {template}")
    dataset = _load_dataset(dataset_file)
    _validate_dataset(dataset, profile)

    period = dataset["period"]
    period_label, period_range = _period_labels(period)
    metrics = dataset["metrics"]
    non_mitigated = metrics.get("non_mitigated") or {}
    top_assets = dataset.get("top_assets") or []
    execution_type = str(dataset.get("execution_type") or "AUTOMATIC_MONTHLY")
    execution_label = (
        "Automática mensal"
        if execution_type == "AUTOMATIC_MONTHLY"
        else "Manual / pontual"
    )

    document = Document(template)
    replacements = {
        "{{CLIENT_NAME}}": profile.display_name,
        "{{PERIOD_LABEL}}": period_label,
        "{{PERIOD_RANGE}}": period_range,
        "{{EXECUTION_LABEL}}": execution_label,
        "{{NM}}": str(_integer(non_mitigated.get("total"), "non_mitigated.total")),
        "{{VA}}": str(
            _integer(non_mitigated.get("vulnerable_assets"), "non_mitigated.vulnerable_assets")
        ),
        "{{EX}}": str(
            _integer(non_mitigated.get("exploitable"), "non_mitigated.exploitable")
        ),
        "{{TA}}": str(len(top_assets)),
        "{{METRIC_DEFINITION_VERSION}}": str(dataset.get("metric_definition_version") or "não informado"),
        "{{RUN_ID}}": str(dataset.get("run_id") or "não informado"),
        "{{COLLECTION_STATUS}}": str((dataset.get("collection_timing") or {}).get("status") or "não informado"),
        "{{TEMPLATE_VERSION}}": TEMPLATE_VERSION,
    }
    _replace_tokens(document, replacements)
    _fill_asset_table(document, top_assets, mask_sensitive=mask_sensitive)
    _enable_field_updates(document)
    document.core_properties.title = f"{REPORT_TITLE} - {profile.display_name} - {period_label}"
    document.core_properties.subject = "Relatório-base automatizado - Fase 5"
    document.core_properties.comments = (
        f"Gerado a partir de {dataset_file.name}; template {TEMPLATE_VERSION}; "
        f"campos sensíveis {'mascarados' if mask_sensitive else 'conforme dataset'}."
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return BaseReportRenderResult(
        output_path=output,
        template_version=TEMPLATE_VERSION,
        client_id=profile.client_id,
        period_id=str(period.get("period_id") or ""),
        top_asset_rows=len(top_assets),
        masked_sensitive_fields=mask_sensitive,
    )
