from __future__ import annotations

import json
import re
from dataclasses import dataclass
from datetime import timedelta
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence
from zoneinfo import ZoneInfo

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from tenable_reports.config.profile import ClientProfile
from tenable_reports.presentation import base_report_docx as base
from tenable_reports.presentation import editorial_catalog as copy
from tenable_reports.presentation.source_filters import add_source_filter_note
from tenable_reports.presentation.translation import TextTranslator, translate_in_chunks


FULL_TEMPLATE_VERSION = "base-fiel-v2.0"
FULL_REPORT_TITLE = "RELATÓRIO DE VULNERABILIDADES TENABLE"


@dataclass(frozen=True, slots=True)
class FullBaseReportRenderResult:
    output_path: Path
    template_version: str
    client_id: str
    period_id: str
    top_asset_rows: int
    top_open_rows: int
    masked_sensitive_fields: bool


def _load_dataset(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except OSError as exc:
        raise ValueError(f"Não foi possível ler o dataset: {path}") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"Dataset JSON inválido na linha {exc.lineno}.") from exc
    if not isinstance(value, dict):
        raise ValueError("O dataset deve conter um objeto JSON na raiz.")
    return value


def _validate_dataset(dataset: Mapping[str, Any], profile: ClientProfile) -> None:
    base._validate_dataset(dataset, profile)
    for field in (
        "top_open_vulnerabilities",
        "top_fixed_vulnerabilities",
        "top_resurfaced_vulnerabilities",
    ):
        value = dataset.get(field, [])
        if not isinstance(value, list):
            raise ValueError(f"{field} deve ser uma lista.")
    coverage = dataset.get("source_coverage")
    coverage = coverage if isinstance(coverage, Mapping) else {}
    if (
        profile.presentation.vm_top5_include_output
        and not coverage.get("plugin_output_included")
    ):
        raise ValueError(
            "presentation.vm_top5_include_output foi habilitado, mas Plugin Output "
            "nao foi coletado no dataset."
        )
    if (
        profile.presentation.was_top5_include_output
        and not coverage.get("was_plugin_output_included")
    ):
        raise ValueError(
            "presentation.was_top5_include_output foi habilitado, mas Plugin Output "
            "WAS nao foi coletado no dataset."
        )


def _clear_body_after_cover_break(document: DocxDocument) -> None:
    body = document._element.body
    found_cover_break = False
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if found_cover_break:
            body.remove(child)
            continue
        if child.tag == qn("w:p") and child.xpath('.//w:br[@w:type="page"]'):
            found_cover_break = True
    if not found_cover_break:
        raise ValueError("O template não contém a quebra de página contratual da capa.")


def _configure_styles(document: DocxDocument) -> None:
    base._configure_styles(document)
    for name, size in (("Heading 1", 14), ("Heading 2", 11), ("Heading 3", 10), ("Heading 4", 9)):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(base.NAVY)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True
    document.styles["Normal"].paragraph_format.space_after = Pt(6)
    document.styles["Normal"].paragraph_format.line_spacing = 1.08


def _sanitize_header_footer(document: DocxDocument, client_name: str) -> None:
    for section in document.sections:
        for table in section.header.tables:
            if len(table.columns) >= 2:
                table.cell(0, 1).text = f"RELATÓRIO DE VULNERABILIDADES\n{client_name}"
        for table in section.footer.tables:
            if table.rows:
                table.cell(0, 0).text = ""


def _sanitize_properties(document: DocxDocument, *, title: str) -> None:
    properties = document.core_properties
    properties.title = title
    properties.subject = "Relatório Tenable"
    properties.author = "ITProtect"
    properties.last_modified_by = ""
    properties.keywords = "Tenable, vulnerabilidades"
    properties.comments = ""
    properties.category = ""


def _set_keep_with_next(paragraph: Any, enabled: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = enabled


def _toc_field(document: DocxDocument) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-4" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def _heading(document: DocxDocument, text: str, level: int = 1) -> Any:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        base._set_paragraph_bottom_border(paragraph, base.BLUE, size=8)
    return paragraph


def _paragraph(document: DocxDocument, text: str = "", *, bold: bool = False) -> Any:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    base._set_run_font(run, size=9, color=base.NAVY, bold=bold)
    base._set_language(run)
    return paragraph


def _bullet(document: DocxDocument, text: str) -> Any:
    paragraph = document.add_paragraph(style="List Paragraph")
    properties = paragraph._p.get_or_add_pPr()
    numbering = properties.find(qn("w:numPr"))
    if numbering is None:
        numbering = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "0")
        number_id = OxmlElement("w:numId")
        number_id.set(qn("w:val"), "1")
        numbering.extend((level, number_id))
        properties.append(numbering)
    run = paragraph.add_run(text)
    base._set_run_font(run, size=9, color=base.NAVY)
    base._set_language(run)
    return paragraph


def _period_dates(period: Mapping[str, Any]) -> tuple[Any, Any]:
    timezone = ZoneInfo(str(period.get("timezone") or "UTC"))
    start = base._parse_utc(str(period["start_at"])).astimezone(timezone)
    end = base._parse_utc(str(period["end_at"])).astimezone(timezone) - timedelta(microseconds=1)
    return start, end


def _written_period(start: Any, end: Any) -> str:
    if start.year == end.year and start.month == end.month:
        return f"{start.day:02d} e {end.day:02d} de {base.MONTHS_PT[start.month].lower()} de {start.year}"
    return (
        f"{start.day:02d} de {base.MONTHS_PT[start.month].lower()} de {start.year} e "
        f"{end.day:02d} de {base.MONTHS_PT[end.month].lower()} de {end.year}"
    )


def _period_paragraph(document: DocxDocument, start: Any, end: Any) -> None:
    paragraph = document.add_paragraph()
    prefix = paragraph.add_run("Período deste relatório compreende-se entre ")
    date_run = paragraph.add_run(f"{start:%d/%m/%Y} e {end:%d/%m/%Y}.")
    base._set_run_font(prefix, size=9, color=base.NAVY)
    base._set_run_font(date_run, size=9, color=base.NAVY, bold=True)
    base._set_language(prefix)
    base._set_language(date_run)


def _overview_paragraph(document: DocxDocument, start: Any, end: Any) -> None:
    paragraph = document.add_paragraph()
    prefix = paragraph.add_run("Segue um Overview, em um período que se compreende entre os dias ")
    dates = paragraph.add_run(_written_period(start, end))
    suffix = paragraph.add_run(copy.OVERVIEW_SUFFIX)
    for run in (prefix, dates, suffix):
        base._set_run_font(run, size=9, color=base.NAVY, bold=run is dates)
        base._set_language(run)


def _cell_margins(cell: Any) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")


def _simple_table(
    document: DocxDocument,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    *,
    widths: Sequence[int] | None = None,
    left_columns: frozenset[int] = frozenset(),
    keep_together: bool = False,
    header_fills: Sequence[str] | None = None,
    empty_message: str | None = copy.EMPTY_TABLE_MONTH,
) -> Any:
    if not headers:
        raise ValueError("Tabela sem cabeçalho.")
    if widths is None:
        widths = tuple(9200 // len(headers) for _ in headers)
    if len(widths) != len(headers):
        raise ValueError("Cabeçalhos e larguras devem ter a mesma quantidade.")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    base._set_table_fixed(table, sum(widths))
    fills = tuple(header_fills or (base.BLUE,) * len(headers))
    for index, (header, width) in enumerate(zip(headers, widths, strict=True)):
        cell = table.cell(0, index)
        cell.text = header
        base._set_cell_width(cell, width)
        base._set_cell_shading(cell, fills[index])
        _cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            dark_text = fills[index] in (base.CRITICAL, base.HIGH, base.MEDIUM, base.LOW)
            base._set_run_font(run, size=7.2, color="000000" if dark_text else base.WHITE, bold=True)
    base._set_repeat_table_header(table.rows[0])
    base._prevent_row_split(table.rows[0])
    for row_index, values in enumerate(rows, start=1):
        if len(values) != len(headers):
            raise ValueError("Linha com quantidade de colunas diferente do cabeçalho.")
        row = table.add_row()
        base._prevent_row_split(row)
        for column, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            cell.text = "" if value is None else str(value)
            base._set_cell_width(cell, widths[column])
            base._set_cell_shading(cell, base.WHITE if row_index % 2 else base.LIGHT_GRAY)
            _cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column in left_columns else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                base._set_run_font(run, size=7.2, color=base.NAVY)
                base._set_language(run)
    if keep_together and rows:
        for row in table.rows[:-1]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _set_keep_with_next(paragraph)
    if not rows and empty_message:
        _paragraph(document, empty_message)
    return table


def _title_table(document: DocxDocument, title: str, headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> Any:
    table = document.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    merged = table.cell(0, 0)
    for cell in table.rows[0].cells[1:]:
        merged = merged.merge(cell)
    merged.text = title
    base._set_cell_shading(merged, base.BLUE)
    for run in merged.paragraphs[0].runs:
        base._set_run_font(run, size=8, color=base.WHITE, bold=True)
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, header in enumerate(headers):
        cell = table.cell(1, index)
        cell.text = header
        base._set_cell_shading(cell, base.LIGHT_BLUE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            base._set_run_font(run, size=7.2, color=base.NAVY, bold=True)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = "" if value is None else str(value)
            for run in row.cells[index].paragraphs[0].runs:
                base._set_run_font(run, size=7.2, color=base.NAVY)
        base._prevent_row_split(row)
    return table


def _control_document(document: DocxDocument, generated_date: str) -> None:
    _heading(document, "CONTROLE DE DOCUMENTO")
    _title_table(document, "Preparação", ("Ação", "Nome", "Data"), (("Criação do Documento", "", generated_date),))
    document.add_paragraph()
    _title_table(
        document,
        "Controle de Versionamento",
        ("Versão", "Data da Versão", "Seções Afetadas", "Alteração", "Alterado por"),
        (("1.0", generated_date, "Todas", "Elaboração do conteúdo", ""),),
    )
    document.add_paragraph()
    _title_table(document, "Lista de Distribuição", ("Nome", "Organização", "E-mail"), (("", "", ""), ("", "", "")))


def _count(value: Any) -> str:
    return "" if value is None else str(int(value))


def _by_severity(metric: Mapping[str, Any], key: str, severity: str) -> Any:
    value = metric.get(key)
    return value.get(severity) if isinstance(value, Mapping) else None


def _overview_table(
    document: DocxDocument,
    metrics: Mapping[str, Any],
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    non = metrics.get("non_mitigated") or {}
    fixed = metrics.get("mitigated") or {}
    severities = (("Crítica", "critical"), ("Alta", "high"), ("Média", "medium"), ("Baixa", "low"))
    rows: list[tuple[Any, ...]] = [(
        "TOTAL",
        _count(fixed.get("total")),
        _count(non.get("total")),
        _count(non.get("exploitable")),
        _count(non.get("patch_available_over_30_days")),
    )]
    for label, key in severities:
        rows.append((
            label,
            _count(_by_severity(fixed, "by_severity", key)),
            _count(_by_severity(non, "by_severity", key)),
            _count(_by_severity(non, "exploitable_by_severity", key)),
            _count(_by_severity(non, "patch_available_over_30_days_by_severity", key)),
        ))
    fills = (base.BLUE, base.BLUE, base.BLUE, base.BLUE, base.BLUE)
    table = _simple_table(
        document,
        ("", "Mitigado", "Não Mitigado", "Explorável", "Patchs disponíveis > 30d"),
        rows,
        widths=(1150, 1450, 1800, 1550, 3250),
        left_columns=frozenset({0}),
        header_fills=fills,
    )
    severity_fills = (base.CRITICAL, base.HIGH, base.MEDIUM, base.LOW)
    for row, fill in zip(table.rows[2:], severity_fills, strict=True):
        base._set_cell_shading(row.cells[0], fill)
    add_source_filter_note(
        document, dataset, "overview", enabled=show_source_filters
    )


def _top_assets_table(document: DocxDocument, rows: Sequence[Mapping[str, Any]], mask_sensitive: bool) -> None:
    values = []
    for item in rows:
        values.append((
            "" if mask_sensitive else item.get("ip_address") or "",
            "" if mask_sensitive else item.get("asset_name") or "",
            _count(item.get("critical")),
            _count(item.get("high")),
            _count(item.get("medium")),
            _count(item.get("low")),
            _count(item.get("total")),
            _count(item.get("exploitable")),
        ))
    _simple_table(
        document,
        base.ASSET_HEADERS,
        values,
        widths=base.ASSET_WIDTHS_TWIPS,
        left_columns=frozenset({0, 1}),
        header_fills=(base.BLUE, base.BLUE, base.CRITICAL, base.HIGH, base.MEDIUM, base.LOW, base.BLUE, base.BLUE),
    )


def _severity_pt(value: Any) -> str:
    return {"CRITICAL": "Crítica", "HIGH": "Alta", "MEDIUM": "Média", "LOW": "Baixa", "INFO": "Informativa"}.get(str(value).upper(), str(value or ""))


def _compact_rows(items: Any) -> list[tuple[Any, ...]]:
    if not isinstance(items, list):
        return []
    return [(
        item.get("plugin_id"),
        item.get("plugin_name") or "",
        item.get("plugin_family") or "",
        _severity_pt(item.get("severity")),
        item.get("finding_instances", item.get("affected_assets", "")),
        item.get("vpr_score") if item.get("vpr_score") is not None else 0,
    ) for item in items[:5] if isinstance(item, Mapping)]


def _principal_vulnerabilities(
    document: DocxDocument, dataset: Mapping[str, Any], *, show_source_filters: bool
) -> None:
    _heading(document, "VISÃO GERAL DAS PRINCIPAIS VULNERABILIDADES")
    _paragraph(document, copy.PRINCIPAL_VULNERABILITIES_INTRO)
    _paragraph(document, copy.FIXED_REMINDER)
    headers = ("Plugin ID", "Nome", "Família OS", "Severidade", "Total", "VPR")
    widths = (900, 3000, 2050, 1050, 850, 850)
    for title, field in (
        ("4.1. Vulnerabilidades Mitigadas", "top_fixed_vulnerabilities"),
        ("4.2. Vulnerabilidades Não Mitigadas", "top_open_vulnerabilities"),
        ("4.3. Vulnerabilidades Ressurgidas", "top_resurfaced_vulnerabilities"),
    ):
        _heading(document, title, 2)
        _simple_table(document, headers, _compact_rows(dataset.get(field)), widths=widths, left_columns=frozenset({1, 2}))
        add_source_filter_note(document, dataset, field, enabled=show_source_filters)


def _description_chunks(text: str, *, max_chars: int = 900) -> tuple[str, ...]:
    normalized = " ".join(str(text or "").split())
    if not normalized:
        return ()
    sentences = re.split(r"(?<=[.!?;:])\s+", normalized)
    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        pieces = [sentence]
        if len(sentence) > max_chars:
            words = sentence.split()
            pieces = []
            part = ""
            for word in words:
                candidate = f"{part} {word}".strip()
                if part and len(candidate) > max_chars:
                    pieces.append(part)
                    part = word
                else:
                    part = candidate
            if part:
                pieces.append(part)
        for piece in pieces:
            candidate = f"{current} {piece}".strip()
            if current and len(candidate) > max_chars:
                chunks.append(current)
                current = piece
            else:
                current = candidate
    if current:
        chunks.append(current)
    return tuple(chunks)


def _labeled_blocks(
    document: DocxDocument,
    label: str,
    text: Any,
    *,
    translator: TextTranslator | None = None,
) -> None:
    _heading(document, label, 4)
    chunks = (
        translate_in_chunks(str(text or ""), translator, max_chars=900)
        if translator is not None
        else _description_chunks(str(text or ""))
    )
    for chunk in chunks:
        _paragraph(document, chunk)


def _host_rows(item: Mapping[str, Any], mask_sensitive: bool, include_output: bool) -> list[tuple[Any, ...]]:
    rows = []
    for host in item.get("hosts") or []:
        if not isinstance(host, Mapping):
            continue
        ips = host.get("ip_addresses") or []
        values: list[Any] = [
            "" if mask_sensitive else host.get("asset_name") or "",
            "" if mask_sensitive else (ips[0] if isinstance(ips, list) and ips else ""),
            host.get("port") if host.get("port") is not None else "",
            host.get("protocol") or "",
        ]
        if include_output:
            values.append(host.get("plugin_output") or "")
        rows.append(tuple(values))
    return rows


def _vulnerability_details(
    document: DocxDocument,
    items: Any,
    *,
    dataset: Mapping[str, Any],
    source_table_id: str,
    show_source_filters: bool,
    heading_level: int,
    number_prefix: str,
    mask_sensitive: bool,
    include_output: bool,
    translator: TextTranslator | None = None,
    protocol_header: str = "PROTOCOLO",
    source_extra_filters: Mapping[str, Any] | None = None,
) -> int:
    if not isinstance(items, list):
        return 0
    rendered = 0
    for index, item in enumerate(items[:5], start=1):
        if not isinstance(item, Mapping):
            continue
        rendered += 1
        _heading(document, f"{number_prefix}.{index}. {item.get('plugin_name') or ''}", heading_level)
        _paragraph(document, f"VPR: {item.get('vpr_score') if item.get('vpr_score') is not None else ''}")
        _labeled_blocks(document, "Descrição:", item.get("description") or item.get("synopsis") or "", translator=translator)
        _labeled_blocks(document, "Solução:", item.get("solution") or "", translator=translator)
        _heading(document, "Mais informações:", 4)
        for url in item.get("reference_urls") or []:
            _paragraph(document, str(url))
        _heading(document, "Host(s) Afetado(s):", 4)
        headers: tuple[str, ...] = ("ASSET NAME", "IP", "PORTA", protocol_header)
        widths: tuple[int, ...] = (2850, 2200, 1200, 1450)
        if include_output:
            headers += ("Output",)
            widths += (3000,)
        _simple_table(
            document,
            headers,
            _host_rows(item, mask_sensitive, include_output),
            widths=widths,
            left_columns=frozenset({0, 1, 4} if include_output else {0, 1}),
        )
        add_source_filter_note(
            document,
            dataset,
            source_table_id,
            enabled=show_source_filters,
            extra_filters={
                **dict(source_extra_filters or {}),
                "Plugin ID": item.get("plugin_id"),
            },
        )
    return rendered


def _web_vulnerability_details(
    document: DocxDocument,
    items: Any,
    *,
    dataset: Mapping[str, Any],
    show_source_filters: bool,
    mask_sensitive: bool,
    include_output: bool,
    translator: TextTranslator | None = None,
) -> int:
    if not isinstance(items, list):
        return 0
    rendered = 0
    for index, item in enumerate(items[:5], start=1):
        if not isinstance(item, Mapping):
            continue
        rendered += 1
        _heading(document, f"6.4.{index}. {item.get('plugin_name') or ''}", 3)
        _paragraph(document, f"VPR: {item.get('vpr_score') if item.get('vpr_score') is not None else ''}")
        _labeled_blocks(document, "Descrição:", item.get("description") or item.get("synopsis") or "", translator=translator)
        _labeled_blocks(document, "Solução:", item.get("solution") or "", translator=translator)
        _heading(document, "Mais informações:", 4)
        for url in item.get("reference_urls") or []:
            _paragraph(document, str(url))
        _heading(document, "URI(s) Afetada(s):", 4)
        headers: tuple[str, ...] = ("URI", "Plugin Output") if include_output else ("URI",)
        widths: tuple[int, ...] = (4500, 4500) if include_output else (9000,)
        rows: list[tuple[Any, ...]] = []
        instances = item.get("applications") or item.get("uris") or item.get("hosts") or []
        for instance in instances:
            if not isinstance(instance, Mapping):
                continue
            uri = "" if mask_sensitive else instance.get("uri") or instance.get("url") or ""
            rows.append((uri, instance.get("plugin_output") or "") if include_output else (uri,))
        _simple_table(document, headers, rows, widths=widths, left_columns=frozenset(range(len(headers))))
        add_source_filter_note(
            document,
            dataset,
            "top_web_vulnerabilities",
            enabled=show_source_filters,
            extra_filters={"Plugin ID": item.get("plugin_id")},
        )
    return rendered


def _was_section(document: DocxDocument, dataset: Mapping[str, Any], profile: ClientProfile, mask_sensitive: bool, translator: TextTranslator | None = None) -> None:
    was = dataset.get("was") if isinstance(dataset.get("was"), Mapping) else {}
    _heading(document, "SENSOR WAS")
    _paragraph(document, copy.WAS_SENSOR)
    _heading(document, "6.1. Saúde Global das aplicações", 2)
    _paragraph(document, copy.WAS_GLOBAL_HEALTH)
    _paragraph(document, copy.WAS_HEALTH_FACTORS)
    _heading(document, "6.2. WAS Aplicações Vulneráveis", 2)
    _paragraph(document, copy.WAS_APPS)
    _paragraph(document, copy.WAS_INFO_NOTE)
    app_rows = []
    for item in was.get("applications") or []:
        if isinstance(item, Mapping):
            app_rows.append((
                "" if mask_sensitive else item.get("uri") or "",
                _count(item.get("critical")), _count(item.get("high")),
                _count(item.get("medium")), _count(item.get("low")), _count(item.get("total")),
            ))
    _simple_table(
        document,
        ("URI", "Crítica", "Alta", "Média", "Baixa", "Total"),
        app_rows,
        widths=(3900, 1050, 1050, 1050, 1050, 1100),
        left_columns=frozenset({0}),
        header_fills=(base.BLUE, base.CRITICAL, base.HIGH, base.MEDIUM, base.LOW, base.BLUE),
    )
    add_source_filter_note(
        document,
        dataset,
        "was_applications",
        enabled=profile.presentation.show_source_filters,
    )
    _heading(document, "6.3. WAS Vulnerabilidades Baseadas em Plugin ID", 2)
    _paragraph(document, copy.WAS_PLUGINS)
    plugin_rows = _compact_rows(was.get("top_vulnerabilities") or [])
    _simple_table(
        document,
        ("Plugin Id", "Nome", "Família", "Severidade", "Total", "VPR"),
        plugin_rows,
        widths=(900, 3000, 2050, 1050, 850, 850),
        left_columns=frozenset({1, 2}),
    )
    add_source_filter_note(
        document,
        dataset,
        "was_top_vulnerabilities",
        enabled=profile.presentation.show_source_filters,
    )
    _heading(document, "6.3.1. OWASP top 10", 3)
    _paragraph(document, copy.OWASP_ORGANIZATION)
    _paragraph(document, copy.OWASP_TOP10)
    _paragraph(document, "Mais informações em: https://owasp.org/Top10/")
    _paragraph(document, copy.OWASP_TABLES)
    owasp = was.get("owasp") if isinstance(was.get("owasp"), Mapping) else {}
    for title, definition in copy.OWASP_CATEGORIES:
        _heading(document, title, 4)
        _paragraph(document, definition)
        category = title[:3]
        rows = []
        for item in owasp.get(category) or []:
            if isinstance(item, Mapping):
                rows.append((item.get("plugin_id") or "", item.get("name") or "", item.get("instances") or ""))
        if rows:
            _simple_table(document, ("Plugin Id", "Nome", "Instâncias"), rows, widths=(1200, 6100, 1400), left_columns=frozenset({1}))
            add_source_filter_note(
                document,
                dataset,
                "was_owasp",
                enabled=profile.presentation.show_source_filters,
                extra_filters={"OWASP 2021": category},
            )
        else:
            _paragraph(document, copy.OWASP_EMPTY_MONTH)
    _heading(document, "6.4. Vulnerabilidades WEB e Suas Correções e/ou Contramedidas Recomendadas", 2)
    _paragraph(document, copy.TOP5_WEB_INTRO)
    rendered_web_details = _web_vulnerability_details(
        document,
        dataset.get("top_web_vulnerabilities", was.get("top_vulnerabilities", [])),
        dataset=dataset,
        show_source_filters=profile.presentation.show_source_filters,
        mask_sensitive=mask_sensitive,
        include_output=profile.presentation.was_top5_include_output,
        translator=translator,
    )
    if rendered_web_details == 0:
        _paragraph(document, copy.TOP5_WEB_EMPTY_MONTH)


def _matrix_rows(value: Any, keys: Sequence[str]) -> list[tuple[Any, ...]]:
    rows = []
    if isinstance(value, list):
        for item in value:
            if isinstance(item, Mapping):
                rows.append(tuple(item.get(key, "") for key in keys))
    return rows


def _security_section(
    document: DocxDocument,
    metrics: Mapping[str, Any],
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    _heading(document, "INCREMENTANDO A SEGURANÇA E PROTEÇÃO DO AMBIENTE")
    _paragraph(document, copy.SECURITY_INCREMENT)
    _paragraph(document, copy.OS_COLUMNS)
    os_rows = _matrix_rows(
        (metrics.get("by_operating_system") or {}).get("rows") or [],
        ("operating_system", "mitigated", "non_mitigated", "exploitable", "patch_available_over_30_days"),
    )
    _simple_table(
        document,
        ("", "Mitigado", "Não Mitigado", "Explorável", "Patch disponível por mais de 30 dias"),
        os_rows,
        widths=(2450, 1350, 1650, 1400, 2800),
        left_columns=frozenset({0}),
    )
    add_source_filter_note(
        document, dataset, "by_operating_system", enabled=show_source_filters
    )
    _paragraph(document, copy.OS_GRAPH)
    _paragraph(document, copy.CVSS)
    cvss_rows = _matrix_rows(metrics.get("by_cvss") or [], ("label", "mitigated", "non_mitigated", "exploitable", "patch_available_over_30_days"))
    _simple_table(
        document,
        ("", "Mitigado", "Não Mitigado", "Explorável", "Patch disponível por mais de 30 dias"),
        cvss_rows,
        widths=(2450, 1350, 1650, 1400, 2800),
        left_columns=frozenset({0}),
    )
    add_source_filter_note(document, dataset, "by_cvss", enabled=show_source_filters)
    _paragraph(document, copy.CVSS_VPR_CORRELATION)
    _paragraph(document, copy.HEATMAP)
    heatmap_rows = _matrix_rows(metrics.get("cvss_vpr_matrix") or [], ("label", "low", "medium", "high", "critical"))
    _simple_table(
        document,
        ("", "Baixo (VPR 0.0-3.9)", "Médio (VPR 4.0-6.9)", "Alto (VPR7.0-8.9)", "Crítico (VPR 9.0-10)"),
        heatmap_rows,
        widths=(2150, 1900, 1900, 1900, 1900),
        left_columns=frozenset({0}),
    )
    add_source_filter_note(
        document, dataset, "cvss_vpr_matrix", enabled=show_source_filters
    )
    _paragraph(document, copy.VPR)
    vpr = metrics.get("vpr_rating") if isinstance(metrics.get("vpr_rating"), Mapping) else {}
    _simple_table(
        document,
        ("RATING 10.0 - 9.0", "RATING 8.9-7.0", "RATING 6.9-4.0", "RATING 3.9-0.1"),
        ((vpr.get("critical", ""), vpr.get("high", ""), vpr.get("medium", ""), vpr.get("low", "")),),
        widths=(2300, 2300, 2300, 2300),
    )
    add_source_filter_note(document, dataset, "vpr_rating", enabled=show_source_filters)


def _summary_section(
    document: DocxDocument,
    metrics: Mapping[str, Any],
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    _heading(document, "RESUMO DE VULNERABILIDADES")
    _paragraph(document, copy.SUMMARY_LIFECYCLE)
    _paragraph(document, copy.SUMMARY_COLUMNS)
    _heading(document, "8.1. Estado das Vulnerabilidades", 2)
    non = metrics.get("non_mitigated") or {}
    fixed = metrics.get("mitigated") or {}
    resurfaced = metrics.get("resurfaced") or {}
    rows = []
    for label, source, exploitable_key in (
        ("Novo", non.get("new_by_severity") or {}, "new_exploitable"),
        ("Ativo", non.get("by_severity") or {}, "exploitable"),
        ("Corrigido", fixed.get("by_severity") or {}, "exploitable"),
        ("Ressurgido", resurfaced.get("by_severity") or {}, "exploitable"),
    ):
        rows.append((label, _count((non if label in ("Novo", "Ativo") else fixed if label == "Corrigido" else resurfaced).get(exploitable_key)), _count(source.get("critical")), _count(source.get("high")), _count(source.get("medium"))))
    _simple_table(document, ("", "Explorável", "Crítica", "Alta", "Média"), rows, widths=(1700, 1800, 1700, 1700, 1700), left_columns=frozenset({0}))
    add_source_filter_note(
        document, dataset, "state_summary", enabled=show_source_filters
    )
    _paragraph(document, copy.AGING)
    _heading(document, "8.2. Idade das Vulnerabilidades", 2)
    aging = non.get("aging_by_severity") if isinstance(non.get("aging_by_severity"), Mapping) else {}
    aging_rows = []
    for label, key in (("Crítica", "critical"), ("Alta", "high"), ("Média", "medium"), ("Baixa", "low")):
        values = aging.get(key) if isinstance(aging.get(key), Mapping) else {}
        aging_rows.append((label, _count(values.get("90_plus_days")), _count(values.get("61_91_days")), _count(values.get("31_60_days")), _count(values.get("15_30_days")), _count(values.get("8_14_days")), _count(values.get("0_7_days"))))
    _simple_table(document, ("", "90+ Dias", "61-91 Dias", "31-60 Dias", "15-30 Dias", "8-14 Dias", "0-7 Dias"), aging_rows, widths=(1250, 1325, 1325, 1325, 1325, 1325, 1325), left_columns=frozenset({0}))
    add_source_filter_note(
        document, dataset, "aging_by_severity", enabled=show_source_filters
    )
    _paragraph(document, copy.FRAMEWORK)
    _heading(document, "8.3. Vulnerabilidades Exploráveis por Framework", 2)
    framework_rows = _matrix_rows(metrics.get("by_exploit_framework") or [], ("framework", "total", "critical", "high", "medium"))
    _simple_table(
        document,
        ("", "Total", "Crítica", "Alta", "Média"),
        framework_rows,
        widths=(2500, 1700, 1700, 1700, 1700),
        left_columns=frozenset({0}),
        empty_message=copy.EXPLOIT_FRAMEWORK_EMPTY_MONTH,
        keep_together=True,
    )
    if framework_rows:
        add_source_filter_note(
            document, dataset, "by_exploit_framework", enabled=show_source_filters
        )


def _back_cover(document: DocxDocument) -> None:
    document.add_page_break()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Cm(9)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(copy.BACK_COVER)
    base._set_run_font(run, size=20, color=base.NAVY, bold=True)


def _body(document: DocxDocument, dataset: Mapping[str, Any], profile: ClientProfile, mask_sensitive: bool, translator: TextTranslator | None = None) -> int:
    period = dataset["period"]
    start, end = _period_dates(period)
    generated = base._parse_utc(str(dataset.get("generated_at") or period["end_at"])).astimezone(ZoneInfo(str(period.get("timezone") or "UTC")))
    _heading(document, "SUMÁRIO")
    _toc_field(document)
    document.add_page_break()
    _control_document(document, generated.strftime("%d/%m/%Y"))
    _heading(document, "OBJETIVO")
    _paragraph(document, copy.OBJECTIVE)
    _period_paragraph(document, start, end)
    _heading(document, "SENSOR NESSUS, NESSUS AGENT E NESSUS NETWORK MONITOR")
    _paragraph(document, copy.NESSUS_SENSOR)
    _paragraph(document, copy.VULNERABLE_ENVIRONMENT)
    _overview_paragraph(document, start, end)
    for text in copy.OVERVIEW_COLUMN_TEXTS:
        _bullet(document, text)
    metrics = dataset.get("metrics") or {}
    _overview_table(
        document,
        metrics,
        dataset,
        show_source_filters=profile.presentation.show_source_filters,
    )
    _heading(document, "3.2. Principais Ativos Vulneráveis", 2)
    _paragraph(document, copy.TOP_ASSETS_INTRO)
    _paragraph(document, copy.TOP_ASSETS_PRIORITY)
    _top_assets_table(document, dataset.get("top_assets") or [], mask_sensitive)
    add_source_filter_note(
        document, dataset, "top_assets",
        enabled=profile.presentation.show_source_filters,
    )
    _principal_vulnerabilities(
        document, dataset,
        show_source_filters=profile.presentation.show_source_filters,
    )
    _heading(document, "VULNERABILIDADES E SUAS CORREÇÕES E/OU CONTRAMEDIDAS RECOMENDADAS")
    _paragraph(document, copy.TOP5_VM_INTRO)
    top_open_count = _vulnerability_details(
        document,
        dataset.get("top_open_vulnerabilities"),
        dataset=dataset,
        source_table_id="top_open_vulnerabilities",
        show_source_filters=profile.presentation.show_source_filters,
        heading_level=2,
        number_prefix="5",
        mask_sensitive=mask_sensitive,
        include_output=profile.presentation.vm_top5_include_output,
        translator=translator,
    )
    _was_section(document, dataset, profile, mask_sensitive, translator)
    _security_section(
        document,
        metrics,
        dataset,
        show_source_filters=profile.presentation.show_source_filters,
    )
    _summary_section(
        document,
        metrics,
        dataset,
        show_source_filters=profile.presentation.show_source_filters,
    )
    _back_cover(document)
    return top_open_count


def generate_full_base_report(
    *,
    template_path: str | Path,
    dataset_path: str | Path,
    profile: ClientProfile,
    output_path: str | Path,
    assets_dir: str | Path | None = None,
    mask_sensitive: bool = False,
    translator: TextTranslator | None = None,
) -> FullBaseReportRenderResult:
    """Gera o relatório-base usando apenas texto editorial dos DOCX de referência."""

    del assets_dir  # A identidade visual já está incorporada ao template aprovado.
    template = Path(template_path)
    dataset_file = Path(dataset_path)
    output = Path(output_path)
    if not template.is_file():
        raise ValueError(f"Template Word não encontrado: {template}")
    dataset = _load_dataset(dataset_file)
    _validate_dataset(dataset, profile)
    document = Document(template)
    _clear_body_after_cover_break(document)
    _configure_styles(document)
    period_label, period_range = base._period_labels(dataset["period"])
    base._replace_tokens(document, {
        "{{CLIENT_NAME}}": profile.display_name,
        "{{PERIOD_LABEL}}": period_label,
        "{{PERIOD_RANGE}}": period_range,
        "{{TEMPLATE_VERSION}}": FULL_TEMPLATE_VERSION,
    })
    _sanitize_header_footer(document, profile.display_name)
    _sanitize_properties(document, title=FULL_REPORT_TITLE)
    top_open_count = _body(document, dataset, profile, mask_sensitive, translator)
    base._enable_field_updates(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return FullBaseReportRenderResult(
        output_path=output,
        template_version=FULL_TEMPLATE_VERSION,
        client_id=profile.client_id,
        period_id=str(dataset["period"].get("period_id") or ""),
        top_asset_rows=len(dataset.get("top_assets") or []),
        top_open_rows=top_open_count,
        masked_sensitive_fields=mask_sensitive,
    )
