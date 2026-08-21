from __future__ import annotations

from dataclasses import dataclass
from copy import deepcopy
from pathlib import Path
import tempfile
from typing import Any, Mapping

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm

from tenable_reports.config.profile import ClientProfile
from tenable_reports.presentation import base_report_docx as base
from tenable_reports.presentation import editorial_catalog as copy
from tenable_reports.presentation.full_base_report_docx import (
    FULL_TEMPLATE_VERSION,
    _clear_body_after_cover_break,
    _compact_rows,
    _configure_styles,
    _heading,
    _load_dataset,
    _paragraph,
    _period_dates,
    _period_paragraph,
    _sanitize_header_footer,
    _sanitize_properties,
    _simple_table,
    _toc_field,
    _top_assets_table,
    _vulnerability_details,
)
from tenable_reports.presentation.source_filters import add_source_filter_note
from tenable_reports.presentation.translation import TextTranslator
from tenable_reports.presentation.monthly_visuals import (
    render_monthly_visual_bundle,
    render_tag_asset_comparison,
)


TAG_REPORT_TITLE = "RELATÓRIO DE VULNERABILIDADES TENABLE POR TAG"


@dataclass(frozen=True, slots=True)
class TagReportRenderResult:
    output_path: Path
    client_id: str
    period_id: str
    tag_uuid: str
    top_asset_rows: int
    top_open_rows: int
    comparison_rendered: bool
    masked_sensitive_fields: bool


def _tag_back_cover(document: Any) -> None:
    first_section = document.sections[0]
    blank_first_references = [
        deepcopy(reference)
        for reference_tag in ("w:headerReference", "w:footerReference")
        for reference in first_section._sectPr.findall(qn(reference_tag))
        if reference.get(qn("w:type")) == "first"
    ]
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.different_first_page_header_footer = True
    for reference in reversed(blank_first_references):
        section._sectPr.insert(0, reference)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Cm(9)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(copy.BACK_COVER)
    base._set_run_font(run, size=20, color=base.NAVY, bold=True)


def _validate_tag_dataset(
    dataset: Mapping[str, Any], profile: ClientProfile
) -> Mapping[str, Any]:
    if dataset.get("client_id") != profile.client_id:
        raise ValueError("O dataset por TAG nao pertence ao cliente selecionado.")
    if dataset.get("document_kind") != "tag":
        raise ValueError("O dataset selecionado nao e um dataset por TAG.")
    if not isinstance(dataset.get("period"), Mapping):
        raise ValueError("Dataset por TAG sem periodo valido.")
    tag = dataset.get("tag")
    if not isinstance(tag, Mapping):
        raise ValueError("Dataset por TAG sem identificacao valida.")
    for field in ("tag_uuid", "category_name", "value"):
        if not str(tag.get(field) or "").strip():
            raise ValueError(f"Dataset por TAG sem {field}.")
    return tag


def _tag_filters(tag: Mapping[str, Any]) -> dict[str, str]:
    return {
        "Tag UUID": str(tag.get("tag_uuid") or ""),
        "Tag": (
            f"{str(tag.get('category_name') or '')}:"
            f"{str(tag.get('value') or '')}"
        ),
    }


def _tag_body(
    document: Any,
    dataset: Mapping[str, Any],
    profile: ClientProfile,
    tag: Mapping[str, Any],
    *,
    mask_sensitive: bool,
    translator: TextTranslator | None,
) -> int:
    start, end = _period_dates(dataset["period"])
    tag_label = f"TAG {tag['category_name']} - {tag['value']}"
    filters = _tag_filters(tag)

    _heading(document, "SUMÁRIO")
    _toc_field(document)
    document.add_page_break()
    _heading(document, tag_label)
    _period_paragraph(document, start, end)

    _heading(document, "3.2. Principais Ativos Vulneráveis", 2)
    _paragraph(document, copy.TOP_ASSETS_INTRO)
    _paragraph(document, copy.TOP_ASSETS_PRIORITY)
    top_assets = dataset.get("top_assets") or []
    if top_assets:
        _top_assets_table(document, top_assets, mask_sensitive)
        add_source_filter_note(
            document,
            dataset,
            "top_assets",
            enabled=profile.presentation.show_source_filters,
            extra_filters=filters,
        )
    else:
        _paragraph(
            document,
            "Neste mês não foram identificados ativos vulneráveis para esta TAG.",
        )

    _heading(document, "VISÃO GERAL DAS PRINCIPAIS VULNERABILIDADES")
    _paragraph(document, copy.PRINCIPAL_VULNERABILITIES_INTRO)
    _heading(document, "4.2. Vulnerabilidades Não Mitigadas", 2)
    top_open = dataset.get("top_open_vulnerabilities") or []
    if top_open:
        _simple_table(
            document,
            ("Plugin ID", "Nome", "Família OS", "Severidade", "Total", "VPR"),
            _compact_rows(top_open),
            widths=(900, 3000, 2050, 1050, 850, 850),
            left_columns=frozenset({1, 2}),
        )
        add_source_filter_note(
            document,
            dataset,
            "top_open_vulnerabilities",
            enabled=profile.presentation.show_source_filters,
            extra_filters=filters,
        )
    else:
        _paragraph(
            document,
            "Neste mês não foram identificadas vulnerabilidades não mitigadas "
            "para esta TAG.",
        )

    _heading(
        document,
        "VULNERABILIDADES E SUAS CORREÇÕES E/OU CONTRAMEDIDAS RECOMENDADAS",
    )
    _paragraph(document, copy.TOP5_VM_INTRO)
    if not top_open:
        _paragraph(
            document,
            "Neste mês não foram identificadas vulnerabilidades não mitigadas "
            "para detalhamento nesta TAG.",
        )
        return 0
    return _vulnerability_details(
        document,
        top_open,
        dataset=dataset,
        source_table_id="top_open_vulnerabilities",
        show_source_filters=profile.presentation.show_source_filters,
        heading_level=2,
        number_prefix="5",
        mask_sensitive=mask_sensitive,
        include_output=profile.presentation.vm_top5_include_output,
        translator=translator,
        protocol_header="PROTOCOLO",
        source_extra_filters=filters,
    )


def _temporal_comparison(
    document: Any,
    dataset: Mapping[str, Any],
    tag: Mapping[str, Any],
    *,
    mask_sensitive: bool,
) -> bool:
    if not bool(tag.get("include_temporal_comparison", False)):
        return False
    status = str(dataset.get("tag_history_status") or "")
    rows = [
        item
        for item in dataset.get("tag_history") or ()
        if isinstance(item, Mapping)
    ]
    _heading(document, "Comparativo Mensal da TAG")
    if status == "INCOMPATIBLE_PERIOD":
        _paragraph(
            document,
            "O período deste relatório não corresponde a um mês civil completo; "
            "a série mensal não é apresentada.",
        )
        return False
    if status != "AVAILABLE" or not rows:
        _paragraph(
            document,
            "Ainda não há histórico mensal comparável para esta TAG.",
        )
        return False
    scope_label = f"TAG {tag['category_name']} - {tag['value']}"
    with tempfile.TemporaryDirectory() as directory:
        visual_count = render_monthly_visual_bundle(
            document,
            rows,
            directory,
            scope_label,
        )
        comparison = dataset.get("tag_comparison")
        if isinstance(comparison, Mapping):
            render_tag_asset_comparison(
                document,
                comparison,
                mask_sensitive=mask_sensitive,
            )
    return visual_count == 5


def generate_tag_report(
    *,
    template_path: str | Path,
    dataset_path: str | Path,
    profile: ClientProfile,
    output_path: str | Path,
    mask_sensitive: bool = False,
    translator: TextTranslator | None = None,
) -> TagReportRenderResult:
    template = Path(template_path)
    if not template.is_file():
        raise ValueError(f"Template Word não encontrado: {template}")
    dataset = _load_dataset(Path(dataset_path))
    tag = _validate_tag_dataset(dataset, profile)
    document = Document(template)
    # The corporate template has an empty even-page header.  Using the same
    # default header on both sides is also more reliable in LibreOffice when a
    # large table or chart is carried to the next page.
    document.settings.odd_and_even_pages_header_footer = False
    _clear_body_after_cover_break(document)
    _configure_styles(document)
    period_label, period_range = base._period_labels(dataset["period"])
    tag_label = f"TAG {tag['category_name']} - {tag['value']}"
    base._replace_tokens(
        document,
        {
            "{{CLIENT_NAME}}": profile.display_name,
            "{{PERIOD_LABEL}}": f"{period_label}\n{tag_label}",
            "{{PERIOD_RANGE}}": period_range,
            "{{TEMPLATE_VERSION}}": FULL_TEMPLATE_VERSION,
        },
    )
    _sanitize_header_footer(document, profile.display_name)
    _sanitize_properties(document, title=TAG_REPORT_TITLE)
    top_open_rows = _tag_body(
        document,
        dataset,
        profile,
        tag,
        mask_sensitive=mask_sensitive,
        translator=translator,
    )
    comparison_rendered = _temporal_comparison(
        document,
        dataset,
        tag,
        mask_sensitive=mask_sensitive,
    )
    _tag_back_cover(document)
    base._enable_field_updates(document)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return TagReportRenderResult(
        output_path=output,
        client_id=profile.client_id,
        period_id=str(dataset["period"].get("period_id") or ""),
        tag_uuid=str(tag["tag_uuid"]),
        top_asset_rows=len(dataset.get("top_assets") or []),
        top_open_rows=top_open_rows,
        comparison_rendered=comparison_rendered,
        masked_sensitive_fields=mask_sensitive,
    )
