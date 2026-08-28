from __future__ import annotations

import hashlib
import re
import tempfile
from dataclasses import dataclass
from datetime import datetime
from enum import StrEnum
from pathlib import Path
from typing import Any, Mapping, Sequence
from zoneinfo import ZoneInfo

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from tenable_reports.application.cloud_report_dataset import (
    load_cloud_report_dataset,
)
from tenable_reports.config.profile import ClientProfile
from tenable_reports.presentation.cloud_report_sections import (
    CloudDocumentBuilder,
    cloud_posture_available,
    render_container_image_vulnerability_overview,
    render_cloud_overview,
    render_cloud_posture,
    render_components_products,
    render_conclusion,
    render_critical_details,
    render_dashboard,
    render_document_control,
    render_executive_overview,
    render_introduction,
    render_monthly_evolution,
    render_remediation_performance,
    render_top_correctable,
    render_top_critical,
    render_top_hosts,
    render_top_images,
    render_vulnerability_aging,
)
from tenable_reports.presentation.translation import TextTranslator


CLOUD_TEMPLATE_VERSION = "cloud-base-v1.0"
STANDARD_SECTION_IDS = (
    "cover",
    "table_of_contents",
    "document_control",
    "objective",
    "cloud_overview",
    "introduction",
    "executive_overview",
    "top_hosts",
    "top_images",
    "container_image_vulnerability_overview",
    "top_critical",
    "critical_details",
    "top_correctable",
    "dashboard",
    "components_products",
    "cloud_posture",
    "vulnerability_aging",
    "remediation_performance",
    "monthly_evolution",
    "conclusion",
    "back_cover",
)
_MONTHS_PT = (
    "",
    "JANEIRO",
    "FEVEREIRO",
    "MARÇO",
    "ABRIL",
    "MAIO",
    "JUNHO",
    "JULHO",
    "AGOSTO",
    "SETEMBRO",
    "OUTUBRO",
    "NOVEMBRO",
    "DEZEMBRO",
)


class CloudReportVariant(StrEnum):
    EXPANDED = "expanded"


@dataclass(frozen=True, slots=True)
class CloudReportRenderResult:
    output_path: Path
    client_id: str
    period_id: str
    variant: CloudReportVariant
    rendered_sections: Sequence[str]
    omitted_sections: Sequence[str]
    dataset_sha256: str


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _period_id(dataset: Mapping[str, Any]) -> str:
    period = dataset.get("period") or {}
    value = str(period.get("period_id") or "").strip()
    if re.fullmatch(r"\d{4}-\d{2}", value):
        return value
    start = datetime.fromisoformat(str(period["start_at"]).replace("Z", "+00:00"))
    timezone = ZoneInfo(str(period.get("timezone") or "UTC"))
    return start.astimezone(timezone).strftime("%Y-%m")


def _month_year(dataset: Mapping[str, Any]) -> str:
    period = dataset.get("period") or {}
    timezone = ZoneInfo(str(period.get("timezone") or "UTC"))
    start = datetime.fromisoformat(str(period["start_at"]).replace("Z", "+00:00"))
    local = start.astimezone(timezone)
    return f"{_MONTHS_PT[local.month]}/{local.year}"


def _replace_text_nodes(element: Any, replacements: Mapping[str, str]) -> None:
    for node in element.iter():
        if node.tag.rsplit("}", 1)[-1] != "t" or node.text is None:
            continue
        value = node.text
        for marker, replacement in replacements.items():
            value = value.replace(marker, replacement)
        node.text = value


def _replace_template_markers(
    document: DocxDocument,
    *,
    client_name: str,
    month_year: str,
) -> None:
    replacements = {
        "{{CLIENT_NAME}}": client_name,
        "{{REPORT_MONTH_YEAR}}": month_year,
    }
    _replace_text_nodes(document._element, replacements)
    for section in document.sections:
        for container in (section.header, section.footer):
            _replace_text_nodes(container._element, replacements)


def _find_paragraph(document: DocxDocument, marker: str):
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            return paragraph
    raise ValueError(f"Marcador obrigatório ausente no template Cloud: {marker}")


def _set_toc_field(
    paragraph: Any,
    *,
    entries: Sequence[str],
) -> None:
    paragraph.clear()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "Sumário automático; a paginação é atualizada ao abrir o documento no Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, value, end):
        run._r.append(element)

    for entry in entries:
        item = paragraph.add_run()
        item.add_break()
        item.add_text(entry)


def _toc_entries(*, include_posture: bool) -> tuple[str, ...]:
    entries = [
        "1. CONTROLE DE DOCUMENTO",
        "2. OBJETIVO",
        "3. TENABLE CLOUD SECURITY",
        "3.1. Introdução",
        "3.1.1. Resumo Executivo do Período",
        "3.2. Principais Hosts Vulneráveis",
        "3.3. Imagens de Contêineres Mais Vulneráveis",
        "3.3.1. Overview das Vulnerabilidades das Imagens de Contêiner",
        "3.4. Principais Vulnerabilidades Críticas (TOP 5 CVEs)",
        "3.5. Principais Vulnerabilidades com Correção Disponível",
        "3.6. Painel de Controle (Dashboards)",
        "3.7. Componentes e Produtos em Maior Risco",
    ]
    if include_posture:
        entries.append("3.8. Postura de Segurança em Nuvem")
    entries.extend(
        (
            "3.9. Envelhecimento das Vulnerabilidades",
            "3.10. Desempenho de Remediação",
            "3.11. Evolução Mensal",
            "4. Conclusão",
        )
    )
    return tuple(entries)

def _set_update_fields(document: DocxDocument) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _configure_heading_styles(document: DocxDocument) -> None:
    for level, size in ((1, 15), (2, 12.5), (3, 11.5), (4, 10.5)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = None
        style.paragraph_format.keep_with_next = True


def _sanitize_properties(document: DocxDocument, profile: ClientProfile) -> None:
    properties = document.core_properties
    properties.author = "ITProtect"
    properties.last_modified_by = "ITProtect"
    properties.title = "Relatório Tenable Cloud Security"
    properties.subject = f"Relatório mensal Cloud Security - {profile.display_name}"
    properties.keywords = "Tenable Cloud Security"
    properties.comments = ""


def generate_cloud_report(
    *,
    template_path: str | Path,
    dataset_path: str | Path,
    profile: ClientProfile,
    output_path: str | Path,
    variant: CloudReportVariant | str = CloudReportVariant.EXPANDED,
    translator: TextTranslator | None = None,
) -> CloudReportRenderResult:
    template = Path(template_path)
    dataset_source = Path(dataset_path)
    output = Path(output_path)
    selected_variant = CloudReportVariant(variant)
    if not template.is_file():
        raise FileNotFoundError(f"Template Cloud não encontrado: {template}")
    dataset = load_cloud_report_dataset(dataset_source)

    document = Document(template)
    _configure_heading_styles(document)
    _replace_template_markers(
        document,
        client_name=profile.display_name,
        month_year=_month_year(dataset),
    )
    include_posture = cloud_posture_available(dataset)
    toc = _find_paragraph(document, "{{TABLE_OF_CONTENTS}}")
    _set_toc_field(
        toc,
        entries=_toc_entries(include_posture=include_posture),
    )
    anchor = _find_paragraph(document, "{{CLOUD_CONTENT_START}}")
    builder = CloudDocumentBuilder(document=document, anchor=anchor)
    rendered_sections = list(STANDARD_SECTION_IDS)
    omitted_sections: list[str] = []

    with tempfile.TemporaryDirectory(prefix="cloud-report-visuals-") as chart_directory:
        chart_dir = Path(chart_directory)
        render_document_control(builder, dataset)
        render_cloud_overview(builder)
        warning = str(
            (dataset.get("snapshot_context") or {}).get("warning") or ""
        ).strip()
        if warning:
            builder.paragraph(warning, bold=True, color="C00000")
        render_introduction(builder)
        show_filters = profile.presentation.show_source_filters
        builder.page_break()
        render_executive_overview(
            builder,
            dataset,
            chart_dir=chart_dir,
        )
        render_top_hosts(builder, dataset, show_source_filters=show_filters)
        render_top_images(builder, dataset, show_source_filters=show_filters)
        render_container_image_vulnerability_overview(
            builder,
            dataset,
            show_source_filters=show_filters,
        )
        render_top_critical(builder, dataset, show_source_filters=show_filters)
        render_critical_details(builder, dataset, translator=translator)
        render_top_correctable(
            builder,
            dataset,
            show_source_filters=show_filters,
            translator=translator,
        )
        render_dashboard(builder, dataset)
        render_components_products(
            builder,
            dataset,
            show_source_filters=show_filters,
        )
        if include_posture:
            render_cloud_posture(
                builder,
                dataset,
                show_source_filters=show_filters,
            )
        else:
            rendered_sections.remove("cloud_posture")
            omitted_sections.append("cloud_posture")
        render_vulnerability_aging(
            builder,
            dataset,
            chart_dir=chart_dir,
            show_source_filters=show_filters,
        )
        render_remediation_performance(
            builder,
            dataset,
            show_source_filters=show_filters,
        )
        history_chart_rendered = render_monthly_evolution(
            builder,
            dataset,
            chart_dir=chart_dir,
        )
        if not history_chart_rendered:
            omitted_sections.append("monthly_evolution_chart")
        builder.page_break()
        render_conclusion(builder)

        anchor.clear()
        _set_update_fields(document)
        _sanitize_properties(document, profile)
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)

    return CloudReportRenderResult(
        output_path=output,
        client_id=profile.client_id,
        period_id=_period_id(dataset),
        variant=selected_variant,
        rendered_sections=tuple(rendered_sections),
        omitted_sections=tuple(omitted_sections),
        dataset_sha256=_sha256(dataset_source),
    )

__all__ = [
    "CLOUD_TEMPLATE_VERSION",
    "STANDARD_SECTION_IDS",
    "CloudReportRenderResult",
    "CloudReportVariant",
    "generate_cloud_report",
]
