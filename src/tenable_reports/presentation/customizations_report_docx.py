from __future__ import annotations

import json
import tempfile
from dataclasses import dataclass
from copy import deepcopy
from pathlib import Path
from typing import Any, Mapping, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

from tenable_reports.config.profile import ClientProfile
from tenable_reports.presentation import base_report_docx as base
from tenable_reports.presentation import editorial_catalog as editorial
from tenable_reports.presentation import full_base_report_docx as faithful
from tenable_reports.presentation.source_filters import add_source_filter_note


CUSTOM_TEMPLATE_VERSION = "customizacoes-fieis-v1.0"
HISTORICAL_INTELLIGENCE_MODULES = frozenset({
    "vm_monthly_volume",
    "vm_previous_period_delta",
    "vm_network_comparison",
    "vm_executive_evolution",
    "vm_monthly_evolution",
})

NO_HISTORY_MESSAGE = "Não há histórico do período imediatamente anterior para comparação."
NO_DATA_MESSAGES = {
    "scan_auth_health": "Dados indisponíveis para este indicador.",
    "vm_plugin_family": "Neste mês não foram identificadas vulnerabilidades mitigadas para agrupamento por família de plugin.",
    "vm_eol_software": "Neste mês não foram identificados sistemas ou softwares sem suporte.",
    "vm_exploit_vector": "Neste mês não foram identificadas vulnerabilidades exploráveis com vetor de ataque classificável.",
    "was_unsupported_tech": "Neste mês não foram identificadas tecnologias WEB sem suporte.",
}

MONTHLY_COMPARISON = (
    "A segurança da informação é uma prioridade essencial para qualquer organização "
    "que deseja proteger seus dados e manter a integridade dos seus sistemas. Para "
    "isso, é vital realizar o acompanhamento contínuo das vulnerabilidades presentes "
    "no ambiente de TI. A seguir, apresentamos uma análise comparativa mensal das "
    "vulnerabilidades mitigadas e não mitigadas, utilizando gráficos para ilustrar "
    "os dados acima e destacando os benefícios dessa abordagem."
)

SCAN_HEALTH = (
    "O widget “Integridade da varredura” fornece um resumo da integridade da varredura "
    "em relação ao sucesso e às falhas de autenticação. Os dispositivos do sistema e "
    "da rede devem ser verificados rotineiramente para garantir que estejam operando "
    "em conformidade com os requisitos organizacionais e normativos para gerenciamento "
    "de vulnerabilidades e configurações."
)

PLUGIN_FAMILY = (
    "A tabela abaixo apresenta um resumo das vulnerabilidades corrigidas nos últimos "
    "30 dias, agrupadas por Família de Plugin, o filtro aplicado foi “Last Fixed: "
    "últimos 30 dias”, o que significa que estão sendo consideradas apenas as "
    "vulnerabilidades que foram efetivamente resolvidas nesse período. Cada linha "
    "representa uma família de plugins, que indica a origem ou o tipo de verificação "
    "utilizada para detectar as vulnerabilidades. O campo “Total” exibe a quantidade "
    "de vulnerabilidades corrigidas para cada família."
)

EOL_INTRO = (
    "A proliferação de produtos sem suporte e em fim de vida (EOL) é um problema de "
    "segurança comum enfrentado por todas as organizações. À medida que os aplicativos "
    "e sistemas operacionais atingem o EOL, os fornecedores param de oferecer suporte, "
    "fazendo com que a segurança e a estabilidade diminuam com o tempo. A identificação "
    "de ativos que executam aplicativos EOL é uma parte importante da avaliação e "
    "minimização do risco organizacional, uma vez que patches, atualizações e correções "
    "de segurança se tornam menos disponíveis. O Center for Internet Security (CIS) "
    "Critical Controls (Sub-control 2.2) afirma que as organizações devem garantir que "
    "apenas os aplicativos de software ou sistemas operacionais que são atualmente "
    "suportados e que recebem atualizações do fornecedor sejam adicionados ao inventário "
    "de software autorizado da organização. As organizações precisam marcar todos os "
    "softwares não suportados no inventário de ativos."
)

EOL_TENABLE = (
    "O Tenable.io permite que as organizações avaliem continuamente a postura de "
    "integridade e segurança da rede, incluindo a identificação e o monitoramento de "
    "softwares sem suporte. A identificação rápida de sistemas operacionais e aplicativos "
    "não suportados permite que os gerentes de risco vejam os riscos associados ao software "
    "EOL. A identificação de exposições fornece orientação às equipes de operações para "
    "implementar, agir e priorizar esforços de remediação para mitigar o risco cibernético. "
    "Os gerentes de risco e as equipes de operações podem comunicar à equipe de liderança "
    "como a atualização de sistemas operacionais e aplicativos não suportados reduz o risco "
    "de rede."
)

EOL_METHOD = (
    "O Tenable.io usa métodos ativos para identificar produtos EOL encontrados no ambiente, "
    "examinando o registro da Microsoft, locais comuns de instalação de software ou usando "
    "utilitários de aplicativos como YUM ou APT em sistemas Linux. Os gerentes de risco são "
    "capazes de verificar as atividades das equipes de operação e identificar áreas para "
    "mitigação de riscos."
)

EOL_ASSETS = (
    "Abaixo será apresentada uma lista dos 20 (vinte) principais Assets, exibindo os ativos "
    "que atualmente possuem a maior contagem de softwares e SOs não suportados no ambiente. "
    "Filtrado para mostrar apenas plug-ins ativos, ressurgidos ou novos, por severidade. "
    "Fornecendo à organização um método claro e simplificado para identificar softwares e "
    "SOs em EOL e permite que os gerentes de seguranças prevejam onde os riscos aumentam, e "
    "desenvolvam um plano de mitigação."
)

EOL_SOFTWARE = (
    "Abaixo será exibido todos os SOs e Softwares, que foram detectados no ambiente, como "
    "sistema operacional não suportado, já que saber quais sistemas operacionais não "
    "suportados ou se estão próximos do fim da vida útil (EOL,) pode melhorar a capacidade "
    "de uma equipe de segurança de mitigar vulnerabilidades e proteger a rede. Os ativos que "
    "executam sistemas operacionais não suportados são mais vulneráveis à exploração. "
    "Identificar e atualizar esses SOs é essencial para um programa de segurança eficaz."
)

EOL_PROTOCOLS = (
    "E exibir informações sobre softwares ou protocolos não suportados. O Tenable.io possui "
    "vários plugins que relatam softwares e configurações não suportados, que não estão "
    "relacionados a um fornecedor de software específico e identificam outliers não "
    "suportados. Essas vulnerabilidades identificadas são frequentemente relacionadas a "
    "Linux/Unix, OpenSSL e serviços web. Esses plugins fornecem a capacidade de agrupar "
    "ativos com base em um software ou protocolo não suportado, ordenados por maior "
    "quantidade de ativos afetados."
)

EXECUTIVE_PARAGRAPHS = (
    "Abaixo o gráfico apresenta uma visão consolidada da evolução das vulnerabilidades identificadas no ambiente ao longo dos períodos analisados, permitindo acompanhar tendências de crescimento, redução e surgimento de novas ocorrências de forma objetiva e visualmente priorizada.",
    "A classificação por cores foi adotada para facilitar a interpretação executiva e o direcionamento das ações de tratamento. Os itens destacados em vermelho representam as vulnerabilidades com maior crescimento no período, exigindo maior atenção devido ao potencial aumento de exposição do ambiente. Em laranja, estão vulnerabilidades com crescimento relevante, porém com impacto inferior ao grupo crítico. Os itens em azul representam vulnerabilidades com comportamento considerado dentro do padrão observado durante a análise.",
    "Adicionalmente, as vulnerabilidades destacadas em roxo correspondem a ocorrências novas identificadas no ciclo atual (“NEW”), indicando itens que não estavam presentes anteriormente. Esse comportamento pode estar associado à inclusão de novos ativos, mudanças no ambiente ou novas detecções realizadas durante o processo de monitoramento.",
    "Os valores apresentados no gráfico representam a variação da quantidade de ocorrências entre os períodos comparados. Valores positivos indicam aumento na incidência da vulnerabilidade, enquanto valores negativos representam redução da exposição identificada anteriormente, normalmente associada à aplicação de correções, atualizações, descontinuação de ativos vulneráveis ou ações de mitigação implementadas no ambiente.",
    "A análise contínua dessas variações permite acompanhar a evolução da postura de segurança do ambiente, apoiar a priorização técnica das ações corretivas e fornecer maior previsibilidade sobre os riscos operacionais e de segurança relacionados aos ativos monitorados.",
)

EVOLUTION_PARAGRAPHS = (
    "Este indicador apresenta a evolução mensal do risco de vulnerabilidades, correlacionando o estoque herdado do período anterior, o volume efetivamente mitigado no mês e o estoque remanescente ao final do período. O objetivo é evidenciar não apenas o esforço de remediação, mas principalmente o volume líquido de vulnerabilidades surgidas no período, refletindo a dinâmica real de exposição do ambiente.",
    "A métrica parte do pressuposto de que, caso não houvesse surgimento de novas vulnerabilidades, o estoque do mês corrente corresponderia ao total não mitigado do mês anterior menos as vulnerabilidades mitigadas no período. A diferença entre esse valor esperado e o estoque real observado representa as vulnerabilidades que ingressaram no ambiente durante o mês, seja por novas descobertas, atualizações de software, inclusão de ativos ou ampliação da cobertura de varredura.",
    "Fórmula de Cálculo",
    "Vulnerabilidades Surgidasₘ = NMₘ − (NMₘ₋₁ − Mₘ)",
    "Onde:",
    "NMₘ₋₁ = Vulnerabilidades não mitigadas no mês anterior",
    "Mₘ = Vulnerabilidades mitigadas no mês atual",
    "NMₘ = Vulnerabilidades não mitigadas no mês atual",
    "Esse indicador permite acompanhar tendências, avaliar a efetividade do programa de gestão de vulnerabilidades e apoiar decisões estratégicas voltadas à redução sustentável do risco.",
)

CONTAINER_INTRO = (
    "O Container Security Scanner permite fazer o escaneamento de imagens em contêineres "
    "com segurança, fazendo inventário inicial, ou instantâneo. Nessas imagens que foram "
    "feitas as checagens, e consequentemente esse inventário é enviado para o Tenable.io "
    "para posterior análise."
)

CONTAINER_CURRENT = (
    "Atualmente no ambiente foi feito a verificação de imagens, segue a lista das que "
    "estão sendo analisadas e que são críticas."
)

ATTACK_VECTOR = (
    "O widget Explorabilidade por Vetor de Ataque (Explore) abaixo exibe três colunas de "
    "vulnerabilidades exploráveis com base nos vetores de métrica de explorabilidade do "
    "CVSS: AV:N (Rede), AV:A (Rede Adjacente) e AV:L (Local).É fornecida uma linha para "
    "cada framework de exploração. O vetor de métrica do CVSS especifica o caminho que "
    "pode ser utilizado para explorar a vulnerabilidade. Vulnerabilidades que podem ser "
    "exploradas remotamente representam um risco maior, pois existe um vetor de ameaça "
    "global. Vulnerabilidades que só podem ser exploradas localmente requerem acesso "
    "local ao sistema por meio de outro mecanismo, como um usuário autorizado ou uma "
    "vulnerabilidade de execução remota."
)

WAS_UNSUPPORTED = (
    "A tabela abaixo apresenta as principais vulnerabilidades sem suporte do fabricante, "
    "detalhando a criticidade e os riscos associados a cada uma delas, por severidade e "
    "VPR. É imperativo que essas vulnerabilidades sejam tratadas com urgência, seja "
    "através de atualizações para versões suportadas, a implementação de soluções "
    "alternativas de segurança, ou a descontinuação dos sistemas afetados."
)


@dataclass(frozen=True, slots=True)
class CustomizationsRenderResult:
    output_path: Path
    template_version: str
    client_id: str
    period_id: str
    requested_modules: tuple[str, ...]
    rendered_modules: tuple[str, ...]
    omitted_modules: tuple[dict[str, str], ...]


def _font(size: int, bold: bool = False) -> Any:
    candidates = (
        Path(r"C:\Windows\Fonts\calibrib.ttf") if bold else Path(r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf") if bold else Path(r"C:\Windows\Fonts\arial.ttf"),
    )
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _bar_chart(path: Path, title: str, rows: Sequence[Mapping[str, Any]], keys: Sequence[tuple[str, str, str]]) -> None:
    width, height = 1400, max(650, 220 + len(rows) * 95)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 38), title, font=_font(34, True), fill="#0B1F4A")
    max_value = max((int(row.get(key, 0) or 0) for row in rows for key, _, _ in keys), default=1)
    y = 130
    for row in rows:
        label = str(row.get("label") or row.get("month") or "")
        draw.text((50, y + 16), label, font=_font(20, True), fill="#0B1F4A")
        x = 310
        for key, caption, color in keys:
            value = int(row.get(key, 0) or 0)
            bar = int(780 * value / max_value) if max_value else 0
            draw.rectangle((x, y, x + bar, y + 25), fill=color)
            draw.text((x + bar + 12, y), f"{caption}: {value}", font=_font(17), fill="#0B1F4A")
            y += 32
        y += 25
    image.save(path)


SEVERITY_SERIES = (
    ("critical", "Crítica", "#FF0000"),
    ("high", "Alta", "#ED7D31"),
    ("medium", "Média", "#FFF200"),
    ("low", "Baixa", "#70AD47"),
)


def _number(value: Any) -> int:
    try:
        return int(value or 0)
    except (TypeError, ValueError):
        return 0


def _format_number(value: int) -> str:
    return f"{value:,}".replace(",", ".")


def _monthly_views(data: Mapping[str, Any]) -> list[tuple[str, list[Mapping[str, Any]]]]:
    views: list[tuple[str, list[Mapping[str, Any]]]] = []
    configured = data.get("monthly_views")
    if isinstance(configured, list):
        for view in configured:
            if not isinstance(view, Mapping):
                continue
            history = view.get("history")
            if isinstance(history, list):
                rows = [row for row in history if isinstance(row, Mapping)]
                if rows:
                    views.append((str(view.get("label") or "Geral"), rows))
    if views:
        return views
    history = data.get("monthly_history")
    if isinstance(history, list):
        rows = [row for row in history if isinstance(row, Mapping)]
        if rows:
            views.append(("Geral", rows))
    return views


def _severity_rows(
    history: Sequence[Mapping[str, Any]],
    *,
    nested_key: str,
    total_key: str,
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for item in history:
        nested = item.get(nested_key)
        if not isinstance(nested, Mapping):
            return []
        row: dict[str, Any] = {
            "label": item.get("label") or item.get("month") or "",
            "total": _number(item.get(total_key)),
        }
        for key, _, _ in SEVERITY_SERIES:
            row[key] = _number(nested.get(key))
        rows.append(row)
    return rows


def _grouped_monthly_chart(
    path: Path,
    title: str,
    rows: Sequence[Mapping[str, Any]],
    series: Sequence[tuple[str, str, str]],
) -> None:
    width, height = 1500, 900
    background = "#292929"
    image = Image.new("RGB", (width, height), background)
    draw = ImageDraw.Draw(image)
    title_font = _font(38, True)
    label_font = _font(19, True)
    small_font = _font(16)
    draw.text((width // 2, 38), title.upper(), font=title_font, fill="white", anchor="ma")
    left, top, right, bottom = 105, 150, width - 55, height - 205
    values = [_number(row.get(key)) for row in rows for key, _, _ in series]
    max_value = max(values, default=1) or 1
    for tick in range(6):
        value = round(max_value * tick / 5)
        y = bottom - int((bottom - top) * tick / 5)
        draw.line((left, y, right, y), fill="#595959", width=1)
        draw.text((left - 18, y), _format_number(value), font=small_font, fill="#E7E7E7", anchor="rm")
    group_width = (right - left) / max(len(rows), 1)
    bar_gap = 5
    bar_width = max(12, min(42, int((group_width * 0.78) / max(len(series), 1)) - bar_gap))
    for row_index, row in enumerate(rows):
        group_center = left + group_width * (row_index + 0.5)
        total_bars_width = len(series) * bar_width + (len(series) - 1) * bar_gap
        start_x = int(group_center - total_bars_width / 2)
        for series_index, (key, _, color) in enumerate(series):
            value = _number(row.get(key))
            bar_height = int((bottom - top) * value / max_value)
            x1 = start_x + series_index * (bar_width + bar_gap)
            y1 = bottom - bar_height
            draw.rectangle((x1, y1, x1 + bar_width, bottom), fill=color)
            if value:
                draw.text(
                    (x1 + bar_width // 2, max(top - 8, y1 - 8)),
                    _format_number(value),
                    font=small_font,
                    fill="white",
                    anchor="ms",
                )
        draw.text(
            (group_center, bottom + 22),
            str(row.get("label") or ""),
            font=label_font,
            fill="white",
            anchor="ma",
        )
    legend_y = height - 92
    legend_width = sum(42 + draw.textlength(label, font=label_font) + 28 for _, label, _ in series)
    legend_x = int((width - legend_width) / 2)
    for _, label, color in series:
        draw.rectangle((legend_x, legend_y, legend_x + 24, legend_y + 24), fill=color)
        draw.text((legend_x + 34, legend_y + 12), label, font=label_font, fill="white", anchor="lm")
        legend_x += int(42 + draw.textlength(label, font=label_font) + 28)
    image.save(path)


def _monthly_line_chart(
    path: Path,
    title: str,
    history: Sequence[Mapping[str, Any]],
    *,
    value_key: str,
) -> None:
    width, height = 1500, 800
    image = Image.new("RGB", (width, height), "#176783")
    draw = ImageDraw.Draw(image)
    title_font = _font(38, True)
    label_font = _font(18, True)
    small_font = _font(16)
    draw.text((width // 2, 35), title.upper(), font=title_font, fill="white", anchor="ma")
    left, top, right, bottom = 120, 150, width - 70, height - 175
    values = [_number(item.get(value_key)) for item in history]
    max_value = max(values, default=1) or 1
    points: list[tuple[int, int]] = []
    for index, value in enumerate(values):
        x = left + int((right - left) * index / max(len(values) - 1, 1))
        y = bottom - int((bottom - top) * value / max_value)
        points.append((x, y))
        draw.line((x, y, x, bottom), fill="#78AFC1", width=1)
    if len(points) > 1:
        draw.line(points, fill="white", width=5, joint="curve")
    for index, ((x, y), value) in enumerate(zip(points, values)):
        draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill="white")
        draw.text((x, y - 22), _format_number(value), font=label_font, fill="white", anchor="ms")
        draw.text(
            (x, bottom + 28),
            str(history[index].get("label") or history[index].get("month") or ""),
            font=small_font,
            fill="white",
            anchor="ma",
        )
    image.save(path)


def _monthly_new_values(history: Sequence[Mapping[str, Any]]) -> list[int]:
    values: list[int] = []
    for index, item in enumerate(history):
        if item.get("new") is not None:
            values.append(_number(item.get("new")))
            continue
        if index == 0:
            values.append(0)
            continue
        previous = _number(history[index - 1].get("non_mitigated"))
        current = _number(item.get("non_mitigated"))
        mitigated = _number(item.get("mitigated"))
        values.append(max(0, current - (previous - mitigated)))
    return values


def _chart(document: Any, path: Path, alt: str) -> Any:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    faithful.base._add_picture(paragraph, path, width=Cm(16.5), alt_text=alt)
    return paragraph


def _custom_data(dataset: Mapping[str, Any]) -> Mapping[str, Any]:
    value = dataset.get("customizations")
    return value if isinstance(value, Mapping) else {}


def resolve_custom_intelligence_modules(
    profile: ClientProfile,
) -> tuple[tuple[str, ...], dict[str, str]]:
    active: list[str] = []
    suppressed: dict[str, str] = {}
    for module in profile.report.intelligence_modules:
        if module == "cloud_container_images" and profile.cloud_security_scope.enabled:
            suppressed[module] = "MOVED_TO_CLOUD_REPORT"
        else:
            active.append(module)
    return tuple(active), suppressed


def _module_enabled(profile: ClientProfile, module: str) -> bool:
    active, _ = resolve_custom_intelligence_modules(profile)
    return module in active


def _cover_title(document: Any) -> None:
    for paragraph in document.paragraphs:
        if "RELATÓRIO DE" in paragraph.text and "VULNERABILIDADES" in paragraph.text:
            paragraph.text = "INTELIGÊNCIA E\nCUSTOMIZAÇÕES\nTENABLE"
            for run in paragraph.runs:
                base._set_run_font(run, size=28, color=base.BLUE, bold=True)
            break


def _new_report_section(document: Any) -> None:
    # These report blocks keep the same page setup. A regular page break also
    # preserves the template header/footer reliably in Word and LibreOffice.
    document.add_page_break()


def _mirror_default_header_footer_on_even_pages(document: Any) -> None:
    section = document.sections[0]
    for source, target in (
        (section.header, section.even_page_header),
        (section.footer, section.even_page_footer),
    ):
        target.is_linked_to_previous = False
        for child in list(target._element):
            target._element.remove(child)
        for child in source._element:
            target._element.append(deepcopy(child))
        target.part.rels.clear()
        for relationship_id, relationship in source.part.rels.items():
            relationship_target = (
                relationship.target_ref
                if relationship.is_external
                else relationship.target_part
            )
            target.part.rels.add_relationship(
                relationship.reltype,
                relationship_target,
                relationship_id,
                relationship.is_external,
            )


def _back_cover(document: Any) -> None:
    first_section = document.sections[0]
    blank_first_references = [
        deepcopy(reference)
        for reference_tag in ("w:headerReference", "w:footerReference")
        for reference in first_section._sectPr.findall(qn(reference_tag))
        if reference.get(qn("w:type")) == "first"
    ]
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.different_first_page_header_footer = True
    for reference in reversed(blank_first_references):
        section._sectPr.insert(0, reference)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Cm(9)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(editorial.BACK_COVER)
    base._set_run_font(run, size=20, color=base.NAVY, bold=True)


def _monthly_modules(document: Any, data: Mapping[str, Any], temp: Path, rendered: list[str]) -> None:
    views = _monthly_views(data)
    if not views:
        return
    faithful._heading(document, "3.1. Comparativo Mensal de Vulnerabilidades Mitigadas e Não Mitigadas.", 2)
    faithful._paragraph(document, MONTHLY_COMPARISON)
    year = str(views[0][1][-1].get("label") or "").split("/")[-1]
    faithful._heading(document, "3.1.1. Vulnerabilidades “Não Mitigadas”.", 3)
    for view_index, (view_label, history) in enumerate(views):
        if view_index:
            _new_report_section(document)
        view_heading = faithful._bullet(document, view_label)
        view_heading.paragraph_format.keep_with_next = True
        severity = _severity_rows(
            history,
            nested_key="non_mitigated_by_severity",
            total_key="non_mitigated",
        )
        if severity:
            comparison = temp / f"non-mitigated-comparison-{view_index}.png"
            _grouped_monthly_chart(
                comparison,
                f"Comparativo de Vulnerabilidades Não Mitigadas {year}",
                severity,
                (*SEVERITY_SERIES, ("total", "Total Vulnerabilidades", "#B244A5")),
            )
            _chart(
                document,
                comparison,
                f"Comparativo mensal de vulnerabilidades não mitigadas por severidade - {view_label}",
            )
        volume = temp / f"non-mitigated-volume-{view_index}.png"
        _monthly_line_chart(
            volume,
            f"Volume de Vuln. Não Mitigadas {year}",
            history,
            value_key="non_mitigated",
        )
        _chart(document, volume, f"Volume mensal de vulnerabilidades não mitigadas - {view_label}")

    faithful._heading(document, "3.1.2. Vulnerabilidades “Mitigadas”.", 3)
    for view_index, (view_label, history) in enumerate(views):
        if view_index:
            _new_report_section(document)
        view_heading = faithful._bullet(document, view_label)
        view_heading.paragraph_format.keep_with_next = True
        severity = _severity_rows(
            history,
            nested_key="mitigated_by_severity",
            total_key="mitigated",
        )
        if severity:
            comparison = temp / f"mitigated-comparison-{view_index}.png"
            _grouped_monthly_chart(
                comparison,
                f"Comparativo de Vuln. Mitigadas {year}",
                severity,
                (*SEVERITY_SERIES, ("total", "Total de Vuln. Mitigadas", "#B244A5")),
            )
            _chart(
                document,
                comparison,
                f"Comparativo mensal de vulnerabilidades mitigadas por severidade - {view_label}",
            )
        volume = temp / f"mitigated-volume-{view_index}.png"
        _monthly_line_chart(
            volume,
            f"Volume de Vuln. Mitigadas {year}",
            history,
            value_key="mitigated",
        )
        _chart(document, volume, f"Volume mensal de vulnerabilidades mitigadas - {view_label}")
    rendered.append("vm_monthly_volume")


def _scan_health(document: Any, data: Mapping[str, Any], temp: Path, rendered: list[str]) -> None:
    health = data.get("scan_auth_health")
    statuses = data.get("customization_statuses")
    known = isinstance(statuses, Mapping) and "scan_auth_health" in statuses
    if not isinstance(health, Mapping) and not known:
        return
    heading = faithful._heading(document, "Integridade da varredura", 2)
    heading.paragraph_format.page_break_before = True
    faithful._paragraph(document, SCAN_HEALTH)
    if not isinstance(health, Mapping) or not _number(health.get("total")):
        faithful._paragraph(document, NO_DATA_MESSAGES["scan_auth_health"])
        rendered.append("scan_auth_health")
        return
    chart = temp / "scan-health.png"
    rows = [{"label": "", "success": health.get("success", 0), "failure": health.get("failure", 0)}]
    _bar_chart(chart, "Integridade da varredura", rows, (("success", "Sucesso", "#00B050"), ("failure", "Falha", "#FF0000")))
    _chart(document, chart, "Integridade da varredura por sucesso e falha de autenticação")
    rendered.append("scan_auth_health")


def _previous_period(
    document: Any,
    data: Mapping[str, Any],
    dataset: Mapping[str, Any],
    profile: ClientProfile,
    rendered: list[str],
) -> None:
    previous = data.get("previous_period_overview")
    if not isinstance(previous, Mapping):
        return
    faithful._paragraph(document, f"Comparativo relatório anterior ({previous.get('label', '')})")
    rows = []
    for label, key in (("TOTAL", "total"), ("Crítica", "critical"), ("Alta", "high"), ("Média", "medium"), ("Baixa", "low")):
        item = previous.get(key) if isinstance(previous.get(key), Mapping) else {}
        rows.append((label, item.get("mitigated", ""), item.get("non_mitigated", ""), item.get("exploitable", ""), item.get("patch_available_over_30_days", "")))
    faithful._simple_table(document, ("", "Mitigado", "Não Mitigado", "Explorável", "Patchs disponíveis > 30d"), rows, widths=(1300, 1500, 1800, 1500, 3000), left_columns=frozenset({0}))
    add_source_filter_note(
        document,
        dataset,
        "previous_period_overview",
        enabled=profile.presentation.show_source_filters,
        period_label=str(previous.get("label") or ""),
    )
    rendered.append("vm_previous_period_delta")


def _network_comparison(
    document: Any,
    data: Mapping[str, Any],
    dataset: Mapping[str, Any],
    profile: ClientProfile,
    mask_sensitive: bool,
    rendered: list[str],
) -> None:
    comparisons = data.get("network_comparisons")
    if not isinstance(comparisons, list) or not comparisons:
        snapshots = data.get("network_tag_snapshots")
        if not isinstance(snapshots, list) or not snapshots:
            return
        comparisons = [
            {
                "network": item.get("network") or item.get("label"),
                "periods": [{**item, "label": "Baseline do período atual"}],
            }
            for item in snapshots if isinstance(item, Mapping)
        ]
    rendered_comparisons = 0
    for comparison in comparisons:
        if not isinstance(comparison, Mapping):
            continue
        periods = comparison.get("periods")
        period_rows = [item for item in periods or [] if isinstance(item, Mapping)]
        if not period_rows:
            continue
        network = str(comparison.get("network") or comparison.get("label") or "").strip()
        faithful._heading(
            document,
            f"3.2.1. Comparativo dos Principais ativos Vulneráveis por Rede: {network}.",
            3,
        )
        faithful._paragraph(
            document,
            f"Segue abaixo TOP 20, agora por segmento de rede do {profile.display_name}: {network}:",
        )
        for period_row in period_rows:
            faithful._bullet(
                document,
                f"{period_row.get('label') or period_row.get('period_id') or ''}:",
            )
            rows = []
            for index, asset in enumerate(period_row.get("assets") or [], start=1):
                if isinstance(asset, Mapping):
                    rows.append((index, "" if mask_sensitive else asset.get("ip_address", ""), "" if mask_sensitive else asset.get("asset_name", ""), asset.get("critical", ""), asset.get("high", ""), asset.get("medium", ""), asset.get("low", ""), asset.get("total", ""), asset.get("exploitable", "")))
            faithful._simple_table(document, ("Nº", "IP Address", "Asset Name", "Crítica", "Alta", "Média", "Baixa", "Total", "Exploitable"), rows, widths=(500, 1350, 1750, 800, 800, 800, 800, 950, 1250), left_columns=frozenset({1, 2}))
            add_source_filter_note(
                document,
                dataset,
                "network_tag_snapshots",
                enabled=profile.presentation.show_source_filters,
                tag_uuid=str(comparison.get("tag_uuid") or period_row.get("tag_uuid") or ""),
                period_label=str(
                    period_row.get("label") or period_row.get("period_id") or ""
                ),
            )
        movement = (
            _asset_movement(
                period_rows[0], period_rows[-1], mask_sensitive=mask_sensitive
            )
            if len(period_rows) >= 2 else []
        )
        if movement:
            faithful._simple_table(
                document,
                (
                    "IP Address", "Asset Name", "Posição anterior",
                    "Posição atual", "Movimentação", "Delta",
                ),
                movement,
                widths=(1400, 2100, 1300, 1200, 1700, 900),
                left_columns=frozenset({0, 1, 4}),
            )
            add_source_filter_note(
                document,
                dataset,
                "network_asset_movement",
                enabled=profile.presentation.show_source_filters,
                tag_uuid=str(comparison.get("tag_uuid") or ""),
                period_labels=tuple(
                    str(item.get("label") or item.get("period_id") or "")
                    for item in (period_rows[0], period_rows[-1])
                ),
            )
        rendered_comparisons += 1
    if rendered_comparisons:
        rendered.append("vm_network_comparison")


def _asset_movement(
    previous: Mapping[str, Any],
    current: Mapping[str, Any],
    *,
    mask_sensitive: bool,
) -> list[tuple[Any, ...]]:
    previous_assets = [
        item for item in previous.get("assets") or [] if isinstance(item, Mapping)
    ]
    current_assets = [
        item for item in current.get("assets") or [] if isinstance(item, Mapping)
    ]
    previous_rank = {
        str(item.get("asset_key") or item.get("source_asset_id") or ""): index
        for index, item in enumerate(previous_assets, start=1)
        if item.get("asset_key") or item.get("source_asset_id")
    }
    previous_total = {
        str(item.get("asset_key") or item.get("source_asset_id") or ""): _number(
            item.get("total")
        )
        for item in previous_assets
        if item.get("asset_key") or item.get("source_asset_id")
    }
    rows: list[tuple[Any, ...]] = []
    for current_position, item in enumerate(current_assets, start=1):
        identity = str(item.get("asset_key") or item.get("source_asset_id") or "")
        if not identity:
            continue
        prior_position = previous_rank.get(identity)
        prior_total = previous_total.get(identity)
        current_total = _number(item.get("total"))
        if prior_position is None:
            movement = "Entrada"
            delta = current_total
        else:
            delta = current_total - int(prior_total or 0)
            movement = (
                "Aumento" if delta > 0 else "Redução" if delta < 0 else "Permaneceu"
            )
        rows.append((
            "" if mask_sensitive else item.get("ip_address", ""),
            "" if mask_sensitive else item.get("asset_name", ""),
            prior_position if prior_position is not None else "-",
            current_position,
            movement,
            delta,
        ))
    return rows


def _plugin_family(
    document: Any,
    data: Mapping[str, Any],
    dataset: Mapping[str, Any],
    profile: ClientProfile,
    rendered: list[str],
) -> None:
    rows = data.get("plugin_family")
    statuses = data.get("customization_statuses")
    known = isinstance(statuses, Mapping) and "vm_plugin_family" in statuses
    if not isinstance(rows, list) and not known:
        return
    faithful._paragraph(document, PLUGIN_FAMILY)
    if not rows:
        faithful._paragraph(document, NO_DATA_MESSAGES["vm_plugin_family"])
        rendered.append("vm_plugin_family")
        return
    values = [(item.get("family", ""), item.get("total", "")) for item in rows if isinstance(item, Mapping)]
    faithful._simple_table(document, ("Família de Plugin", "Total"), values, widths=(7200, 1800), left_columns=frozenset({0}))
    add_source_filter_note(
        document,
        dataset,
        "plugin_family",
        enabled=profile.presentation.show_source_filters,
    )
    rendered.append("vm_plugin_family")


def _eol(
    document: Any,
    data: Mapping[str, Any],
    dataset: Mapping[str, Any],
    profile: ClientProfile,
    mask_sensitive: bool,
    rendered: list[str],
) -> None:
    assets = data.get("eol_assets")
    software = data.get("eol_software")
    statuses = data.get("customization_statuses")
    known = isinstance(statuses, Mapping) and "vm_eol_software" in statuses
    if not known and not (isinstance(assets, list) and assets) and not (isinstance(software, list) and software):
        return
    faithful._heading(document, "Sistemas operacionais e software sem suportes")
    for text in (EOL_INTRO, EOL_TENABLE, EOL_METHOD):
        faithful._paragraph(document, text)
    if not assets and not software:
        faithful._paragraph(document, NO_DATA_MESSAGES["vm_eol_software"])
        rendered.append("vm_eol_software")
        return
    plugin_ids = ", ".join(
        str(item.get("plugin_id"))
        for item in software or []
        if isinstance(item, Mapping) and item.get("plugin_id") is not None
    )
    plugin_filter = {"Plugin ID": plugin_ids} if plugin_ids else None
    faithful._heading(document, "6.1. Ativos com SOs e Softwares sem suportes.", 2)
    faithful._paragraph(document, EOL_ASSETS)
    asset_rows = []
    for item in assets or []:
        if isinstance(item, Mapping):
            asset_rows.append((
                "" if mask_sensitive else item.get("ip_address", ""),
                "" if mask_sensitive else item.get("asset_name", ""),
                item.get("critical", ""), item.get("high", ""), item.get("medium", ""), item.get("low", ""), item.get("total", ""),
            ))
    faithful._simple_table(document, ("IP Address", "Asset Name", "Crítica", "Alta", "Média", "Baixa", "Total"), asset_rows, widths=(1500, 2200, 1000, 1000, 1000, 1000, 1100), left_columns=frozenset({0, 1}))
    add_source_filter_note(
        document,
        dataset,
        "eol_assets",
        enabled=profile.presentation.show_source_filters,
        extra_filters=plugin_filter,
    )
    faithful._heading(document, "6.1. Principais Softwares e SOs sem suportes por vulnerabilidades", 2)
    faithful._paragraph(document, EOL_SOFTWARE)
    faithful._paragraph(document, EOL_PROTOCOLS)
    software_rows = []
    for item in software or []:
        if isinstance(item, Mapping):
            software_rows.append((item.get("plugin_id", ""), item.get("name", ""), item.get("family", ""), item.get("severity", ""), item.get("total", "")))
    faithful._simple_table(document, ("Plugin ID", "Nome", "Família OS", "Severidade", "Total"), software_rows, widths=(900, 3400, 2300, 1300, 1000), left_columns=frozenset({1, 2}))
    add_source_filter_note(
        document,
        dataset,
        "eol_software",
        enabled=profile.presentation.show_source_filters,
        extra_filters=plugin_filter,
    )
    rendered.append("vm_eol_software")


def _executive(document: Any, data: Mapping[str, Any], temp: Path, rendered: list[str]) -> None:
    rows = data.get("vulnerability_evolution")
    if not isinstance(rows, list) or not rows:
        return
    faithful._heading(document, "6.3. Análise Executiva da Evolução de Vulnerabilidades e Criticidade dos Ativos", 2)
    for text in EXECUTIVE_PARAGRAPHS:
        faithful._paragraph(document, text)
    chart = temp / "executive.png"
    _bar_chart(chart, "Evolução de Vulnerabilidades", rows, (("change", "Variação", "#2E59FC"),))
    _chart(document, chart, "Evolução de vulnerabilidades e criticidade dos ativos")
    rendered.append("vm_executive_evolution")


def _evolution(document: Any, data: Mapping[str, Any], temp: Path, rendered: list[str]) -> None:
    views = _monthly_views(data)
    if not views:
        return
    history = views[0][1]
    faithful._heading(document, "Evolução mensal de Vulnerabilidades")
    for text in EVOLUTION_PARAGRAPHS:
        faithful._paragraph(document, text)
    evolution_rows = []
    for item, new_value in zip(history, _monthly_new_values(history)):
        evolution_rows.append({
            "label": item.get("label") or item.get("month") or "",
            "non_mitigated": _number(item.get("non_mitigated")),
            "mitigated": _number(item.get("mitigated")),
            "new": new_value,
        })
    year = str(history[-1].get("label") or "").split("/")[-1]
    chart = temp / "monthly-evolution.png"
    _grouped_monthly_chart(
        chart,
        f"Evolução Mensal Vulnerabilidades {year}",
        evolution_rows,
        (
            ("non_mitigated", "Total de Vuln. Não Mitigadas", "#FF0000"),
            ("mitigated", "Total Mitigadas", "#00B050"),
            ("new", "Vuln. Novas", "#FFF200"),
        ),
    )
    document.add_page_break()
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    _chart(
        document,
        chart,
        "Evolução mensal de vulnerabilidades não mitigadas, mitigadas e novas",
    )
    new_rows = _severity_rows(history, nested_key="new_by_severity", total_key="new")
    if new_rows:
        new_chart = temp / "monthly-new-by-severity.png"
        _grouped_monthly_chart(
            new_chart,
            f"Comparativo de Vulnerabilidades “Novas” por níveis elevados de Risco {year}",
            new_rows,
            SEVERITY_SERIES,
        )
        _chart(
            document,
            new_chart,
            "Comparativo mensal de vulnerabilidades novas por severidade",
        )
    rendered.append("vm_monthly_evolution")


def _containers(
    document: Any,
    data: Mapping[str, Any],
    dataset: Mapping[str, Any],
    profile: ClientProfile,
    mask_sensitive: bool,
    rendered: list[str],
) -> None:
    images = data.get("container_images")
    if not isinstance(images, list) or not images:
        return
    heading = faithful._heading(document, "TENABLE CLOUD SECURITY (CONTAINER IMAGES)")
    heading.paragraph_format.page_break_before = True
    faithful._paragraph(document, CONTAINER_INTRO)
    faithful._paragraph(document, CONTAINER_CURRENT)
    faithful._heading(document, "Top 5 Imagens de container mais vulneráveis", 2)
    values = []
    for item in images[:5]:
        if isinstance(item, Mapping):
            values.append(("" if mask_sensitive else item.get("repository", ""), item.get("tag", ""), item.get("critical", ""), item.get("high", ""), item.get("medium", ""), item.get("low", "")))
    faithful._simple_table(document, ("Container Repository", "Tag", "Crítica", "Alta", "Média", "Baixa"), values, widths=(3300, 1100, 1150, 1150, 1150, 1150), left_columns=frozenset({0, 1}))
    add_source_filter_note(
        document,
        dataset,
        "container_images",
        enabled=profile.presentation.show_source_filters,
    )
    faithful._heading(document, "Overview das vulnerabilidades das imagens de container", 2)
    for item in images[:5]:
        if not isinstance(item, Mapping):
            continue
        title = "" if mask_sensitive else str(item.get("repository") or "")
        faithful._paragraph(document, f"{title}: {item.get('tag', '')}")
        findings = []
        for finding in item.get("findings") or []:
            if isinstance(finding, Mapping):
                findings.append((finding.get("cve", ""), finding.get("severity", ""), finding.get("vpr", ""), finding.get("software", ""), finding.get("fixed_by", "")))
        faithful._simple_table(document, ("CVE", "Severidade (VPR)", "VPR", "Software", "Fixed by"), findings, widths=(1500, 1650, 1000, 3100, 1850), left_columns=frozenset({0, 3, 4}))
        extra_filters = {}
        if not mask_sensitive:
            extra_filters = {
                "Repository": item.get("repository"),
                "Tag": item.get("tag"),
            }
        add_source_filter_note(
            document,
            dataset,
            "container_findings",
            enabled=profile.presentation.show_source_filters,
            extra_filters=extra_filters,
        )
    rendered.append("cloud_container_images")


def _attack_vector(
    document: Any,
    data: Mapping[str, Any],
    dataset: Mapping[str, Any],
    profile: ClientProfile,
    rendered: list[str],
) -> None:
    vectors = data.get("attack_vectors")
    statuses = data.get("customization_statuses")
    known = isinstance(statuses, Mapping) and "vm_exploit_vector" in statuses
    if not isinstance(vectors, list) and not known:
        return
    faithful._heading(document, "9.4. Vulnerabilidades Exploráveis por Vetor de Ataque", 2)
    faithful._paragraph(document, ATTACK_VECTOR)
    faithful._paragraph(document, "Os vetores de ameaça são designados como: Rede (AV:N), Adjacente (AV:A) e Local (AV:L).")
    if not vectors:
        faithful._paragraph(document, NO_DATA_MESSAGES["vm_exploit_vector"])
        rendered.append("vm_exploit_vector")
        return
    rows = [(item.get("framework", ""), item.get("local", ""), item.get("network", ""), item.get("adjacent_network", "")) for item in vectors if isinstance(item, Mapping)]
    faithful._simple_table(document, ("", "Local", "Network", "Adj. Network"), rows, widths=(3000, 2000, 2000, 2000), left_columns=frozenset({0}))
    add_source_filter_note(
        document,
        dataset,
        "attack_vectors",
        enabled=profile.presentation.show_source_filters,
    )
    rendered.append("vm_exploit_vector")


def _was_unsupported(
    document: Any,
    data: Mapping[str, Any],
    dataset: Mapping[str, Any],
    profile: ClientProfile,
    rendered: list[str],
) -> None:
    items = data.get("was_unsupported_tech")
    statuses = data.get("customization_statuses")
    known = isinstance(statuses, Mapping) and "was_unsupported_tech" in statuses
    if not isinstance(items, list) and not known:
        return
    faithful._heading(document, "WAS Vulnerabilidades WEB – Principais Aplicações “Unsupported”", 3)
    faithful._paragraph(document, WAS_UNSUPPORTED)
    if not items:
        faithful._paragraph(
            document,
            editorial.WAS_COLLECTION_UNAVAILABLE
            if isinstance(statuses, Mapping)
            and statuses.get("was_unsupported_tech") == "DATA_UNAVAILABLE"
            else NO_DATA_MESSAGES["was_unsupported_tech"],
        )
        rendered.append("was_unsupported_tech")
        return
    rows = []
    for item in items:
        if isinstance(item, Mapping):
            rows.append((item.get("plugin_id", ""), item.get("name", ""), item.get("family", ""), item.get("severity", ""), item.get("total", ""), item.get("applications", ""), item.get("vpr", "")))
    plugin_ids = ", ".join(
        str(item.get("plugin_id"))
        for item in items
        if isinstance(item, Mapping) and item.get("plugin_id") is not None
    )
    faithful._simple_table(document, ("Plugin Id", "Nome", "Família", "Severidade", "Total", "App.", "VPR"), rows, widths=(850, 2750, 1850, 1050, 800, 800, 800), left_columns=frozenset({1, 2}))
    add_source_filter_note(
        document,
        dataset,
        "was_unsupported_tech",
        enabled=profile.presentation.show_source_filters,
        extra_filters={"Plugin ID": plugin_ids} if plugin_ids else None,
    )
    rendered.append("was_unsupported_tech")


def generate_customizations_report(
    *,
    template_path: str | Path,
    dataset_path: str | Path,
    profile: ClientProfile,
    output_path: str | Path,
    mask_sensitive: bool = False,
) -> CustomizationsRenderResult:
    template = Path(template_path)
    dataset_file = Path(dataset_path)
    output = Path(output_path)
    if not template.is_file():
        raise ValueError(f"Template Word não encontrado: {template}")
    dataset = json.loads(dataset_file.read_text(encoding="utf-8"))
    faithful._validate_dataset(dataset, profile)
    document = Document(template)
    document.settings.odd_and_even_pages_header_footer = True
    faithful._clear_body_after_cover_break(document)
    faithful._configure_styles(document)
    period_label, period_range = base._period_labels(dataset["period"])
    base._replace_tokens(document, {"{{CLIENT_NAME}}": profile.display_name, "{{PERIOD_LABEL}}": period_label, "{{PERIOD_RANGE}}": period_range, "{{TEMPLATE_VERSION}}": CUSTOM_TEMPLATE_VERSION})
    _cover_title(document)
    faithful._sanitize_header_footer(document, profile.display_name)
    _mirror_default_header_footer_on_even_pages(document)
    faithful._sanitize_properties(document, title="INTELIGÊNCIA E CUSTOMIZAÇÕES TENABLE")
    faithful._heading(document, "SUMÁRIO")
    faithful._toc_field(document)
    document.add_page_break()
    data = _custom_data(dataset)
    rendered: list[str] = []
    history_status = data.get("history_status")
    if (
        isinstance(history_status, Mapping)
        and history_status.get("status") != "COMPATIBLE_PREDECESSOR"
        and any(_module_enabled(profile, module) for module in HISTORICAL_INTELLIGENCE_MODULES)
    ):
        faithful._paragraph(document, NO_HISTORY_MESSAGE)
    with tempfile.TemporaryDirectory() as directory:
        temp = Path(directory)
        if _module_enabled(profile, "vm_monthly_volume"):
            _monthly_modules(document, data, temp, rendered)
        if _module_enabled(profile, "scan_auth_health"):
            _scan_health(document, data, temp, rendered)
        if _module_enabled(profile, "vm_previous_period_delta"):
            _previous_period(document, data, dataset, profile, rendered)
        if _module_enabled(profile, "vm_plugin_family"):
            _plugin_family(document, data, dataset, profile, rendered)
        if _module_enabled(profile, "vm_eol_software"):
            _eol(document, data, dataset, profile, mask_sensitive, rendered)
        if _module_enabled(profile, "vm_executive_evolution"):
            _executive(document, data, temp, rendered)
        if _module_enabled(profile, "vm_monthly_evolution"):
            _evolution(document, data, temp, rendered)
        if _module_enabled(profile, "cloud_container_images"):
            _containers(document, data, dataset, profile, mask_sensitive, rendered)
        if _module_enabled(profile, "vm_exploit_vector"):
            _attack_vector(document, data, dataset, profile, rendered)
        if _module_enabled(profile, "was_unsupported_tech"):
            _was_unsupported(document, data, dataset, profile, rendered)
    if rendered:
        _back_cover(document)
    base._enable_field_updates(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    requested = tuple(profile.report.intelligence_modules)
    _, suppressed = resolve_custom_intelligence_modules(profile)
    omitted = tuple(
        {
            "module_id": module,
            "reason": (
                suppressed[module]
                if module in suppressed
                else "MOVED_TO_TAG_REPORT"
                if module == "vm_network_comparison"
                else "NO_COMPATIBLE_HISTORY"
                if module in HISTORICAL_INTELLIGENCE_MODULES
                else "NO_COMPATIBLE_DATA"
            ),
        }
        for module in requested
        if module not in rendered
    )
    return CustomizationsRenderResult(
        output,
        CUSTOM_TEMPLATE_VERSION,
        profile.client_id,
        str(dataset["period"].get("period_id") or ""),
        requested,
        tuple(rendered),
        omitted,
    )
