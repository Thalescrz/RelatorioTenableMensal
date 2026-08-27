from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Mapping, Sequence
from zoneinfo import ZoneInfo

from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from tenable_reports.presentation import base_report_docx as base
from tenable_reports.presentation import cloud_editorial_catalog as copy
from tenable_reports.presentation.cloud_visuals import (
    normalize_history_series,
    render_aging_chart,
    render_monthly_history_chart,
    render_severity_chart,
)
from tenable_reports.presentation.source_filters import format_source_filter_note
from tenable_reports.presentation.translation import (
    TextTranslator,
    split_translation_chunks,
    translate_in_chunks,
)


BLUE = "2E59FC"
NAVY = "101326"
LIGHT_BLUE = "EAF0FF"
LIGHT_GRAY = "F1F3F6"
MID_GRAY = "68728A"
WHITE = "FFFFFF"
SEVERITY_FILLS = {
    "CRITICAL": "C00000",
    "HIGH": "F26B00",
    "MEDIUM": "FFF200",
    "LOW": "00B050",
    "NONE": "B7C9C5",
}
SEVERITY_LABELS = {
    "CRITICAL": "Crítica",
    "HIGH": "Alta",
    "MEDIUM": "Média",
    "LOW": "Baixa",
    "NONE": "Sem vulnerabilidade",
}


@dataclass(slots=True)
class CloudDocumentBuilder:
    document: DocxDocument
    anchor: Any

    def _move(self, element: Any) -> Any:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
        self.anchor._p.addprevious(element)
        return element

    def paragraph(
        self,
        text: str = "",
        *,
        bold: bool = False,
        size: float = 10.5,
        color: str = NAVY,
        align: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
        keep_with_next: bool = False,
        space_before: float = 0,
        space_after: float = 6,
        font_name: str = "Times New Roman",
    ) -> Any:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.line_spacing = 1.08
        paragraph.paragraph_format.keep_with_next = keep_with_next
        run = paragraph.add_run(str(text))
        run.bold = bold
        run.font.name = font_name
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        self._move(paragraph._p)
        return paragraph

    def heading(self, text: str, level: int = 1) -> Any:
        sizes = {1: 15, 2: 12.5, 3: 11.5, 4: 10.5}
        paragraph = self.paragraph(
            text,
            bold=True,
            size=sizes.get(level, 10.5),
            color=BLUE if level <= 2 else NAVY,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            keep_with_next=True,
            space_before=10 if level <= 2 else 6,
            space_after=5,
            font_name="Arial",
        )
        properties = paragraph._p.get_or_add_pPr()
        outline = properties.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            properties.append(outline)
        outline.set(qn("w:val"), str(max(0, min(level, 9) - 1)))
        return paragraph

    def standard_paragraph(
        self,
        text: str = "",
        *,
        bold: bool = False,
        color: str = NAVY,
        align: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
    ) -> Any:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = align
        run = paragraph.add_run(str(text))
        base._set_run_font(run, size=9, color=color, bold=bold)
        base._set_language(run)
        self._move(paragraph._p)
        return paragraph

    def standard_heading(self, text: str, level: int = 1) -> Any:
        paragraph = self.document.add_paragraph(
            str(text),
            style=f"Heading {level}",
        )
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.keep_with_next = True
        size = {1: 14, 2: 11, 3: 10, 4: 9}.get(level, 9)
        for run in paragraph.runs:
            base._set_run_font(
                run,
                size=size,
                color=base.NAVY,
                bold=True,
            )
            base._set_language(run)
        self._move(paragraph._p)
        return paragraph

    def bullet(self, text: str) -> Any:
        return self.paragraph(
            f"• {text}",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=4,
        )

    def page_break(self) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        self._move(paragraph._p)

    def image(
        self,
        path: str | Path,
        *,
        alt_text: str,
        width_cm: float = 16.0,
    ) -> Any:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(8)
        base._add_picture(
            paragraph,
            Path(path),
            width=Cm(width_cm),
            alt_text=alt_text,
        )
        self._move(paragraph._p)
        return paragraph

    def table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[Any]],
        *,
        widths: Sequence[int] | None = None,
        left_columns: frozenset[int] = frozenset(),
        empty_message: str = copy.EMPTY_TABLE_MONTH,
    ) -> Any | None:
        if not rows:
            self.paragraph(
                empty_message,
                color=MID_GRAY,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            return None
        if widths is None:
            widths = tuple(9200 // len(headers) for _ in headers)
        if len(widths) != len(headers):
            raise ValueError("Cabeçalhos e larguras da tabela Cloud são incompatíveis.")
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        base._set_table_fixed(table, sum(widths))
        header = table.rows[0]
        base._set_repeat_table_header(header)
        base._prevent_row_split(header)
        for index, (cell, label, width) in enumerate(zip(header.cells, headers, widths)):
            base._set_cell_width(cell, width)
            base._set_cell_shading(cell, BLUE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(label))
            base._set_run_font(run, size=7.5, color=WHITE, bold=True, name="Arial")
        for row_index, values in enumerate(rows):
            row = table.add_row()
            base._prevent_row_split(row)
            for column, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
                base._set_cell_width(cell, width)
                if row_index % 2:
                    base._set_cell_shading(cell, LIGHT_GRAY)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column in left_columns
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                run = paragraph.add_run("" if value is None else str(value))
                base._set_run_font(run, size=7.2, color=NAVY, name="Arial")
        self._move(table._tbl)
        return table

    def source_note(
        self,
        dataset: Mapping[str, Any],
        table_id: str,
        *,
        enabled: bool,
    ) -> None:
        if not enabled:
            return
        note = format_source_filter_note(dataset, table_id)
        if note:
            self.paragraph(
                note,
                size=8,
                color=MID_GRAY,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_name="Arial",
                space_before=2,
                space_after=6,
            )


def _period_labels(dataset: Mapping[str, Any]) -> tuple[str, str, str]:
    period = dataset.get("period") or {}
    timezone_name = str(period.get("timezone") or "UTC")
    timezone = ZoneInfo(timezone_name)
    start = datetime.fromisoformat(str(period["start_at"]).replace("Z", "+00:00"))
    end = datetime.fromisoformat(str(period["end_at"]).replace("Z", "+00:00"))
    start_local = start.astimezone(timezone)
    end_local = (end.astimezone(timezone) - timedelta(microseconds=1))
    return (
        start_local.strftime("%d/%m/%Y"),
        end_local.strftime("%d/%m/%Y"),
        start_local.strftime("%m/%Y"),
    )


def _asset_identity(row: Mapping[str, Any]) -> str:
    if row.get("kind") == "container_image":
        return str(row.get("repository_uri") or row.get("name") or "")
    return str(row.get("name") or "")


def render_document_control(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
) -> None:
    start, end, _ = _period_labels(dataset)
    builder.heading("1. CONTROLE DE DOCUMENTO", 1)
    builder.table(
        ("Preparação", "Nome", "Data"),
        (("Preparação", "", ""),),
        widths=(2600, 4200, 2400),
        left_columns=frozenset({0, 1}),
    )
    builder.table(
        ("Controle de Versionamento", "Versão", "Data"),
        (("Relatório mensal", "1.0", end),),
        widths=(4600, 2200, 2400),
        left_columns=frozenset({0}),
    )
    builder.table(
        ("Lista de Distribuição", "Nome", "Organização"),
        (("Distribuição", "", ""),),
        widths=(3200, 3000, 3000),
        left_columns=frozenset({0, 1, 2}),
    )
    builder.heading("2. OBJETIVO", 1)
    builder.paragraph(copy.OBJECTIVE)
    builder.paragraph(
        f"Período deste relatório compreende-se entre {start} e {end}.",
        bold=True,
    )


def render_cloud_overview(builder: CloudDocumentBuilder) -> None:
    builder.heading("3. TENABLE CLOUD SECURITY", 1)
    builder.paragraph(copy.CLOUD_OVERVIEW)
    builder.paragraph(copy.CLOUD_INTEGRATION)


def render_introduction(builder: CloudDocumentBuilder) -> None:
    builder.heading("3.1. Introdução", 2)
    builder.paragraph(copy.REPORT_OBJECTIVES_INTRO)
    for item in copy.REPORT_OBJECTIVES:
        builder.bullet(item)
    builder.paragraph(copy.DETECTION_INTRO)
    for item in copy.DETECTION_CAPABILITIES:
        builder.bullet(item)


def render_top_hosts(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    builder.heading("3.2. Principais Hosts Vulneráveis", 2)
    builder.paragraph(copy.TOP_HOSTS_INTRO)
    builder.paragraph(copy.TOP_HOSTS_DETAILS)
    rows = [
        (
            row.get("name") or "",
            ", ".join(row.get("ip_addresses") or ()),
            row.get("critical", 0),
            row.get("high", 0),
            row.get("medium", 0),
            row.get("low", 0),
            row.get("vulnerabilities", 0),
        )
        for row in dataset.get("top_vulnerable_hosts") or ()
    ]
    builder.table(
        ("Asset Name", "IP Address", "Crítica", "Alta", "Média", "Baixa", "Total"),
        rows,
        widths=(2100, 1600, 1050, 1050, 1050, 1050, 1300),
        left_columns=frozenset({0, 1}),
    )
    builder.source_note(dataset, "cloud_top_hosts", enabled=show_source_filters)


def render_top_images(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    builder.heading("3.3. Imagens de Contêineres Mais Vulneráveis", 2)
    builder.paragraph(copy.TOP_IMAGES_INTRO)
    builder.paragraph(copy.TOP_IMAGES_TABLE_INTRO)
    rows = [
        (
            row.get("name") or "",
            row.get("repository_uri") or row.get("digest") or "",
            row.get("critical", 0),
            row.get("high", 0),
            row.get("medium", 0),
            row.get("low", 0),
            row.get("vulnerabilities", 0),
        )
        for row in dataset.get("top_vulnerable_images") or ()
    ]
    builder.table(
        ("Imagem", "Repositório / Digest", "Crítica", "Alta", "Média", "Baixa", "Total"),
        rows,
        widths=(1500, 2700, 950, 950, 950, 950, 1200),
        left_columns=frozenset({0, 1}),
    )
    builder.source_note(dataset, "cloud_top_images", enabled=show_source_filters)


def render_top_critical(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    builder.heading("3.4. Principais Vulnerabilidades Críticas (TOP 5 CVEs)", 2)
    builder.paragraph(copy.TOP_CRITICAL_INTRO)
    builder.paragraph(copy.TOP_CRITICAL_PRIORITY)
    rows = [
        (
            index,
            item.get("cve") or "",
            item.get("vpr_display") or "N/D",
            item.get("cvss_display") or "N/D",
            SEVERITY_LABELS.get(str(item.get("severity") or ""), item.get("severity") or ""),
            item.get("affected_assets", 0),
            ", ".join(item.get("components") or ()) or "N/D",
        )
        for index, item in enumerate(dataset.get("top_critical_cves") or (), start=1)
    ]
    builder.table(
        ("Rank", "CVE", "VPR", "CVSS", "Severidade", "Ativos afetados", "Componente / Produto"),
        rows,
        widths=(600, 1450, 700, 700, 1200, 1250, 3300),
        left_columns=frozenset({1, 6}),
        empty_message=copy.EMPTY_CRITICAL_MONTH,
    )
    builder.source_note(dataset, "cloud_top_critical_cves", enabled=show_source_filters)


def _translated_description(
    text: str,
    translator: TextTranslator | None,
) -> tuple[tuple[str, ...], bool]:
    if translator is None:
        return split_translation_chunks(text, max_chars=900), False
    try:
        return translate_in_chunks(text, translator, max_chars=900), False
    except (TypeError, ValueError, RuntimeError):
        return split_translation_chunks(text, max_chars=900), True


def render_critical_details(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    translator: TextTranslator | None,
) -> None:
    critical = dataset.get("top_critical_cves") or ()
    corrections = {
        str(item.get("cve") or ""): item
        for item in dataset.get("top_correctable_vulnerabilities") or ()
        if isinstance(item, Mapping)
    }
    for index, item in enumerate(critical[:5], start=1):
        cve = str(item.get("cve") or "CVE não informada")
        builder.standard_heading(f"3.4.{index}. {cve}", 3)
        builder.standard_paragraph(
            " | ".join(
                (
                    f"VPR: {item.get('vpr_display') or 'N/D'}",
                    f"CVSS: {item.get('cvss_display') or 'N/D'}",
                    "Severidade: "
                    + SEVERITY_LABELS.get(
                        str(item.get("severity") or ""),
                        str(item.get("severity") or "N/D"),
                    ),
                )
            ),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        builder.standard_paragraph(
            "Componente / produto: "
            + (", ".join(item.get("components") or ()) or "N/D"),
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        builder.standard_heading("Descrição:", 4)
        description = str(item.get("description") or "").strip()
        if description:
            chunks, failed = _translated_description(description, translator)
            for chunk in chunks:
                builder.standard_paragraph(chunk)
            if failed:
                builder.standard_paragraph(
                    copy.TRANSLATION_UNAVAILABLE,
                    color=MID_GRAY,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )
        else:
            builder.standard_paragraph(copy.SOURCE_UNAVAILABLE, color=MID_GRAY)
        correction = corrections.get(cve)
        builder.standard_heading("Correção ou contramedida recomendada:", 4)
        if correction and correction.get("recommended_action"):
            action_chunks, action_failed = _translated_description(
                str(correction["recommended_action"]),
                translator,
            )
            for chunk in action_chunks:
                builder.standard_paragraph(chunk)
            if action_failed:
                builder.standard_paragraph(
                    copy.TRANSLATION_UNAVAILABLE,
                    color=MID_GRAY,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )
            builder.standard_paragraph(
                "Tipo de correção: "
                + str(correction.get("correction_type_display") or "Não determinado"),
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
        else:
            builder.standard_paragraph(
                "Não determinada pelos dados coletados neste mês.",
                color=MID_GRAY,
            )
        builder.standard_heading("Ativos afetados:", 4)
        asset_rows = []
        for asset in item.get("assets") or ():
            if not isinstance(asset, Mapping):
                continue
            kind = "Imagem" if asset.get("kind") == "container_image" else "Máquina virtual"
            address = (
                asset.get("repository_uri")
                or asset.get("digest")
                or ", ".join(asset.get("ip_addresses") or ())
            )
            asset_rows.append(
                (
                    kind,
                    _asset_identity(asset),
                    address or "",
                    asset.get("account_id") or "",
                    ", ".join(asset.get("components") or ()) or "N/D",
                )
            )
        builder.table(
            ("Tipo", "Ativo", "IP / Repositório", "Conta", "Componente"),
            asset_rows,
            widths=(1200, 2100, 2500, 1500, 1900),
            left_columns=frozenset({0, 1, 2, 3, 4}),
        )


def render_top_correctable(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
    translator: TextTranslator | None,
) -> None:
    builder.heading("3.5. Principais Vulnerabilidades com Correção Disponível", 2)
    rows = [
        (
            item.get("cve") or "",
            item.get("vpr_display") or "N/D",
            item.get("cvss_display") or "N/D",
            SEVERITY_LABELS.get(
                str(item.get("severity") or ""),
                item.get("severity") or "",
            ),
            item.get("affected_assets", 0),
            item.get("correction_type_display") or "Não determinado",
        )
        for item in dataset.get("top_correctable_vulnerabilities") or ()
    ]
    builder.table(
        (
            "CVE",
            "VPR",
            "CVSS",
            "Severidade",
            "Ativos afetados",
            "Tipo de correção",
        ),
        rows,
        widths=(1250, 700, 700, 1200, 1450, 3900),
        left_columns=frozenset({0, 5}),
        empty_message=copy.EMPTY_CORRECTABLE_MONTH,
    )
    builder.source_note(
        dataset,
        "cloud_top_correctable",
        enabled=show_source_filters,
    )

def render_dashboard(builder: CloudDocumentBuilder, dataset: Mapping[str, Any]) -> None:
    builder.heading("3.6. Painel de Controle (Dashboards) – Informações Rápidas", 2)
    builder.paragraph(copy.DASHBOARD_INTRO)
    builder.heading("3.6.1. Proteção de Workloads (Workload Protection)", 3)
    builder.paragraph(copy.WORKLOAD_STATUS)
    workload = dataset.get("workload_status")
    if not isinstance(workload, Mapping) or not isinstance(
        workload.get("by_max_severity"),
        Mapping,
    ):
        builder.paragraph(copy.SOURCE_UNAVAILABLE, color=MID_GRAY)
    else:
        counts = workload["by_max_severity"]
        rows = [
            (SEVERITY_LABELS[severity], counts.get(severity, 0))
            for severity in ("CRITICAL", "HIGH", "MEDIUM", "LOW", "NONE")
        ]
        builder.table(
            ("Maior severidade no workload", "Máquinas virtuais"),
            rows,
            widths=(5800, 3400),
            left_columns=frozenset({0}),
        )
    builder.heading("3.6.2. Status dos Sistemas Operacionais", 3)
    builder.paragraph(copy.OPERATING_SYSTEM_STATUS)
    builder.paragraph(
        "Inserir print da plataforma aqui",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )


def _source_status(
    dataset: Mapping[str, Any],
    *names: str,
) -> str:
    statuses = dataset.get("source_status")
    statuses = statuses if isinstance(statuses, Mapping) else {}
    for name in names:
        value = statuses.get(name)
        if value is not None:
            return str(value).upper()
    return "UNKNOWN"


def _capability_available(
    dataset: Mapping[str, Any],
    source: str,
) -> bool:
    capabilities = dataset.get("capabilities")
    capabilities = capabilities if isinstance(capabilities, Mapping) else {}
    sources = capabilities.get("sources")
    sources = sources if isinstance(sources, Mapping) else {}
    value = sources.get(source)
    if isinstance(value, Mapping):
        value = value.get("status")
    return str(value or "").upper() == "AVAILABLE"


def cloud_posture_available(dataset: Mapping[str, Any]) -> bool:
    return _capability_available(dataset, "findings") and bool(
        dataset.get("top_posture_findings")
    )


def render_executive_overview(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    chart_dir: str | Path,
) -> None:
    builder.heading("3.1.1. Resumo Executivo do Período", 3)
    builder.paragraph(copy.EXECUTIVE_OVERVIEW)
    overview = dataset.get("overview")
    overview = overview if isinstance(overview, Mapping) else {}
    severity = overview.get("severity_counts")
    severity = severity if isinstance(severity, Mapping) else {}
    rows = (
        ("Ativos Cloud", overview.get("assets", 0)),
        ("Máquinas virtuais", overview.get("virtual_machines", 0)),
        ("Imagens de contêiner", overview.get("container_images", 0)),
        ("CVEs únicas", overview.get("unique_cves", 0)),
        ("Ocorrências de vulnerabilidade", overview.get("vulnerability_occurrences", 0)),
        ("Vulnerabilidades com correção mapeada", len(dataset.get("top_correctable_vulnerabilities") or ())),
    )
    builder.table(
        ("Indicador", "Resultado"),
        rows,
        widths=(6500, 2700),
        left_columns=frozenset({0}),
    )
    chart = render_severity_chart(
        severity,
        Path(chart_dir) / "cloud-severity-overview.png",
    )
    builder.image(
        chart,
        alt_text="Distribuição das ocorrências Cloud por severidade",
    )


def render_components_products(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    builder.heading("3.7. Componentes e Produtos em Maior Risco", 2)
    builder.paragraph(copy.COMPONENTS_PRODUCTS)
    rows = [
        (
            item.get("component") or "Não informado",
            item.get("affected_assets", 0),
            item.get("vulnerabilities", 0),
            item.get("occurrences", 0),
        )
        for item in dataset.get("top_components") or ()
        if isinstance(item, Mapping)
    ]
    builder.table(
        ("Componente / Produto", "Ativos afetados", "CVEs únicas", "Ocorrências"),
        rows,
        widths=(4800, 1500, 1400, 1500),
        left_columns=frozenset({0}),
    )
    builder.source_note(dataset, "cloud_top_components", enabled=show_source_filters)


def render_cloud_posture(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    builder.heading("3.8. Postura de Segurança em Nuvem", 2)
    builder.paragraph(copy.CLOUD_POSTURE)
    rows = [
        (
            item.get("policy") or "",
            item.get("category") or "",
            SEVERITY_LABELS.get(str(item.get("severity") or ""), item.get("severity") or ""),
            item.get("provider") or "N/D",
            item.get("findings", 0),
            item.get("affected_resources", 0),
        )
        for item in dataset.get("top_posture_findings") or ()
        if isinstance(item, Mapping)
    ]
    builder.table(
        ("Política", "Categoria", "Severidade", "Provedor", "Achados", "Recursos"),
        rows,
        widths=(3000, 1600, 1200, 1400, 1000, 1000),
        left_columns=frozenset({0, 1, 3}),
    )
    builder.source_note(dataset, "cloud_posture", enabled=show_source_filters)


def render_vulnerability_aging(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    chart_dir: str | Path,
    show_source_filters: bool,
) -> None:
    builder.heading("3.9. Envelhecimento das Vulnerabilidades", 2)
    builder.paragraph(copy.VULNERABILITY_AGING)
    if _source_status(dataset, "vulnerability_lifecycle", "lifecycle") == "UNAVAILABLE":
        builder.paragraph(copy.SOURCE_UNAVAILABLE, color=MID_GRAY)
        return
    aging = dataset.get("aging")
    if not isinstance(aging, Mapping):
        builder.paragraph(copy.SOURCE_UNAVAILABLE, color=MID_GRAY)
        return
    rows = (
        ("0 a 30 dias", aging.get("0-30", 0)),
        ("31 a 60 dias", aging.get("31-60", 0)),
        ("61 a 90 dias", aging.get("61-90", 0)),
        ("91 a 180 dias", aging.get("91-180", 0)),
        ("Mais de 180 dias", aging.get(">180", 0)),
        ("Data indisponível", aging.get("data_indisponivel", 0)),
    )
    builder.table(
        ("Faixa", "Vulnerabilidades abertas"),
        rows,
        widths=(6000, 3200),
        left_columns=frozenset({0}),
    )
    chart = render_aging_chart(
        aging,
        Path(chart_dir) / "cloud-vulnerability-aging.png",
    )
    builder.image(
        chart,
        alt_text="Distribuição das vulnerabilidades Cloud abertas por faixa de idade",
    )
    builder.source_note(dataset, "cloud_aging", enabled=show_source_filters)


def render_remediation_performance(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    show_source_filters: bool,
) -> None:
    builder.heading("3.10. Desempenho de Remediação", 2)
    builder.paragraph(copy.REMEDIATION_PERFORMANCE)
    if _source_status(dataset, "vulnerability_lifecycle", "lifecycle") == "UNAVAILABLE":
        builder.paragraph(copy.SOURCE_UNAVAILABLE, color=MID_GRAY)
        return
    performance = dataset.get("remediation_performance")
    if not isinstance(performance, Mapping):
        builder.paragraph(copy.SOURCE_UNAVAILABLE, color=MID_GRAY)
        return
    average = performance.get("average_resolution_days")
    builder.table(
        ("Indicador", "Resultado"),
        (
            ("Vulnerabilidades resolvidas no período", performance.get("resolved", 0)),
            ("Tempo médio até resolução", f"{average} dias" if average is not None else "N/D"),
        ),
        widths=(6500, 2700),
        left_columns=frozenset({0}),
    )
    builder.source_note(
        dataset,
        "cloud_remediation_performance",
        enabled=show_source_filters,
    )


def _current_period_history(dataset: Mapping[str, Any]) -> Mapping[str, Any]:
    period = dataset.get("period")
    period = period if isinstance(period, Mapping) else {}
    overview = dataset.get("overview")
    overview = overview if isinstance(overview, Mapping) else {}
    period_id = str(period.get("period_id") or "").strip()
    label = period_id or "Período atual"
    parts = period_id.split("-", maxsplit=1)
    if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
        month = int(parts[1])
        month_labels = (
            "",
            "Jan",
            "Fev",
            "Mar",
            "Abr",
            "Mai",
            "Jun",
            "Jul",
            "Ago",
            "Set",
            "Out",
            "Nov",
            "Dez",
        )
        if 1 <= month <= 12:
            label = f"{month_labels[month]}/{parts[0][-2:]}"
    return {
        "period_id": period_id,
        "label": label,
        "availability": "AVAILABLE",
        "overview": overview,
    }


def render_monthly_evolution(
    builder: CloudDocumentBuilder,
    dataset: Mapping[str, Any],
    *,
    chart_dir: str | Path,
) -> bool:
    history = [
        item
        for item in dataset.get("history") or ()
        if isinstance(item, Mapping)
    ]
    if not history:
        builder.page_break()
    builder.heading("3.11. Evolução Mensal", 2)
    builder.paragraph(copy.MONTHLY_EVOLUTION)
    if not history:
        history = [_current_period_history(dataset)]
        builder.paragraph(
            "Sem histórico mensal anterior, este primeiro ponto representa "
            "somente a fotografia atual do período.",
            color=MID_GRAY,
        )
    points = normalize_history_series(history)
    if not points:
        builder.paragraph(copy.HISTORY_UNAVAILABLE, color=MID_GRAY)
        return False
    builder.table(
        ("Mês", "Ocorrências de vulnerabilidade"),
        [
            (point.label, point.value if point.value is not None else "N/D")
            for point in points
        ],
        widths=(4200, 5000),
        left_columns=frozenset({0}),
    )
    chart = render_monthly_history_chart(
        history,
        Path(chart_dir) / "cloud-monthly-evolution.png",
    )
    if chart is None:
        builder.paragraph(copy.HISTORY_UNAVAILABLE, color=MID_GRAY)
        return False
    builder.image(
        chart,
        alt_text=(
            "Evolução mensal das ocorrências Cloud; meses indisponíveis "
            "aparecem como lacunas"
        ),
        width_cm=12.5,
    )
    return True

def render_conclusion(builder: CloudDocumentBuilder) -> None:
    builder.heading("4. Conclusão", 1)
    builder.paragraph(copy.CONCLUSION_OVERVIEW)
    builder.paragraph(copy.CONCLUSION_VALUE_INTRO)
    for item in copy.CONCLUSION_VALUES:
        builder.bullet(item)
    builder.paragraph(copy.CONCLUSION_PLATFORM)
    builder.paragraph(copy.CONCLUSION_SUMMARY)


__all__ = [
    "CloudDocumentBuilder",
    "cloud_posture_available",
    "render_cloud_overview",
    "render_cloud_posture",
    "render_components_products",
    "render_conclusion",
    "render_critical_details",
    "render_dashboard",
    "render_document_control",
    "render_executive_overview",
    "render_introduction",
    "render_monthly_evolution",
    "render_remediation_performance",
    "render_top_correctable",
    "render_top_critical",
    "render_top_hosts",
    "render_top_images",
    "render_vulnerability_aging",
]
