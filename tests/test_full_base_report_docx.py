from __future__ import annotations

import json
import re
import tempfile
import unittest
import zipfile
from dataclasses import replace
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from tenable_reports.config.profile import load_client_profile
from tenable_reports.presentation import base_report_docx as base
from tenable_reports.presentation.base_report_docx import ASSET_HEADERS
from tenable_reports.presentation.full_base_report_docx import (
    FULL_TEMPLATE_VERSION,
    _description_chunks,
    _simple_table,
    generate_full_base_report,
)


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "templates/corporate/base-v1.docx"
ASSETS = ROOT / "templates/corporate/assets"
FIXTURE = ROOT / "tests/fixtures/report-dataset-phase5.json"
PROFILE = ROOT / "clients/examples/client-profile.json"


def all_document_text(document):
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for container in (section.header, section.footer):
            chunks.extend(paragraph.text for paragraph in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def _cell_fill(cell) -> str | None:
    shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return None if shading is None else shading.get(qn("w:fill"))


class FullBaseReportDocxTests(unittest.TestCase):
    def test_risk_band_labels_in_highlight_tables_use_approved_palette(self) -> None:
        document = Document()
        aging = _simple_table(
            document,
            ("", "90+ Dias"),
            (("Crítica", 1), ("Alta", 2), ("Média", 3), ("Baixa", 4),
             ("Texto crítico para análise", 5)),
        )
        self.assertEqual(
            [_cell_fill(row.cells[0]) for row in aging.rows[1:5]],
            [base.CRITICAL, base.HIGH, base.MEDIUM, base.LOW],
        )
        self.assertNotIn(
            _cell_fill(aging.rows[5].cells[0]),
            {base.CRITICAL, base.HIGH, base.MEDIUM, base.LOW},
        )

        cvss = _simple_table(
            document,
            ("", "Total"),
            (("CVSSv3 10.0", 1), ("CVSSv3 7.0 - 9.9", 2),
             ("CVSSv3 4.0 - 6.9", 3), ("CVSSv3 0.0 - 3.9", 4)),
        )
        self.assertEqual(
            [_cell_fill(row.cells[0]) for row in cvss.rows[1:]],
            [base.CRITICAL, base.HIGH, base.MEDIUM, base.LOW],
        )

        heatmap = _simple_table(
            document,
            (
                "",
                "Baixo (VPR 0.0-3.9)",
                "Médio (VPR 4.0-6.9)",
                "Alto (VPR7.0-8.9)",
                "Crítico (VPR 9.0-10)",
            ),
            (("CVSSv3 Baixo (0.0-3.9)", 1, 2, 3, 4),
             ("CVSSv3 Médio (4.0-6.9)", 1, 2, 3, 4),
             ("CVSSv3 Alto (VPR7.0-8.9)", 1, 2, 3, 4),
             ("CVSSv3 Crítico (VPR 9.0-10)", 1, 2, 3, 4)),
        )
        self.assertEqual(
            [_cell_fill(cell) for cell in heatmap.rows[0].cells[1:]],
            [base.LOW, base.MEDIUM, base.HIGH, base.CRITICAL],
        )
        self.assertEqual(
            [_cell_fill(row.cells[0]) for row in heatmap.rows[1:]],
            [base.LOW, base.MEDIUM, base.HIGH, base.CRITICAL],
        )

        rating = _simple_table(
            document,
            (
                "RATING 10.0 - 9.0",
                "RATING 8.9-7.0",
                "RATING 6.9-4.0",
                "RATING 3.9-0.1",
            ),
            ((1, 2, 3, 4),),
        )
        self.assertEqual(
            [_cell_fill(cell) for cell in rating.rows[0].cells],
            [base.CRITICAL, base.HIGH, base.MEDIUM, base.LOW],
        )

    def test_compact_vulnerability_table_displays_missing_and_numeric_zero_vpr_as_zero(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            dataset = json.loads(FIXTURE.read_text(encoding="utf-8"))
            dataset["top_open_vulnerabilities"][0]["vpr_score"] = 0
            dataset["top_open_vulnerabilities"][1]["vpr_score"] = None
            dataset_path = Path(directory) / "vpr-zero.json"
            dataset_path.write_text(json.dumps(dataset), encoding="utf-8")
            output = Path(directory) / "vpr-zero.docx"

            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=dataset_path,
                profile=load_client_profile(PROFILE),
                output_path=output,
                assets_dir=ASSETS,
                mask_sensitive=True,
            )

            rows_by_plugin_id = {
                row.cells[0].text: row
                for table in Document(output).tables
                if table.rows
                and tuple(cell.text for cell in table.rows[0].cells)
                == ("Plugin ID", "Nome", "Família OS", "Severidade", "Total", "VPR")
                for row in table.rows[1:]
            }
            self.assertEqual(rows_by_plugin_id["900001"].cells[5].text, "0")
            self.assertEqual(rows_by_plugin_id["900002"].cells[5].text, "0")

    def test_full_sanitised_report_contains_contractual_sections(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "full.docx"
            result = generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=FIXTURE,
                profile=load_client_profile(PROFILE),
                output_path=output,
                assets_dir=ASSETS,
                mask_sensitive=True,
            )
            self.assertEqual(result.template_version, FULL_TEMPLATE_VERSION)
            self.assertEqual(result.top_open_rows, 5)
            self.assertTrue(result.masked_sensitive_fields)

            document = Document(output)
            text = all_document_text(document)
            self.assertNotIn("{{", text)
            self.assertIn("CONTROLE DE DOCUMENTO", text)
            self.assertIn("SUMÁRIO", text)
            self.assertIn("3.2. Principais Ativos Vulneráveis", text)
            self.assertIn("VULNERABILIDADES E SUAS CORREÇÕES", text)
            self.assertIn(
                "6.4. Vulnerabilidades WEB e Suas Correções e/ou Contramedidas Recomendadas",
                text,
            )
            self.assertIn("SUA MELHOR ALIADA NA JORNADA DA PROTEÇÃO DIGITAL.", text)
            self.assertNotIn("METODOLOGIA, QUALIDADE E LIMITAÇÕES", text)
            self.assertNotIn("RELATÓRIO-BASE CONCLUÍDO", text)
            self.assertNotIn("Sem dados disponíveis", text)
            self.assertGreaterEqual(len(document.tables), 20)

            headers = [tuple(cell.text for cell in table.rows[0].cells) for table in document.tables]
            self.assertIn(ASSET_HEADERS, headers)
            self.assertNotIn(("Output",), headers)
            for header in headers:
                self.assertNotIn("Output", header)

            asset_table = next(table for table in document.tables if tuple(cell.text for cell in table.rows[0].cells) == ASSET_HEADERS)
            for row in asset_table.rows[1:]:
                self.assertEqual(row.cells[0].text, "")
                self.assertEqual(row.cells[1].text, "")
                self.assertLessEqual(int(row.cells[7].text), int(row.cells[6].text))

            host_header = ("ASSET NAME", "IP", "PORTA", "PROTOCOLO")
            host_tables = [table for table in document.tables if tuple(cell.text for cell in table.rows[0].cells) == host_header]
            self.assertEqual(len(host_tables), 5)
            for table in host_tables:
                for row in table.rows[1:]:
                    self.assertEqual(row.cells[0].text, "")
                    self.assertEqual(row.cells[1].text, "")
            web_tables = [
                table for table in document.tables
                if tuple(cell.text for cell in table.rows[0].cells) == ("URI",)
            ]
            self.assertEqual(len(web_tables), 5)

            self.assertTrue(all(len(paragraph.text) <= 1250 for paragraph in document.paragraphs))
            self.assertAlmostEqual(document.sections[0].page_width.cm, 21.0, places=1)
            self.assertAlmostEqual(document.sections[0].page_height.cm, 29.7, places=1)

            period = next(
                paragraph for paragraph in document.paragraphs
                if paragraph.text.startswith("Período deste relatório")
            )
            self.assertTrue(period.runs[-1].bold)
            overview = next(
                paragraph for paragraph in document.paragraphs
                if paragraph.text.startswith("Segue um Overview")
            )
            self.assertIn("01 e 31 de julho de 2026", overview.text)
            self.assertTrue(any(run.bold and "julho" in run.text for run in overview.runs))

            with zipfile.ZipFile(output) as package:
                xml_parts = {
                    name: package.read(name).decode("utf-8")
                    for name in package.namelist()
                    if name.endswith(".xml")
                }
            all_xml = "\n".join(xml_parts.values())
            self.assertIn(" TOC ", all_xml)
            self.assertIn(" PAGE ", all_xml)
            self.assertIn("w:tblHeader", all_xml)
            self.assertIn("w:numPr", all_xml)
            self.assertNotRegex(all_xml, r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}")
            image_nodes = re.findall(r"<wp:docPr\b[^>]*>", all_xml)
            self.assertTrue(image_nodes)
            self.assertTrue(all('descr="' in node and 'descr=""' not in node for node in image_nodes))

    def test_exploitability_matrix_is_kept_together_across_pages(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            dataset = json.loads(FIXTURE.read_text(encoding="utf-8"))
            labels = (
                "Exploitable",
                "Malware",
                "Core Impact",
                "Canvas",
                "D2 Elliot",
                "ExploitHub",
                "Metasploit",
            )
            dataset["metrics"]["by_exploit_framework"] = [
                {
                    "framework": label,
                    "total": 1,
                    "critical": 0,
                    "high": 1,
                    "medium": 0,
                }
                for label in labels
            ]
            dataset_path = Path(directory) / "dataset.json"
            dataset_path.write_text(json.dumps(dataset), encoding="utf-8")
            output = Path(directory) / "frameworks.docx"
            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=dataset_path,
                profile=load_client_profile(PROFILE),
                output_path=output,
                mask_sensitive=True,
            )

            document = Document(output)
            table = next(
                table
                for table in document.tables
                if len(table.rows) > 1 and table.cell(1, 0).text == "Exploitable"
            )
            for row in table.rows[:-1]:
                self.assertTrue(all(cell.paragraphs[0].paragraph_format.keep_with_next for cell in row.cells))

    def test_output_column_is_controlled_by_profile(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            profile = load_client_profile(PROFILE)
            output = Path(directory) / "without-output.docx"
            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=FIXTURE,
                profile=profile,
                output_path=output,
                mask_sensitive=True,
            )
            headers = [
                tuple(cell.text for cell in table.rows[0].cells)
                for table in Document(output).tables
            ]
            self.assertNotIn("Output", {cell for row in headers for cell in row})

    def test_source_filter_notes_are_opt_in_and_sanitized(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            dataset = json.loads(FIXTURE.read_text(encoding="utf-8"))
            dataset["table_provenance"] = {"version": "table-provenance-v1", "tables": {
                "top_assets": {
                    "source": "Tenable VM", "period_start_at": "2026-07-01T00:00:00Z",
                    "period_end_at": "2026-08-01T00:00:00Z", "states": ["OPEN"],
                    "severities": ["CRITICAL", "HIGH"], "tag_value": "NÃO DEVE APARECER",
                    "secret_key": "NÃO DEVE VAZAR",
                },
            }}
            dataset_path = Path(directory) / "dataset.json"
            dataset_path.write_text(json.dumps(dataset), encoding="utf-8")
            profile = load_client_profile(PROFILE)
            profile = replace(
                profile,
                presentation=replace(profile.presentation, show_source_filters=True),
            )
            output = Path(directory) / "filters.docx"
            generate_full_base_report(
                template_path=TEMPLATE, dataset_path=dataset_path, profile=profile,
                output_path=output, mask_sensitive=True,
            )
            text = all_document_text(Document(output))
            self.assertIn("Validação rápida na Tenable: fonte Tenable VM", text)
            self.assertNotIn("NÃO DEVE APARECER", text)
            self.assertNotIn("NÃO DEVE VAZAR", text)

    def test_source_filter_notes_cover_base_report_data_tables(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            dataset = json.loads(FIXTURE.read_text(encoding="utf-8"))
            common = {
                "view": "Explore > Findings > Vulnerabilities",
                "period_start_at": "2026-07-01T00:00:00Z",
                "period_end_at": "2026-08-01T00:00:00Z",
                "severities": ["CRITICAL", "HIGH", "MEDIUM", "LOW"],
            }
            dataset["table_provenance"] = {
                "version": "table-provenance-v1",
                "tables": {
                    "overview": {**common, "validation_queries": [
                        {"label": "Não mitigadas", "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"]},
                        {"label": "Mitigadas", "states": ["FIXED"], "date_fields": ["Last Fixed"]},
                    ], "rule": "Resumo geral do período"},
                    "top_assets": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Agrupar por ativo"},
                    "top_open_vulnerabilities": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Agrupar por Plugin ID"},
                    "top_fixed_vulnerabilities": {**common, "states": ["FIXED"], "date_fields": ["Last Fixed"], "rule": "Agrupar por Plugin ID mitigado"},
                    "top_resurfaced_vulnerabilities": {**common, "states": ["REOPENED"], "date_fields": ["Resurfaced Date"], "rule": "Agrupar por Plugin ID ressurgido"},
                    "was_applications": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Agrupar por aplicação WEB"},
                    "was_top_vulnerabilities": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Agrupar por Plugin ID WEB"},
                    "was_owasp": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Agrupar pela categoria OWASP 2021"},
                    "top_web_vulnerabilities": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Hosts por Plugin ID WEB"},
                    "by_operating_system": {**common, "validation_queries": [
                        {"label": "Não mitigadas", "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"]},
                        {"label": "Mitigadas", "states": ["FIXED"], "date_fields": ["Last Fixed"]},
                    ], "rule": "Agrupar por sistema operacional"},
                    "by_cvss": {**common, "validation_queries": [
                        {"label": "Não mitigadas", "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"]},
                        {"label": "Mitigadas", "states": ["FIXED"], "date_fields": ["Last Fixed"]},
                    ], "rule": "Faixas de CVSS v3"},
                    "cvss_vpr_matrix": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Cruzar CVSS v3 e VPR"},
                    "vpr_rating": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Faixas de VPR"},
                    "state_summary": {**common, "validation_queries": [
                        {"label": "Novo", "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen", "First Seen"]},
                        {"label": "Ativo/não mitigado", "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"]},
                        {"label": "Corrigido", "states": ["FIXED"], "date_fields": ["Last Fixed"]},
                        {"label": "Ressurgido", "states": ["REOPENED"], "date_fields": ["Last Seen", "Resurfaced Date"]},
                    ], "rule": "Separar pelo estado do finding"},
                    "aging_by_severity": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Faixas calculadas desde First Seen"},
                    "by_exploit_framework": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Flags individuais de framework"},
                },
            }
            dataset.setdefault("metrics", {})["by_exploit_framework"] = [
                {
                    "framework": "Metasploit",
                    "total": 3,
                    "critical": 1,
                    "high": 2,
                    "medium": 0,
                }
            ]
            dataset.setdefault("was", {}).setdefault("owasp", {})["A01"] = [
                {"plugin_id": 980001, "name": "Teste OWASP", "instances": 1}
            ]
            dataset_path = Path(directory) / "dataset.json"
            dataset_path.write_text(json.dumps(dataset), encoding="utf-8")
            profile = load_client_profile(PROFILE)
            profile = replace(
                profile,
                presentation=replace(profile.presentation, show_source_filters=True),
            )
            output = Path(directory) / "all-base-filters.docx"
            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=dataset_path,
                profile=profile,
                output_path=output,
                mask_sensitive=True,
            )
            text = all_document_text(Document(output))
            for marker in (
                "Resumo geral do período",
                "Plugin ID = 900001",
                "OWASP 2021 = A01",
                "Faixas de CVSS v3",
                "Faixas calculadas desde First Seen",
                "Flags individuais de framework",
                "State = Fixed; Severity = Critical, High, Medium, Low; Last Fixed = 01/07/2026 a 31/07/2026",
                "State = Resurfaced; Severity = Critical, High, Medium, Low; Resurfaced Date = 01/07/2026 a 31/07/2026",
            ):
                self.assertIn(marker, text)

    def test_output_column_can_be_enabled_for_vm_and_web(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            profile = load_client_profile(PROFILE)
            profile = replace(
                profile,
                presentation=replace(
                    profile.presentation,
                    vm_top5_include_output=True,
                    was_top5_include_output=True,
                ),
            )
            dataset = json.loads(FIXTURE.read_text(encoding="utf-8"))
            dataset["source_coverage"]["plugin_output_included"] = True
            dataset["source_coverage"]["was_plugin_output_included"] = True
            dataset_path = Path(directory) / "with-output.json"
            dataset_path.write_text(json.dumps(dataset), encoding="utf-8")
            output = Path(directory) / "with-output.docx"
            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=dataset_path,
                profile=profile,
                output_path=output,
                mask_sensitive=True,
            )
            headers = [
                tuple(cell.text for cell in table.rows[0].cells)
                for table in Document(output).tables
            ]
            self.assertEqual(headers.count(("ASSET NAME", "IP", "PORTA", "PROTOCOLO", "Output")), 5)
            self.assertEqual(headers.count(("URI", "Plugin Output")), 5)

    def test_output_column_fails_when_plugin_output_was_not_collected(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            profile = load_client_profile(PROFILE)
            profile = replace(
                profile,
                presentation=replace(
                    profile.presentation,
                    vm_top5_include_output=True,
                ),
            )
            with self.assertRaisesRegex(ValueError, "nao foi coletado"):
                generate_full_base_report(
                    template_path=TEMPLATE,
                    dataset_path=FIXTURE,
                    profile=profile,
                    output_path=Path(directory) / "must-not-exist.docx",
                    mask_sensitive=True,
                )

    def test_long_description_is_split_without_truncation(self) -> None:
        text = " ".join(f"Sentença {index} com conteúdo técnico." for index in range(200))
        chunks = _description_chunks(text, max_chars=300)
        self.assertGreater(len(chunks), 1)
        self.assertTrue(all(len(chunk) <= 300 for chunk in chunks))
        self.assertEqual(" ".join(chunks), text)

    def test_empty_framework_and_owasp_categories_have_monthly_message(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "empty-sections.docx"
            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=FIXTURE,
                profile=load_client_profile(PROFILE),
                output_path=output,
                assets_dir=ASSETS,
                mask_sensitive=True,
            )
            text = all_document_text(Document(output))
            self.assertIn(
                "Neste mês não foram identificadas vulnerabilidades exploráveis por frameworks conhecidos.",
                text,
            )
            self.assertEqual(
                text.count("Neste mês não foram identificadas vulnerabilidades relacionadas a esta categoria OWASP."),
                10,
            )

    def test_empty_framework_does_not_add_an_orphan_validation_note(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            dataset = json.loads(FIXTURE.read_text(encoding="utf-8"))
            dataset["table_provenance"] = {"tables": {
                "by_exploit_framework": {
                    "view": "Explore > Findings > Vulnerabilities",
                    "states": ["OPEN", "REOPENED"],
                    "date_fields": ["Last Seen"],
                    "rule": "MARCADOR-FRAMEWORK-SEM-DADOS",
                }
            }}
            dataset_path = Path(directory) / "empty-framework.json"
            dataset_path.write_text(json.dumps(dataset), encoding="utf-8")
            profile = load_client_profile(PROFILE)
            profile = replace(
                profile,
                presentation=replace(profile.presentation, show_source_filters=True),
            )
            output = Path(directory) / "empty-framework.docx"
            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=dataset_path,
                profile=profile,
                output_path=output,
                assets_dir=ASSETS,
                mask_sensitive=True,
            )

            text = all_document_text(Document(output))
            self.assertIn(
                "Neste mês não foram identificadas vulnerabilidades exploráveis por frameworks conhecidos.",
                text,
            )
            self.assertNotIn("MARCADOR-FRAMEWORK-SEM-DADOS", text)

    def test_empty_web_top5_has_monthly_message(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            dataset = json.loads(FIXTURE.read_text(encoding="utf-8"))
            dataset["top_web_vulnerabilities"] = []
            dataset.setdefault("was", {})["top_vulnerabilities"] = []
            dataset_path = Path(directory) / "empty-web.json"
            dataset_path.write_text(json.dumps(dataset), encoding="utf-8")
            output = Path(directory) / "empty-web.docx"
            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=dataset_path,
                profile=load_client_profile(PROFILE),
                output_path=output,
                assets_dir=ASSETS,
                mask_sensitive=True,
            )
            self.assertIn(
                "Neste mês não foram identificadas vulnerabilidades WEB não mitigadas para detalhamento neste item.",
                all_document_text(Document(output)),
            )

    def test_unavailable_web_collection_is_not_presented_as_zero_findings(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            dataset = json.loads(FIXTURE.read_text(encoding="utf-8"))
            dataset["was"] = {
                "availability": "NOT_COLLECTED",
                "applications": [],
                "top_vulnerabilities": [],
                "owasp": {},
            }
            dataset["top_web_vulnerabilities"] = []
            dataset_path = Path(directory) / "unavailable-web.json"
            dataset_path.write_text(json.dumps(dataset), encoding="utf-8")
            output = Path(directory) / "unavailable-web.docx"

            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=dataset_path,
                profile=load_client_profile(PROFILE),
                output_path=output,
                assets_dir=ASSETS,
                mask_sensitive=True,
            )

            text = all_document_text(Document(output))
            self.assertEqual(text.count("Não foi possível concluir a coleta WEB"), 1)
            self.assertNotIn(
                "Neste mês não foram identificadas vulnerabilidades WEB não mitigadas",
                text,
            )
            self.assertNotIn(
                "Neste mês não foram identificadas vulnerabilidades relacionadas a esta categoria OWASP.",
                text,
            )

    def test_every_empty_table_receives_a_monthly_message(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "empty-tables.docx"
            generate_full_base_report(
                template_path=TEMPLATE,
                dataset_path=FIXTURE,
                profile=load_client_profile(PROFILE),
                output_path=output,
                assets_dir=ASSETS,
                mask_sensitive=True,
            )
            document = Document(output)
            empty_tables = sum(
                len(table.rows) == 1 and any(cell.text.strip() for cell in table.rows[0].cells)
                for table in document.tables
            )
            text = all_document_text(document)
            generic = text.count("Neste mês não foram identificados registros para este item.")
            framework = text.count(
                "Neste mês não foram identificadas vulnerabilidades exploráveis por frameworks conhecidos."
            )
            self.assertEqual(generic + framework, empty_tables)


if __name__ == "__main__":
    unittest.main()
