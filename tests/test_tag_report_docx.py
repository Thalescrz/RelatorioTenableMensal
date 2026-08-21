from __future__ import annotations

import json
import zipfile
from dataclasses import replace
from pathlib import Path

from docx import Document

from tenable_reports.config.profile import load_client_profile
from tenable_reports.presentation.tag_report_docx import generate_tag_report


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "templates/corporate/base-v1.docx"
FIXTURE = ROOT / "tests/fixtures/report-dataset-phase5.json"
PROFILE = ROOT / "clients/examples/client-profile.json"


def _all_text(document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def _tag_dataset(
    tmp_path: Path,
    *,
    empty: bool = False,
    with_history: bool = False,
    include_comparison: bool = False,
    gap: bool = False,
) -> Path:
    data = json.loads(FIXTURE.read_text(encoding="utf-8"))
    data["document_kind"] = "tag"
    data["tag"] = {
        "tag_uuid": "tag-a",
        "category_uuid": "category-team",
        "category_name": "Equipe",
        "value": "Infraestrutura",
        "include_temporal_comparison": include_comparison,
    }
    data["table_provenance"] = {
        "version": "table-provenance-v1",
        "tables": {
            "top_assets": {
                "source": "Tenable Vulnerability Management",
                "view": "Explore > Findings > Vulnerabilities",
                "states": ["OPEN", "REOPENED"],
                "severities": ["CRITICAL", "HIGH", "MEDIUM", "LOW"],
                "date_fields": ["Last Seen"],
                "period_start_at": data["period"]["start_at"],
                "period_end_at": data["period"]["end_at"],
                "timezone": data["period"]["timezone"],
            },
            "top_open_vulnerabilities": {
                "source": "Tenable Vulnerability Management",
                "view": "Explore > Findings > Vulnerabilities",
                "states": ["OPEN", "REOPENED"],
                "severities": ["CRITICAL", "HIGH", "MEDIUM", "LOW"],
                "date_fields": ["Last Seen"],
                "period_start_at": data["period"]["start_at"],
                "period_end_at": data["period"]["end_at"],
                "timezone": data["period"]["timezone"],
            },
        },
    }
    if empty:
        data["top_assets"] = []
        data["top_open_vulnerabilities"] = []
        data["metrics"]["non_mitigated"]["total"] = 0
        data["metrics"]["non_mitigated"]["by_severity"] = {
            "critical": 0,
            "high": 0,
            "medium": 0,
            "low": 0,
        }
    if with_history:
        def month(
            period_id: str,
            label: str,
            non_mitigated: int,
            mitigated: int,
            new: int,
        ) -> dict:
            return {
                "period_id": period_id,
                "label": label,
                "availability": "AVAILABLE",
                "non_mitigated": non_mitigated,
                "non_mitigated_by_severity": {
                    "critical": non_mitigated // 4,
                    "high": non_mitigated // 4,
                    "medium": non_mitigated // 4,
                    "low": non_mitigated - 3 * (non_mitigated // 4),
                },
                "mitigated": mitigated,
                "mitigated_by_severity": {
                    "critical": 0,
                    "high": mitigated // 2,
                    "medium": mitigated - mitigated // 2,
                    "low": 0,
                },
                "new": new,
                "new_by_severity": {
                    "critical": 0,
                    "high": new // 2,
                    "medium": new - new // 2,
                    "low": 0,
                },
                "top_assets": data["top_assets"][:3],
            }

        data["tag_history_status"] = "AVAILABLE"
        data["tag_history"] = [
            month("2026-01", "Janeiro/2026", 40, 5, 8),
            (
                {
                    "period_id": "2026-02",
                    "label": "Fevereiro/2026",
                    "availability": "UNAVAILABLE",
                }
                if gap
                else month("2026-02", "Fevereiro/2026", 35, 7, 2)
            ),
            month("2026-07", "Julho/2026", 28, 9, 4),
        ]
        if not gap:
            data["tag_comparison"] = {
                "periods": [
                    {
                        "period_id": "2026-02",
                        "label": "Fevereiro/2026",
                        "top_assets": data["top_assets"][:3],
                    },
                    {
                        "period_id": "2026-07",
                        "label": "Julho/2026",
                        "top_assets": data["top_assets"][:3],
                    },
                ]
            }
    path = tmp_path / "tag-dataset.json"
    path.write_text(json.dumps(data), encoding="utf-8")
    return path


def _profile(*, include_output: bool = False, show_filters: bool = False):
    profile = load_client_profile(PROFILE)
    return replace(
        profile,
        presentation=replace(
            profile.presentation,
            vm_top5_include_output=include_output,
            show_source_filters=show_filters,
        ),
    )


def _headers(document) -> list[tuple[str, ...]]:
    return [
        tuple(cell.text for cell in table.rows[0].cells)
        for table in document.tables
        if table.rows
    ]


def _count_document_images(path: Path) -> int:
    with zipfile.ZipFile(path) as package:
        return sum(
            name.startswith("word/media/")
            for name in package.namelist()
        )


def test_tag_report_contains_only_approved_operational_sections(tmp_path: Path) -> None:
    output = tmp_path / "tag.docx"
    result = generate_tag_report(
        template_path=TEMPLATE,
        dataset_path=_tag_dataset(tmp_path),
        profile=_profile(),
        output_path=output,
        mask_sensitive=False,
    )

    document = Document(result.output_path)
    text = _all_text(document)
    headers = _headers(document)
    assert result.tag_uuid == "tag-a"
    assert result.top_asset_rows == 10
    assert result.top_open_rows == 5
    assert "TAG Equipe - Infraestrutura" in text
    assert "Principais Ativos Vulneráveis" in text
    assert "VULNERABILIDADES E SUAS CORREÇÕES" in text
    assert "SENSOR WAS" not in text
    assert "CLOUD SECURITY" not in text
    assert ("ASSET NAME", "IP", "PORTA", "PROTOCOLO") in headers
    assert all("Output" not in header for header in headers)


def test_tag_report_empty_blocks_have_monthly_message(tmp_path: Path) -> None:
    output = tmp_path / "tag-empty.docx"
    generate_tag_report(
        template_path=TEMPLATE,
        dataset_path=_tag_dataset(tmp_path, empty=True),
        profile=_profile(),
        output_path=output,
    )
    text = _all_text(Document(output))

    assert "Neste mês não foram identificados ativos vulneráveis" in text
    assert "Neste mês não foram identificadas vulnerabilidades não mitigadas" in text


def test_output_column_and_tag_validation_filters_are_optional(tmp_path: Path) -> None:
    output = tmp_path / "tag-output.docx"
    generate_tag_report(
        template_path=TEMPLATE,
        dataset_path=_tag_dataset(tmp_path),
        profile=_profile(include_output=True, show_filters=True),
        output_path=output,
    )
    document = Document(output)
    text = _all_text(document)
    host_header = next(header for header in _headers(document) if header[:4] == (
        "ASSET NAME",
        "IP",
        "PORTA",
        "PROTOCOLO",
    ))

    assert host_header[-1] == "Output"
    assert "Tag UUID = tag-a" in text
    assert "Tag = Equipe:Infraestrutura" in text


def test_enabled_tag_comparison_renders_tables_and_five_charts(tmp_path: Path) -> None:
    output = tmp_path / "tag-history.docx"
    result = generate_tag_report(
        template_path=TEMPLATE,
        dataset_path=_tag_dataset(
            tmp_path,
            with_history=True,
            include_comparison=True,
        ),
        profile=_profile(),
        output_path=output,
    )
    document = Document(output)
    headers = _headers(document)

    assert ("Mês", "Crítica", "Alta", "Média", "Baixa", "Total") in headers
    assert result.comparison_rendered is True
    assert _count_document_images(result.output_path) >= 8
    assert _all_text(document).count("Comparativo Mensal") == 1


def test_missing_month_is_unavailable_and_not_plotted_as_zero(tmp_path: Path) -> None:
    output = tmp_path / "tag-history-gap.docx"
    result = generate_tag_report(
        template_path=TEMPLATE,
        dataset_path=_tag_dataset(
            tmp_path,
            with_history=True,
            include_comparison=True,
            gap=True,
        ),
        profile=_profile(),
        output_path=output,
    )
    document = Document(result.output_path)

    assert any(
        "Fevereiro/2026" in row.cells[0].text
        and "Indisponível" in row.cells[-1].text
        for table in document.tables
        for row in table.rows
    )


def test_disabled_comparison_has_no_empty_heading_or_chart(tmp_path: Path) -> None:
    output = tmp_path / "tag-history-disabled.docx"
    result = generate_tag_report(
        template_path=TEMPLATE,
        dataset_path=_tag_dataset(
            tmp_path,
            with_history=True,
            include_comparison=False,
        ),
        profile=_profile(),
        output_path=output,
    )

    assert result.comparison_rendered is False
    assert "Comparativo Mensal" not in _all_text(Document(result.output_path))
