from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from tenable_reports.application.publishing import (
    PublicationDocument,
    PublicationDocumentReplacement,
    create_publication_manifest,
    replace_publication_documents_atomically,
    sha256_file,
)


def _document(path: Path, text: str) -> Path:
    document = Document()
    document.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


class AtomicPublicationTests(unittest.TestCase):
    def _manifest(self, directory: Path) -> tuple[Path, Path, Path, Path]:
        dataset = directory / "dataset.json"
        dataset.write_text("{}", encoding="utf-8")
        base = _document(directory / "base.docx", "Documento original")
        cloud = _document(directory / "cloud.docx", "Documento Cloud")
        manifest = create_publication_manifest(
            output_path=directory / "publication-manifest.json",
            client_id="cliente-a",
            tenant_id="tenant-a",
            run_id="run-a",
            execution_type="AUTOMATIC_MONTHLY",
            period={
                "period_id": "2026-07",
                "mode": "PREVIOUS_CALENDAR_MONTH",
                "timezone": "America/Fortaleza",
                "start_at": "2026-07-01T03:00:00Z",
                "end_at": "2026-08-01T03:00:00Z",
            },
            dataset_path=dataset,
            documents=(
                PublicationDocument(base, "base"),
                PublicationDocument(cloud, "cloud", document_variant="expanded"),
            ),
            history_database=None,
        )
        return manifest, dataset, base, cloud

    def test_invalid_staged_document_preserves_original_and_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as directory_name:
            directory = Path(directory_name)
            manifest, dataset, base, _ = self._manifest(directory)
            original_hash = sha256_file(base)
            original_manifest = manifest.read_bytes()
            invalid = directory / ".repair" / "base.docx"
            invalid.parent.mkdir(parents=True)
            invalid.write_bytes(b"not-a-docx")

            with self.assertRaisesRegex(ValueError, "DOCX invalido"):
                replace_publication_documents_atomically(
                    manifest_path=manifest,
                    dataset_path=dataset,
                    replacements=(PublicationDocumentReplacement(
                        staged_path=invalid,
                        destination=PublicationDocument(base, "base"),
                    ),),
                )

            self.assertEqual(sha256_file(base), original_hash)
            self.assertEqual(manifest.read_bytes(), original_manifest)

    def test_valid_vm_replacement_preserves_cloud_entry(self) -> None:
        with tempfile.TemporaryDirectory() as directory_name:
            directory = Path(directory_name)
            manifest, dataset, base, cloud = self._manifest(directory)
            old_hash = sha256_file(base)
            staged = _document(directory / ".repair" / "base.docx", "Documento reparado")

            replace_publication_documents_atomically(
                manifest_path=manifest,
                dataset_path=dataset,
                replacements=(PublicationDocumentReplacement(
                    staged_path=staged,
                    destination=PublicationDocument(base, "base"),
                ),),
            )

            payload = json.loads(manifest.read_text(encoding="utf-8"))
            documents = payload["documents"]
            base_entry = next(item for item in documents if item["document_kind"] == "base")
            cloud_entry = next(item for item in documents if item["document_kind"] == "cloud")
            self.assertNotEqual(sha256_file(base), old_hash)
            self.assertEqual(base_entry["sha256"], sha256_file(base))
            self.assertEqual(cloud_entry["path"], str(cloud.resolve()))
            self.assertTrue(cloud.is_file())

    def test_commit_failure_restores_documents_and_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as directory_name:
            directory = Path(directory_name)
            manifest, dataset, base, _ = self._manifest(directory)
            original_hash = sha256_file(base)
            original_manifest = manifest.read_bytes()
            staged = _document(
                directory / ".repair" / "base.docx",
                "Documento que nao pode ser confirmado",
            )

            def fail_commit() -> None:
                raise RuntimeError("falha de persistencia")

            with self.assertRaisesRegex(RuntimeError, "falha de persistencia"):
                replace_publication_documents_atomically(
                    manifest_path=manifest,
                    dataset_path=dataset,
                    replacements=(PublicationDocumentReplacement(
                        staged_path=staged,
                        destination=PublicationDocument(base, "base"),
                    ),),
                    commit_callback=fail_commit,
                )

            self.assertEqual(sha256_file(base), original_hash)
            self.assertEqual(manifest.read_bytes(), original_manifest)


if __name__ == "__main__":
    unittest.main()
