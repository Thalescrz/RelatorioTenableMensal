from __future__ import annotations

import json
import re
import zipfile
from dataclasses import replace
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from tenable_reports.config.profile import load_client_profile
from tenable_reports.presentation.cloud_editorial_catalog import (
    approved_cloud_editorial_paragraphs,
)
from tenable_reports.presentation.cloud_report_docx import (
    CloudReportVariant,
    generate_cloud_report,
)


ROOT = Path(__file__).resolve().parents[1]
CLOUD_TEMPLATE = ROOT / "templates/corporate/cloud-base-v1.docx"
PROFILE = ROOT / "clients/examples/client-profile.json"


def _all_text(path: Path) -> str:
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for container in (section.header, section.footer):
            chunks.extend(paragraph.text for paragraph in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def _profile():
    profile = load_client_profile(PROFILE)
    return replace(
        profile,
        cloud_security_scope=replace(
            profile.cloud_security_scope,
            enabled=True,
            layout="comparison",
        ),
    )


def _dataset(tmp_path: Path, *, populated: bool = True) -> Path:
    cve = {
        "cve": "CVE-2099-1000",
        "severity": "CRITICAL",
        "vpr": 0.0,
        "vpr_display": "0",
        "cvss": 9.8,
        "cvss_display": "9.8",
        "affected_assets": 2,
        "affected_virtual_machines": 1,
        "affected_container_images": 1,
        "components": ["fixture-component"],
        "description": " ".join(
            f"Long technical sentence {index}." for index in range(80)
        ),
        "assets": [
            {
                "kind": "virtual_machine",
                "asset_id": "vm-fixture",
                "name": "vm-fixture.invalid",
                "account_id": "account-fixture",
                "ip_addresses": ["192.0.2.10"],
                "repository_uri": None,
                "digest": None,
                "components": ["fixture-component"],
            },
            {
                "kind": "container_image",
                "asset_id": "image-fixture",
                "name": "image-fixture",
                "account_id": "account-fixture",
                "ip_addresses": [],
                "repository_uri": "registry.invalid/image-fixture",
                "digest": "sha256:fixture",
                "components": ["fixture-component"],
            },
        ],
    }
    hosts = [
        {
            "kind": "virtual_machine",
            "asset_id": "vm-fixture",
            "name": "vm-fixture.invalid",
            "account_id": "account-fixture",
            "ip_addresses": ["192.0.2.10"],
            "repository_uri": None,
            "digest": None,
            "components": ["fixture-component"],
            "vulnerabilities": 4,
            "critical": 1,
            "high": 2,
            "medium": 1,
            "low": 0,
        }
    ]
    images = [
        {
            "kind": "container_image",
            "asset_id": "image-fixture",
            "name": "image-fixture",
            "account_id": "account-fixture",
            "ip_addresses": [],
            "repository_uri": "registry.invalid/image-fixture",
            "digest": "sha256:fixture",
            "components": ["fixture-component"],
            "vulnerabilities": 3,
            "critical": 1,
            "high": 1,
            "medium": 1,
            "low": 0,
        }
    ]
    correctable = {
        **cve,
        "correction_type": "patch_update",
        "correction_type_display": "Patch/Atualização",
        "correction_origin": "explicit",
        "recommended_action": "Apply the vendor security patch.",
        "remediation_steps": ["Apply the vendor security patch."],
        "correlated_findings": 2,
    }
    payload = {
        "schema_version": 1,
        "document_kind": "cloud",
        "metric_definition_version": "cloud-metrics-v1",
        "connector_version": "cloud-graphql-v1",
        "period": {
            "start_at": "2026-07-01T00:00:00+00:00",
            "end_at": "2026-08-01T00:00:00+00:00",
            "reference_at": "2026-08-26T00:00:00+00:00",
            "timezone": "UTC",
            "period_id": "2026-07",
        },
        "collected_at": "2026-08-26T12:00:00+00:00",
        "snapshot_context": {
            "historical_reconstruction": "EXACT_SNAPSHOT",
            "warning": None,
        },
        "overview": {
            "assets": 2 if populated else 0,
            "virtual_machines": 1 if populated else 0,
            "container_images": 1 if populated else 0,
            "vulnerability_occurrences": 7 if populated else 0,
            "unique_cves": 1 if populated else 0,
            "severity_counts": {
                "CRITICAL": 2 if populated else 0,
                "HIGH": 3 if populated else 0,
                "MEDIUM": 2 if populated else 0,
                "LOW": 0,
            },
            "posture_findings": 0,
        },
        "top_critical_cves": [cve] if populated else [],
        "top_vulnerable_hosts": hosts if populated else [],
        "top_vulnerable_images": images if populated else [],
        "workload_status": {
            "total_virtual_machines": 1 if populated else 0,
            "by_max_severity": {
                "CRITICAL": 1 if populated else 0,
                "HIGH": 0,
                "MEDIUM": 0,
                "LOW": 0,
                "NONE": 0,
            },
        },
        "top_components": [],
        "top_posture_findings": [],
        "top_correctable_vulnerabilities": [correctable] if populated else [],
        "aging": {
            "0-30": 1 if populated else 0,
            "31-60": 0,
            "61-90": 0,
            "91-180": 0,
            ">180": 0,
            "data_indisponivel": 0,
        },
        "remediation_performance": {
            "resolved": 0,
            "average_resolution_days": None,
            "period_interval": "[start_at, end_at)",
        },
        "inventory": {
            "total_resources": 0,
            "by_provider": [],
            "by_region": [],
        },
        "source_status": {
            "virtual_machines": "COMPLETE",
            "container_images": "COMPLETE",
            "findings": "UNAVAILABLE",
            "lifecycle": "COMPLETE",
        },
        "quality_issues": [],
        "capabilities": {},
        "history": [],
        "table_provenance": {"schema_version": 1, "tables": {}},
    }
    path = tmp_path / "cloud-dataset.json"
    path.write_text(json.dumps(payload), encoding="utf-8")
    return path


def test_sanitized_cloud_template_keeps_three_page_families() -> None:
    assert CLOUD_TEMPLATE.is_file()
    document = Document(CLOUD_TEMPLATE)
    text = _all_text(CLOUD_TEMPLATE)

    assert len(document.sections) == 3
    assert "{{CLIENT_NAME}}" in text
    assert "{{REPORT_MONTH_YEAR}}" in text
    assert "{{CLOUD_CONTENT_START}}" in text
    assert "01/07/2026" not in text
    assert "TRT8" not in text.upper()
    assert len(document.inline_shapes) == 0

    with zipfile.ZipFile(CLOUD_TEMPLATE) as package:
        all_xml = "\n".join(
            package.read(name).decode("utf-8", errors="ignore")
            for name in package.namelist()
            if name.endswith(".xml")
        )
        header_text = "\n".join(
            "".join(
                node.text or ""
                for node in ET.fromstring(package.read(name)).iter()
                if node.tag.rsplit("}", 1)[-1] == "t"
            )
            for name in package.namelist()
            if name.startswith("word/header") and name.endswith(".xml")
        )
    assert "TRT8" not in all_xml.upper()
    image_nodes = re.findall(r"<wp:docPr\b[^>]*>", all_xml)
    assert image_nodes
    assert all(
        'descr="' in node and 'descr=""' not in node
        for node in image_nodes
    )
    assert not re.search(
        r"n[ºo]\s*\d+\s*/\s*\d{4}",
        header_text,
        re.IGNORECASE,
    )
    assert "dc:creator></dc:creator" in all_xml or "dc:creator/>" in all_xml


def test_base_cloud_report_keeps_approved_sections_and_detailed_top_five(
    tmp_path: Path,
) -> None:
    output = tmp_path / "cloud-base.docx"
    calls: list[str] = []

    def translator(text: str, source: str, target: str) -> str:
        calls.append(text)
        return f"TRADUZIDO: {text}"

    result = generate_cloud_report(
        template_path=CLOUD_TEMPLATE,
        dataset_path=_dataset(tmp_path),
        profile=_profile(),
        output_path=output,
        variant=CloudReportVariant.BASE,
        translator=translator,
    )
    text = _all_text(output)

    assert output.is_file()
    assert result.variant is CloudReportVariant.BASE
    assert tuple(result.rendered_sections) == (
        "cover",
        "table_of_contents",
        "document_control",
        "objective",
        "cloud_overview",
        "introduction",
        "top_hosts",
        "top_images",
        "top_critical",
        "critical_details",
        "top_correctable",
        "dashboard",
        "conclusion",
        "back_cover",
    )
    assert "Principais Vulnerabilidades Críticas" in text
    assert "Principais Vulnerabilidades com Correção Disponível" in text
    assert "CVE-2099-1000" in text
    assert "Tipo de correção" in text
    assert "Ativos afetados" in text
    assert "VPR: 0" in text
    assert "TRADUZIDO:" in text
    assert "TRADUZIDO: Apply the vendor security patch." in text
    assert len(calls) > 1
    assert "{{" not in text
    for paragraph in approved_cloud_editorial_paragraphs():
        assert paragraph in text


def test_empty_cloud_table_has_monthly_message_not_blank_page(
    tmp_path: Path,
) -> None:
    output = tmp_path / "cloud-empty.docx"
    generate_cloud_report(
        template_path=CLOUD_TEMPLATE,
        dataset_path=_dataset(tmp_path, populated=False),
        profile=_profile(),
        output_path=output,
        variant=CloudReportVariant.BASE,
    )

    text = _all_text(output)
    assert "Neste mês não foram identificadas" in text
    assert "{{" not in text


def test_empty_translation_preserves_original_with_explicit_notice(
    tmp_path: Path,
) -> None:
    output = tmp_path / "cloud-translation-fallback.docx"

    generate_cloud_report(
        template_path=CLOUD_TEMPLATE,
        dataset_path=_dataset(tmp_path),
        profile=_profile(),
        output_path=output,
        variant=CloudReportVariant.BASE,
        translator=lambda *_: "",
    )

    text = _all_text(output)
    assert "Long technical sentence 0." in text
    assert "A tradução automática não pôde ser concluída" in text
