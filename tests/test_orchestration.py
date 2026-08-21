from __future__ import annotations

import json
import subprocess
import tempfile
import unittest
from zipfile import ZipFile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from tenable_reports.application.history import (
    finalize_history_publication,
    prepare_dataset_history,
)
from tenable_reports.application.orchestration import (
    OrchestrationRequest,
    build_client_command,
    load_orchestration_config,
    run_orchestration,
)
from tenable_reports.application.publishing import (
    PublicationDocument,
    create_publication_manifest,
    validate_docx_package,
)
from tenable_reports.application.retention import apply_retention, plan_retention
from tenable_reports.application.report_registry import InMemoryReportRegistry
from tenable_reports.cli import main
from tenable_reports.presentation.full_base_report_docx import generate_full_base_report
from tenable_reports.presentation.customizations_report_docx import (
    generate_customizations_report,
)
from tenable_reports.config.profile import load_client_profile


ROOT = Path(__file__).resolve().parents[1]
EXAMPLE_CONFIG = ROOT / "orchestration/clients.example.json"


class _MemorySnapshotRepository:
    location = "memory://orchestration-history"

    def __init__(self) -> None:
        self.published = []

    def publish(self, snapshot) -> None:
        self.published.append(snapshot)

    def compatible_snapshots(self, compatibility, *, before_period_end_at):
        return ()


def _monthly_dataset(
    *,
    run_id: str,
    period_id: str,
    start_at: str,
    end_at: str,
    total: int,
) -> dict:
    payload = json.loads(
        (ROOT / "tests/fixtures/report-dataset-phase5.json").read_text(
            encoding="utf-8"
        )
    )
    payload["metric_definition_version"] = "report-definition-v1.2"
    payload["run_id"] = run_id
    payload["execution_type"] = "AUTOMATIC_MONTHLY"
    payload["period"].update({
        "period_id": period_id,
        "mode": "PREVIOUS_CALENDAR_MONTH",
        "timezone": "America/Fortaleza",
        "start_at": start_at,
        "end_at": end_at,
    })
    payload["metrics"]["non_mitigated"]["total"] = total
    payload["customizations"] = {
        "network_tag_snapshots": [{
            "tag_uuid": "tag-rede-a",
            "category": "Rede",
            "network": "Rede A",
            "period_id": period_id,
            "assets": [{
                "asset_key": "cliente-exemplo:tenable_vm:asset-a",
                "asset_name": "",
                "ip_address": "",
                "critical": 1,
                "high": 2,
                "medium": 3,
                "low": 1,
                "total": total,
                "exploitable": 2,
            }],
        }],
    }
    return payload


def _write_month(directory: Path, payload: dict) -> tuple[Path, Path]:
    dataset = directory / f"{payload['run_id']}.json"
    findings = directory / f"{payload['run_id']}.jsonl"
    dataset.write_text(json.dumps(payload), encoding="utf-8")
    findings.write_text("", encoding="utf-8")
    return dataset, findings


def _docx_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


class OrchestrationTests(unittest.TestCase):
    def test_example_uses_ephemeral_storage_defaults(self) -> None:
        config = load_orchestration_config(EXAMPLE_CONFIG)

        self.assertEqual(config.failed_staging_days, 7)
        self.assertEqual(config.logs_days, 90)
        self.assertTrue(config.cleanup_after_publish)

    def test_tiered_retention_is_applied_automatically_after_orchestration(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            payload = json.loads(EXAMPLE_CONFIG.read_text(encoding="utf-8"))
            output_root = Path(directory) / "data"
            payload["defaults"]["output_root"] = str(output_root)
            payload["clients"] = payload["clients"][:1]
            config_path = ROOT / "orchestration" / "clients.test-tiered-retention.json"
            expired = (
                output_root / "automatic-monthly" / "raw" / "cliente-a-exemplo" / "old-run"
            )
            expired.mkdir(parents=True)
            old_timestamp = datetime(2026, 5, 1, tzinfo=timezone.utc).timestamp()
            import os
            os.utime(expired, (old_timestamp, old_timestamp))

            def runner(command, _):
                return subprocess.CompletedProcess(
                    command,
                    0,
                    stdout=json.dumps({
                        "status": "complete",
                        "client_id": "cliente-a-exemplo",
                        "publication_manifest": "manifest-a.json",
                    }) + "\n",
                    stderr="",
                )

            try:
                config_path.write_text(json.dumps(payload), encoding="utf-8")
                result = run_orchestration(
                    config=load_orchestration_config(config_path),
                    request=OrchestrationRequest(mode="automatic"),
                    runner=runner,
                    sleeper=lambda _: None,
                    now=datetime(2026, 8, 15, 12, tzinfo=timezone.utc),
                    run_status={"old-run": "COMPLETE"},
                    history_confirmed_run_ids={"old-run"},
                )
            finally:
                config_path.unlink(missing_ok=True)

        self.assertFalse(expired.exists())
        self.assertTrue(any("old-run" in path for path in result.retention_removed))
    def test_scheduled_transient_failure_succeeds_on_second_attempt(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            payload = json.loads(EXAMPLE_CONFIG.read_text(encoding="utf-8"))
            payload["defaults"]["output_root"] = str(Path(directory) / "data")
            payload["clients"] = payload["clients"][:1]
            config_path = ROOT / "orchestration" / "clients.test-retry.json"
            calls = []

            def runner(command, _):
                calls.append(tuple(command))
                if len(calls) == 1:
                    return subprocess.CompletedProcess(
                        command,
                        2,
                        stdout=json.dumps({
                            "status": "failed",
                            "error_code": "TENABLE_RATE_LIMIT",
                            "retryable": True,
                            "message": "Limite temporário da API.",
                        }) + "\n",
                        stderr="",
                    )
                return subprocess.CompletedProcess(
                    command,
                    0,
                    stdout=json.dumps({
                        "status": "complete",
                        "client_id": "cliente-a-exemplo",
                        "publication_manifest": "manifest-a.json",
                    }) + "\n",
                    stderr="",
                )

            try:
                config_path.write_text(json.dumps(payload), encoding="utf-8")
                result = run_orchestration(
                    config=load_orchestration_config(config_path),
                    request=OrchestrationRequest(mode="automatic"),
                    runner=runner,
                    sleeper=lambda _: None,
                    now=datetime(2026, 8, 1, 12, tzinfo=timezone.utc),
                )
            finally:
                config_path.unlink(missing_ok=True)

        self.assertEqual(result.clients[0].status, "COMPLETE")
        self.assertEqual(
            [attempt.origin for attempt in result.clients[0].attempts],
            ["SCHEDULED", "AUTOMATIC_RETRY"],
        )
        self.assertEqual([attempt.attempt_number for attempt in result.clients[0].attempts], [1, 2])
        self.assertEqual(len(calls), 2)

    def test_invalid_credentials_do_not_retry(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            payload = json.loads(EXAMPLE_CONFIG.read_text(encoding="utf-8"))
            payload["defaults"]["output_root"] = str(Path(directory) / "data")
            payload["clients"] = payload["clients"][:1]
            config_path = ROOT / "orchestration" / "clients.test-no-retry.json"
            calls = []

            def runner(command, _):
                calls.append(tuple(command))
                return subprocess.CompletedProcess(
                    command,
                    2,
                    stdout=json.dumps({
                        "status": "failed",
                        "error_code": "TENABLE_AUTH_INVALID",
                        "retryable": False,
                        "message": "Credenciais inválidas.",
                    }) + "\n",
                    stderr="",
                )

            try:
                config_path.write_text(json.dumps(payload), encoding="utf-8")
                result = run_orchestration(
                    config=load_orchestration_config(config_path),
                    request=OrchestrationRequest(mode="automatic"),
                    runner=runner,
                    sleeper=lambda _: None,
                    now=datetime(2026, 8, 1, 12, tzinfo=timezone.utc),
                )
            finally:
                config_path.unlink(missing_ok=True)

        self.assertEqual(result.clients[0].status, "FAILED")
        self.assertEqual(len(result.clients[0].attempts), 1)
        self.assertEqual(len(calls), 1)
    def test_example_configuration_resolves_two_isolated_clients(self) -> None:
        config = load_orchestration_config(EXAMPLE_CONFIG)
        self.assertEqual(config.orchestration_id, "carteira-tenable")
        self.assertEqual(config.max_parallel, 2)
        self.assertEqual(config.database_env_file.name, "database.env")
        self.assertEqual(len(config.clients), 2)
        self.assertNotEqual(config.clients[0].env_file, config.clients[1].env_file)
        self.assertEqual(config.clients[0].tags, ("Rede: Matriz",))

    def test_configuration_rejects_embedded_credentials(self) -> None:
        profile = ROOT / "clients/examples/orchestration/client-a.json"
        with tempfile.TemporaryDirectory() as directory:
            config_path = Path(directory) / "clients.json"
            config_path.write_text(json.dumps({
                "schema_version": 1,
                "orchestration_id": "unsafe",
                "defaults": {"tenable_secret": "do-not-store-here"},
                "clients": [{
                    "client_id": "cliente-a-exemplo",
                    "profile": str(profile),
                    "env_file": str(Path(directory) / "client-a.env"),
                }],
            }), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "Credencial embutida"):
                load_orchestration_config(config_path)

    def test_dry_run_builds_manual_commands_without_credentials(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            payload = json.loads(EXAMPLE_CONFIG.read_text(encoding="utf-8"))
            payload["defaults"]["output_root"] = str(Path(directory) / "data")
            config_path = ROOT / "orchestration" / "clients.test-dry-run.json"
            try:
                config_path.write_text(json.dumps(payload), encoding="utf-8")
                config = load_orchestration_config(config_path)
                result = run_orchestration(
                    config=config,
                    request=OrchestrationRequest(
                        mode="manual",
                        days=10,
                        dry_run=True,
                    ),
                    now=datetime(2026, 8, 13, 12, tzinfo=timezone.utc),
                )
            finally:
                config_path.unlink(missing_ok=True)
            self.assertEqual(result.status, "DRY_RUN")
            self.assertEqual(len(result.clients), 2)
            for client in result.clients:
                command_text = " ".join(client.command)
                self.assertIn("--mode manual", command_text)
                self.assertIn("--days 10", command_text)
                self.assertNotIn("do-not-store-here", command_text)
                self.assertTrue(client.log_path.is_file())
            self.assertTrue(result.manifest_path.is_file())
            self.assertTrue(result.notification_path.is_file())

    def test_real_orchestration_isolates_failure_and_continues_other_client(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            payload = json.loads(EXAMPLE_CONFIG.read_text(encoding="utf-8"))
            payload["defaults"]["output_root"] = str(Path(directory) / "data")
            config_path = ROOT / "orchestration" / "clients.test-runner.json"
            try:
                config_path.write_text(json.dumps(payload), encoding="utf-8")
                config = load_orchestration_config(config_path)

                def runner(command: list[str], _: Path) -> subprocess.CompletedProcess[str]:
                    client_id = "cliente-a-exemplo" if "client-a.json" in " ".join(command) else "cliente-b-exemplo"
                    if client_id == "cliente-a-exemplo":
                        return subprocess.CompletedProcess(
                            command,
                            0,
                            stdout=json.dumps({
                                "status": "complete",
                                "client_id": client_id,
                                "publication_manifest": "manifest-a.json",
                            }) + "\n",
                            stderr="",
                        )
                    return subprocess.CompletedProcess(
                        command,
                        2,
                        stdout="",
                        stderr="erro: export indisponivel",
                    )

                result = run_orchestration(
                    config=config,
                    request=OrchestrationRequest(mode="automatic"),
                    runner=runner,
                    now=datetime(2026, 8, 13, 12, tzinfo=timezone.utc),
                )
            finally:
                config_path.unlink(missing_ok=True)
            self.assertEqual(result.status, "PARTIAL_FAILURE")
            self.assertEqual(result.failed_count, 1)
            statuses = {item.client_id: item.status for item in result.clients}
            self.assertEqual(statuses["cliente-a-exemplo"], "COMPLETE")
            self.assertEqual(statuses["cliente-b-exemplo"], "FAILED")

    def test_client_command_uses_process_safe_env_path_and_non_interactive_tags(self) -> None:
        config = load_orchestration_config(EXAMPLE_CONFIG)
        command = build_client_command(
            config=config,
            client=config.clients[0],
            request=OrchestrationRequest(mode="automatic"),
            client_run_id="run-client-a",
        )
        command_text = " ".join(command)
        self.assertIn("run-client", command)
        self.assertIn(str(config.clients[0].env_file), command)
        self.assertIn(str(config.database_env_file), command)
        self.assertIn("Rede: Matriz", command)
        self.assertNotIn("--select-tags", command)
        self.assertNotIn("TENABLE_ACCESS", command_text)

    def test_cli_orchestration_requires_confirmation_except_dry_run(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            payload = json.loads(EXAMPLE_CONFIG.read_text(encoding="utf-8"))
            payload["defaults"]["output_root"] = str(Path(directory) / "data")
            config_path = ROOT / "orchestration" / "clients.test-cli.json"
            try:
                config_path.write_text(json.dumps(payload), encoding="utf-8")
                self.assertEqual(main(["orchestrate", "--config", str(config_path)]), 2)
                self.assertEqual(main([
                    "orchestrate",
                    "--config",
                    str(config_path),
                    "--mode",
                    "manual",
                    "--days",
                    "5",
                    "--dry-run",
                ]), 0)
            finally:
                config_path.unlink(missing_ok=True)

    def test_publication_manifest_validates_docx_and_hashes(self) -> None:
        profile = load_client_profile(ROOT / "clients/examples/client-profile.json")
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory)
            base = generate_full_base_report(
                template_path=ROOT / "templates/corporate/base-v1.docx",
                dataset_path=ROOT / "tests/fixtures/report-dataset-phase5.json",
                profile=profile,
                output_path=output / "base.docx",
                assets_dir=ROOT / "templates/corporate/assets",
                mask_sensitive=True,
            )
            custom = generate_customizations_report(
                template_path=ROOT / "templates/corporate/base-v1.docx",
                dataset_path=ROOT / "tests/fixtures/report-dataset-phase5.json",
                profile=profile,
                output_path=output / "custom.docx",
                mask_sensitive=True,
            )
            manifest = create_publication_manifest(
                output_path=output / "publication.json",
                client_id=profile.client_id,
                tenant_id=profile.tenant_id,
                run_id="offline-proof",
                execution_type="MANUAL",
                period={"period_id": "2026-07"},
                dataset_path=ROOT / "tests/fixtures/report-dataset-phase5.json",
                documents=(base.output_path, custom.output_path),
                history_database=None,
            )
            payload = json.loads(manifest.read_text(encoding="utf-8"))
            self.assertEqual(payload["status"], "READY_FOR_CONTROLLED_DISTRIBUTION")
            self.assertEqual(len(payload["documents"]), 2)
            self.assertTrue(all(len(item["sha256"]) == 64 for item in payload["documents"]))
            self.assertEqual(validate_docx_package(base.output_path)["package_status"], "VALID")

    def test_publication_manifest_records_tag_document_metadata(self) -> None:
        profile = load_client_profile(ROOT / "clients/examples/client-profile.json")
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory)
            tag_document = generate_full_base_report(
                template_path=ROOT / "templates/corporate/base-v1.docx",
                dataset_path=ROOT / "tests/fixtures/report-dataset-phase5.json",
                profile=profile,
                output_path=output / "tag.docx",
                assets_dir=ROOT / "templates/corporate/assets",
                mask_sensitive=True,
            )
            manifest = create_publication_manifest(
                output_path=output / "publication.json",
                client_id=profile.client_id,
                tenant_id=profile.tenant_id,
                run_id="tag-proof",
                execution_type="MANUAL",
                period={"period_id": "2026-07"},
                dataset_path=ROOT / "tests/fixtures/report-dataset-phase5.json",
                documents=(PublicationDocument(
                    path=tag_document.output_path,
                    document_kind="tag",
                    tag_uuid="tag-a",
                    tag_category="Equipe",
                    tag_value="Infra",
                ),),
                history_database=None,
            )
            document = json.loads(manifest.read_text(encoding="utf-8"))["documents"][0]
            self.assertEqual(document["document_kind"], "tag")
            self.assertEqual(document["tag_uuid"], "tag-a")
            self.assertEqual(document["tag_category"], "Equipe")
            self.assertEqual(document["tag_value"], "Infra")

    def test_publication_rejects_structurally_empty_docx_before_cleanup(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            document = Path(directory) / "empty.docx"
            with ZipFile(document, "w") as package:
                package.writestr("[Content_Types].xml", "")
                package.writestr("word/document.xml", "")
                package.writestr("word/styles.xml", "")

            with self.assertRaisesRegex(ValueError, "DOCX"):
                validate_docx_package(document)

    def test_retention_is_plan_only_until_explicit_application(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory) / "manual"
            expired = root / "reports" / "cliente" / "run-expired"
            expired.mkdir(parents=True)
            (expired / "proof.txt").write_text("proof", encoding="utf-8")
            old_timestamp = datetime(2025, 1, 1, tzinfo=timezone.utc).timestamp()
            import os
            os.utime(expired, (old_timestamp, old_timestamp))
            plan = plan_retention(
                scoped_output_root=root,
                retention_days=30,
                now=datetime(2026, 8, 13, tzinfo=timezone.utc),
            )
            self.assertEqual(len(plan), 1)
            self.assertTrue(expired.is_dir())
            removed = apply_retention(scoped_output_root=root, candidates=plan)
            self.assertEqual(removed, (expired.resolve(),))
            self.assertFalse(expired.exists())

    def test_two_month_main_reference_end_to_end(self) -> None:
        profile = load_client_profile(
            ROOT / "clients/examples/client-profile-all-customizations.json"
        )
        registry = InMemoryReportRegistry()
        snapshots = _MemorySnapshotRepository()

        with tempfile.TemporaryDirectory() as directory_name:
            directory = Path(directory_name)
            july_v1, july_v1_findings = _write_month(
                directory,
                _monthly_dataset(
                    run_id="run-july-v1",
                    period_id="2026-07",
                    start_at="2026-07-01T03:00:00Z",
                    end_at="2026-08-01T03:00:00Z",
                    total=100,
                ),
            )
            first = prepare_dataset_history(
                profile=profile,
                dataset_path=july_v1,
                normalized_findings_path=july_v1_findings,
                output_path=directory / "july-v1-enriched.json",
                registry=registry,
                repository=snapshots,
            )
            first_payload = json.loads(
                first.enriched_dataset_path.read_text(encoding="utf-8")
            )
            first_customizations = first_payload["customizations"]
            self.assertEqual(first.history_status, "NO_IMMEDIATE_MAIN")
            self.assertEqual(
                first_customizations["monthly_history"][0]["non_mitigated"],
                100,
            )
            self.assertEqual(
                first_customizations["network_tag_snapshots"][0]["tag_uuid"],
                "tag-rede-a",
            )
            first_docx = generate_customizations_report(
                template_path=ROOT / "templates/corporate/base-v1.docx",
                dataset_path=first.enriched_dataset_path,
                profile=profile,
                output_path=directory / "july-v1-custom.docx",
                mask_sensitive=True,
            )
            first_text = _docx_text(first_docx.output_path)
            self.assertNotIn("Baseline do período atual", first_text)
            self.assertIn(
                "Não há histórico do período imediatamente anterior para comparação.",
                first_text,
            )

            finalize_history_publication(
                first,
                snapshot_repository=snapshots,
                registry=registry,
                publication_validated=True,
                auto_promote=True,
            )
            self.assertEqual(
                registry.get_main(first.reference_key).run_id,
                "run-july-v1",
            )

            july_v2, july_v2_findings = _write_month(
                directory,
                _monthly_dataset(
                    run_id="run-july-v2",
                    period_id="2026-07",
                    start_at="2026-07-01T03:00:00Z",
                    end_at="2026-08-01T03:00:00Z",
                    total=125,
                ),
            )
            second = prepare_dataset_history(
                profile=profile,
                dataset_path=july_v2,
                normalized_findings_path=july_v2_findings,
                output_path=directory / "july-v2-enriched.json",
                registry=registry,
                repository=snapshots,
            )
            finalize_history_publication(
                second,
                snapshot_repository=snapshots,
                registry=registry,
                publication_validated=True,
                auto_promote=True,
            )
            self.assertEqual(
                registry.get_main(second.reference_key).run_id,
                "run-july-v1",
            )

            registry.promote_main(
                second.reference_key,
                "run-july-v2",
                actor="analyst",
                reason="versão revisada",
            )
            self.assertEqual(
                registry.get_main(second.reference_key).run_id,
                "run-july-v2",
            )

            august_v1, august_v1_findings = _write_month(
                directory,
                _monthly_dataset(
                    run_id="run-august-v1",
                    period_id="2026-08",
                    start_at="2026-08-01T03:00:00Z",
                    end_at="2026-09-01T03:00:00Z",
                    total=140,
                ),
            )
            august = prepare_dataset_history(
                profile=profile,
                dataset_path=august_v1,
                normalized_findings_path=august_v1_findings,
                output_path=directory / "august-v1-enriched.json",
                registry=registry,
                repository=snapshots,
            )
            self.assertEqual(august.predecessor.run_id, "run-july-v2")
            august_payload = json.loads(
                august.enriched_dataset_path.read_text(encoding="utf-8")
            )
            self.assertEqual(
                august_payload["customizations"]["previous_period_overview"]
                ["total"]["non_mitigated"],
                125,
            )

            registry.soft_delete(
                "run-july-v2",
                actor="analyst",
                reason="versão retirada",
                replacement_run_id="run-july-v1",
            )
            self.assertEqual(
                registry.get_main(first.reference_key).run_id,
                "run-july-v1",
            )

            august_v2_payload = _monthly_dataset(
                run_id="run-august-v2",
                period_id="2026-08",
                start_at="2026-08-01T03:00:00Z",
                end_at="2026-09-01T03:00:00Z",
                total=150,
            )
            august_v2, august_v2_findings = _write_month(
                directory,
                august_v2_payload,
            )
            after_delete = prepare_dataset_history(
                profile=profile,
                dataset_path=august_v2,
                normalized_findings_path=august_v2_findings,
                output_path=directory / "august-v2-enriched.json",
                registry=registry,
                repository=snapshots,
            )
            self.assertEqual(after_delete.predecessor.run_id, "run-july-v1")
            after_delete_payload = json.loads(
                after_delete.enriched_dataset_path.read_text(encoding="utf-8")
            )
            self.assertEqual(
                after_delete_payload["customizations"]["previous_period_overview"]
                ["total"]["non_mitigated"],
                100,
            )


if __name__ == "__main__":
    unittest.main()
