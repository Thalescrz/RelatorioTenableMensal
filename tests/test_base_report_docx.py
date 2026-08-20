from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from tenable_reports.config.profile import load_client_profile
from tenable_reports.presentation.base_report_docx import (
    ASSET_HEADERS,
    TEMPLATE_VERSION,
    create_base_template,
    generate_base_report,
)


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "templates/corporate/assets"
FIXTURE = ROOT / "tests/fixtures/report-dataset-phase5.json"
PROFILE = ROOT / "clients/examples/client-profile.json"


def all_text(document):
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for container in (section.header, section.footer):
            chunks.extend(paragraph.text for paragraph in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


class BaseReportDocxTests(unittest.TestCase):
    def test_template_and_sanitised_proof_follow_contract(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            template = Path(directory) / "base-v1.docx"
            output = Path(directory) / "proof.docx"
            create_base_template(template, assets_dir=ASSETS)
            result = generate_base_report(
                template_path=template,
                dataset_path=FIXTURE,
                profile=load_client_profile(PROFILE),
                output_path=output,
                mask_sensitive=True,
            )

            self.assertTrue(output.is_file())
            self.assertEqual(result.template_version, TEMPLATE_VERSION)
            self.assertEqual(result.period_id, "2026-07")
            self.assertEqual(result.top_asset_rows, 10)
            self.assertTrue(result.masked_sensitive_fields)

            document = Document(output)
            text = all_text(document)
            self.assertNotIn("{{", text)
            self.assertIn("JULHO/2026", text)
            self.assertIn("01/07/2026 a 31/07/2026", text)
            self.assertIn("Exploitable", text)
            self.assertNotIn("Output", text)
            self.assertGreaterEqual(len(document.inline_shapes), 3)

            asset_table = next(
                table
                for table in document.tables
                if tuple(cell.text for cell in table.rows[0].cells) == ASSET_HEADERS
            )
            self.assertEqual(len(asset_table.rows), 11)
            for row in asset_table.rows[1:]:
                self.assertEqual(row.cells[0].text, "")
                self.assertEqual(row.cells[1].text, "")
                self.assertLessEqual(int(row.cells[7].text), int(row.cells[6].text))

            section = document.sections[0]
            self.assertAlmostEqual(section.page_width.cm, 21.0, places=1)
            self.assertAlmostEqual(section.page_height.cm, 29.7, places=1)

            with zipfile.ZipFile(output) as package:
                document_xml = package.read("word/document.xml").decode("utf-8")
                core_xml = package.read("docProps/core.xml").decode("utf-8")
                self.assertGreaterEqual(document_xml.count('descr="'), 3)
                self.assertIn("w:tblHeader", document_xml)
                self.assertNotIn("{{", document_xml)
                self.assertNotIn("@", document_xml)
                self.assertIn("ITProtect", core_xml)

    def test_inconsistent_exploitable_count_is_rejected(self) -> None:
        import json

        with tempfile.TemporaryDirectory() as directory:
            directory_path = Path(directory)
            template = directory_path / "base-v1.docx"
            dataset = directory_path / "invalid.json"
            output = directory_path / "proof.docx"
            create_base_template(template, assets_dir=ASSETS)
            payload = json.loads(FIXTURE.read_text(encoding="utf-8"))
            payload["top_assets"][0]["exploitable"] = payload["top_assets"][0]["total"] + 1
            dataset.write_text(json.dumps(payload), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "não pode exceder total"):
                generate_base_report(
                    template_path=template,
                    dataset_path=dataset,
                    profile=load_client_profile(PROFILE),
                    output_path=output,
                )


if __name__ == "__main__":
    unittest.main()
