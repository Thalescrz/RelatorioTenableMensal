from __future__ import annotations

import json
import tempfile
from dataclasses import replace
from pathlib import Path

from docx import Document

from tenable_reports.config.profile import load_client_profile
from tenable_reports.presentation.customizations_report_docx import (
    generate_customizations_report,
)


ROOT = Path(__file__).resolve().parents[1]


def _text(document):
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_customizations_are_kept_outside_the_base_document() -> None:
    with tempfile.TemporaryDirectory() as directory:
        output = Path(directory) / "custom.docx"
        result = generate_customizations_report(
            template_path=ROOT / "templates/corporate/base-v1.docx",
            dataset_path=ROOT / "tests/fixtures/report-dataset-phase5.json",
            profile=load_client_profile(
                ROOT / "clients/examples/client-profile-all-customizations.json"
            ),
            output_path=output,
            mask_sensitive=True,
        )
        assert set(result.rendered_modules) == {
            "vm_monthly_volume",
            "vm_previous_period_delta",
            "vm_network_comparison",
            "scan_auth_health",
            "vm_plugin_family",
            "vm_eol_software",
            "vm_executive_evolution",
            "vm_monthly_evolution",
            "cloud_container_images",
            "vm_exploit_vector",
            "was_unsupported_tech",
        }
        assert result.requested_modules == tuple(
            load_client_profile(
                ROOT / "clients/examples/client-profile-all-customizations.json"
            ).report.intelligence_modules
        )
        assert result.omitted_modules == ()
        document = Document(output)
        text = _text(document)
        assert "JULHO/2026" in text
        assert "01/07/2026 a 31/07/2026" in text
        assert "Comparativo Mensal de Vulnerabilidades" in text
        assert "Vulnerabilidades “Não Mitigadas”" in text
        assert "Vulnerabilidades “Mitigadas”" in text
        assert "Geral" in text
        assert "Servidores" in text
        assert "Sistemas operacionais e software sem suportes" in text
        assert "TENABLE CLOUD SECURITY (CONTAINER IMAGES)" in text
        assert "Vulnerabilidades Exploráveis por Vetor de Ataque" in text
        assert "Dados indisponíveis para este indicador." in text
        assert "Principais ativos Vulneráveis por Rede" in text
        assert "Rede de exemplo A" in text
        assert "Rede de exemplo B" not in text
        assert "Junho/26" in text
        assert "Julho/26" in text
        assert "Exploitable" in text
        assert "Posição anterior" in text
        assert "Movimentação" in text
        assert "Principais Aplicações “Unsupported”" in text
        assert "SUA MELHOR ALIADA NA JORNADA DA PROTEÇÃO DIGITAL." in text
        assert "METODOLOGIA, QUALIDADE E LIMITAÇÕES" not in text
        # Capa/logos mais os gráficos mensais observados nos documentos de referência:
        # não mitigadas, mitigadas, evolução e novas por severidade.
        assert len(document.inline_shapes) >= 12


def test_customization_modules_without_data_are_omitted_with_reason() -> None:
    with tempfile.TemporaryDirectory() as directory:
        source = json.loads(
            (ROOT / "tests/fixtures/report-dataset-phase5.json").read_text(
                encoding="utf-8"
            )
        )
        source["customizations"] = {}
        dataset = Path(directory) / "without-history.json"
        dataset.write_text(json.dumps(source), encoding="utf-8")
        result = generate_customizations_report(
            template_path=ROOT / "templates/corporate/base-v1.docx",
            dataset_path=dataset,
            profile=load_client_profile(
                ROOT / "clients/examples/client-profile-intelligence-expanded.json"
            ),
            output_path=Path(directory) / "custom.docx",
            mask_sensitive=True,
        )
        assert result.rendered_modules == ()
        reasons = {item["module_id"]: item["reason"] for item in result.omitted_modules}
        assert reasons["vm_monthly_volume"] == "NO_COMPATIBLE_HISTORY"
        assert reasons["vm_network_comparison"] == "NO_COMPATIBLE_HISTORY"
        assert reasons["cloud_container_images"] == "NO_COMPATIBLE_DATA"


def test_profile_without_customizations_produces_only_cover_and_summary() -> None:
    with tempfile.TemporaryDirectory() as directory:
        output = Path(directory) / "custom.docx"
        result = generate_customizations_report(
            template_path=ROOT / "templates/corporate/base-v1.docx",
            dataset_path=ROOT / "tests/fixtures/report-dataset-phase5.json",
            profile=load_client_profile(
                ROOT / "clients/examples/client-profile-vm-standard.json"
            ),
            output_path=output,
            mask_sensitive=True,
        )
        assert result.rendered_modules == ()
        document = Document(output)
        assert len(document.sections) == 1
        assert "SUA MELHOR ALIADA" not in _text(document)


def test_first_month_renders_current_baseline_and_explicit_no_data_messages() -> None:
    with tempfile.TemporaryDirectory() as directory:
        source = json.loads(
            (ROOT / "tests/fixtures/report-dataset-phase5.json").read_text(encoding="utf-8")
        )
        source["customizations"] = {
            "monthly_history": [source["customizations"]["monthly_history"][-1]],
            "monthly_views": [{
                "id": "general", "label": "Geral",
                "history": [source["customizations"]["monthly_history"][-1]],
            }],
            "network_tag_snapshots": [{
                "tag_uuid": "tag-a", "network": "Rede A", "period_id": "2026-07",
                "assets": [{"asset_key": "a", "total": 3, "exploitable": 1}],
            }],
            "plugin_family": [], "eol_assets": [], "eol_software": [],
            "attack_vectors": [], "was_unsupported_tech": [],
            "scan_auth_health": {"success": 0, "failure": 0, "total": 0},
            "customization_statuses": {
                "scan_auth_health": "NO_OCCURRENCES",
                "vm_plugin_family": "NO_OCCURRENCES",
                "vm_eol_software": "NO_OCCURRENCES",
                "vm_exploit_vector": "NO_OCCURRENCES",
                "was_unsupported_tech": "NO_OCCURRENCES",
            },
            "history_status": {"status": "NO_IMMEDIATE_MAIN"},
        }
        dataset = Path(directory) / "first-month.json"
        dataset.write_text(json.dumps(source), encoding="utf-8")
        result = generate_customizations_report(
            template_path=ROOT / "templates/corporate/base-v1.docx",
            dataset_path=dataset,
            profile=load_client_profile(ROOT / "clients/examples/client-profile-all-customizations.json"),
            output_path=Path(directory) / "custom.docx",
            mask_sensitive=True,
        )
        text = _text(Document(result.output_path))
        assert "Não há histórico do período imediatamente anterior para comparação." in text
        assert "Baseline do período atual" in text
        assert "Movimentação" not in text
        assert "Neste mês não foram identificadas vulnerabilidades mitigadas" in text
        assert "Neste mês não foram identificados sistemas ou softwares sem suporte." in text
        assert "Neste mês não foram identificadas vulnerabilidades exploráveis" in text
        assert "Neste mês não foram identificadas tecnologias WEB sem suporte." in text


def test_source_filter_notes_cover_custom_data_tables() -> None:
    with tempfile.TemporaryDirectory() as directory:
        source = json.loads(
            (ROOT / "tests/fixtures/report-dataset-phase5.json").read_text(
                encoding="utf-8"
            )
        )
        common = {
            "view": "Explore > Findings > Vulnerabilities",
            "period_start_at": "2026-07-01T00:00:00Z",
            "period_end_at": "2026-08-01T00:00:00Z",
            "severities": ["CRITICAL", "HIGH", "MEDIUM", "LOW"],
        }
        source["table_provenance"] = {
            "version": "table-provenance-v1",
            "tables": {
                "previous_period_overview": {
                    **common,
                    "validation_queries": [
                        {"label": "Não mitigadas", "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"]},
                        {"label": "Mitigadas", "states": ["FIXED"], "date_fields": ["Last Fixed"]},
                    ],
                    "rule": "Período main anterior",
                },
                "network_tag_snapshots": [{
                    **common,
                    "states": ["OPEN", "REOPENED"],
                    "date_fields": ["Last Seen"],
                    "tag_uuid": "tag-rede-exemplo-a",
                    "tag_category": "Rede",
                    "tag_value": "Rede de exemplo A",
                    "rule": "Mesma rede em dois períodos",
                }],
                "network_asset_movement": [{
                    **common,
                    "tag_uuid": "tag-rede-exemplo-a",
                    "tag_category": "Rede",
                    "tag_value": "Rede de exemplo A",
                    "validation_queries": [
                        {"label": "Consulta 1", "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"]},
                        {"label": "Consulta 2", "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"]},
                    ],
                    "rule": "Variação de posição do ativo",
                }],
                "plugin_family": {**common, "states": ["FIXED"], "date_fields": ["Last Fixed"], "rule": "Agrupar por família de plugin"},
                "eol_assets": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Catálogo textual de fim de suporte"},
                "eol_software": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Plugins de fim de suporte"},
                "container_images": {**common, "platform_validation_available": False, "rule": "Agrupar por imagem de container"},
                "container_findings": {**common, "platform_validation_available": False, "rule": "Findings da imagem selecionada"},
                "attack_vectors": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "platform_filters": {"Plugin > Exploit Available": "Yes"}, "rule": "Exploit Available e vetor CVSS v3"},
                "was_unsupported_tech": {**common, "states": ["OPEN", "REOPENED"], "date_fields": ["Last Seen"], "rule": "Tecnologias WEB sem suporte"},
            },
        }
        dataset = Path(directory) / "custom-filters.json"
        dataset.write_text(json.dumps(source), encoding="utf-8")
        profile = load_client_profile(
            ROOT / "clients/examples/client-profile-all-customizations.json"
        )
        profile = replace(
            profile,
            presentation=replace(profile.presentation, show_source_filters=True),
        )
        output = Path(directory) / "custom-filters.docx"
        generate_customizations_report(
            template_path=ROOT / "templates/corporate/base-v1.docx",
            dataset_path=dataset,
            profile=profile,
            output_path=output,
            mask_sensitive=True,
        )
        text = _text(Document(output))
        for marker in (
            "Período main anterior",
            "Tag = Rede:Rede de exemplo A",
            "Variação de posição do ativo",
            "Agrupar por família de plugin",
            "Catálogo textual de fim de suporte",
            "Exploit Available e vetor CVSS v3",
            "Tecnologias WEB sem suporte",
            "Mitigadas: State = Fixed; Last Fixed = Junho 2026",
            "Consulta 1: State = Active, New, Resurfaced; Last Seen = Junho/26",
            "Consulta 2: State = Active, New, Resurfaced; Last Seen = Julho/26",
            "State = Fixed; Severity = Critical, High, Medium, Low; Last Fixed = 01/07/2026 a 31/07/2026",
            "Plugin > Exploit Available = Yes",
            "Plugin ID = 990001, 990002",
            "Plugin ID = 981001, 981002",
        ):
            assert marker in text
        assert "Validação rápida na Tenable: Cloud Security" not in text
