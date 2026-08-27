from __future__ import annotations

import argparse
import hashlib
import posixpath
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


CLIENT_PLACEHOLDER = "{{CLIENT_NAME}}"
MONTH_PLACEHOLDER = "{{REPORT_MONTH_YEAR}}"
CONTENT_PLACEHOLDER = "{{CLOUD_CONTENT_START}}"
TOC_PLACEHOLDER = "{{TABLE_OF_CONTENTS}}"

_MONTH_YEAR = re.compile(
    r"^(?:JANEIRO|FEVEREIRO|MAR[ÇC]O|ABRIL|MAIO|JUNHO|JULHO|AGOSTO|"
    r"SETEMBRO|OUTUBRO|NOVEMBRO|DEZEMBRO)\s*/\s*\d{4}$",
    re.IGNORECASE,
)
_CLIENT_CODE = re.compile(r"^TRT\s*\d+$", re.IGNORECASE)
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_OFFICE_REL_NS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
)
_CONTENT_TYPES_NS = (
    "http://schemas.openxmlformats.org/package/2006/content-types"
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _read_forbidden_terms(path: str | Path | None) -> tuple[str, ...]:
    if path is None:
        return ()
    source = Path(path)
    if not source.is_file():
        raise FileNotFoundError(f"Arquivo de termos proibidos não encontrado: {source}")
    return tuple(
        dict.fromkeys(
            line.strip()
            for line in source.read_text(encoding="utf-8-sig").splitlines()
            if line.strip()
        )
    )


def _set_all_text_nodes(element, value: str) -> None:
    nodes = [node for node in element.iter() if node.tag.rsplit("}", 1)[-1] == "t"]
    if not nodes:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = value
        run.append(text)
        element.append(run)
        return
    nodes[0].text = value
    for node in nodes[1:]:
        node.text = ""


def _replace_forbidden_terms(element, forbidden_terms: Iterable[str]) -> None:
    for node in element.iter():
        if node.tag.rsplit("}", 1)[-1] != "t" or not node.text:
            continue
        value = node.text
        for term in forbidden_terms:
            value = re.sub(re.escape(term), CLIENT_PLACEHOLDER, value, flags=re.IGNORECASE)
        value = re.sub(
            r"\s*[–—-]\s*CONTRATO\b.*$",
            "",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(
            r"\s*n[ºo]\s*\d+\s*/\s*\d{4}",
            "",
            value,
            flags=re.IGNORECASE,
        )
        node.text = value


def _paragraph_plain_text(element) -> str:
    return " ".join(
        "".join(node.text or "" for node in element.iter() if node.tag == qn("w:t")).split()
    )


def _find_section_boundaries(document: Document) -> list[int]:
    body = document._element.body
    boundaries: list[int] = []
    for index, element in enumerate(body):
        if element.tag == qn("w:sectPr"):
            boundaries.append(index)
            continue
        if element.find(".//" + qn("w:sectPr")) is not None:
            boundaries.append(index)
    return boundaries


def _new_placeholder_paragraph(text: str):
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    paragraph.append(run)
    return paragraph


def _distill_body(document: Document, forbidden_terms: tuple[str, ...]) -> None:
    body = document._element.body
    boundaries = _find_section_boundaries(document)
    if len(boundaries) < 3:
        raise ValueError(
            "O modelo Cloud precisa conter as três seções aprovadas: capa/sumário, miolo e contracapa."
        )
    first_boundary, content_boundary = boundaries[0], boundaries[1]

    for element in list(body):
        _replace_forbidden_terms(element, forbidden_terms)

    cover_elements = list(body)[:first_boundary]
    non_empty_cover = [
        element for element in cover_elements if _paragraph_plain_text(element)
    ]
    if len(non_empty_cover) < 3:
        raise ValueError("Não foi possível identificar os campos de capa do modelo Cloud.")

    month_element = next(
        (
            element
            for element in cover_elements
            if _MONTH_YEAR.fullmatch(_paragraph_plain_text(element))
        ),
        None,
    )
    if month_element is None:
        raise ValueError("Mês/ano da capa não foi identificado.")
    _set_all_text_nodes(month_element, MONTH_PLACEHOLDER)

    title_position = next(
        index
        for index, element in enumerate(non_empty_cover)
        if "CLOUD SECURITY" in _paragraph_plain_text(element).upper()
    )
    if title_position == 0:
        raise ValueError("Campo de cliente da capa não foi identificado.")
    client_element = non_empty_cover[title_position - 1]
    _set_all_text_nodes(client_element, CLIENT_PLACEHOLDER)

    for element in list(body)[first_boundary + 1 : content_boundary]:
        body.remove(element)

    boundary_element = list(body)[first_boundary + 1]
    Paragraph(boundary_element, document._body).add_run(CONTENT_PLACEHOLDER)

    for element in list(body)[: first_boundary + 1]:
        if element.tag.rsplit("}", 1)[-1] == "sdt":
            position = list(body).index(element)
            body.remove(element)
            body.insert(position, _new_placeholder_paragraph(TOC_PLACEHOLDER))


def _sanitize_headers_and_footers(
    document: Document,
    forbidden_terms: tuple[str, ...],
) -> None:
    seen: set[int] = set()
    for section in document.sections:
        for container in (section.header, section.footer):
            identity = id(container._element)
            if identity in seen:
                continue
            seen.add(identity)
            _replace_forbidden_terms(container._element, forbidden_terms)
            text_nodes = [
                node
                for node in container._element.iter()
                if node.tag.rsplit("}", 1)[-1] == "t"
            ]
            after_client = False
            for node in text_nodes:
                value = node.text or ""
                if CLIENT_PLACEHOLDER in value:
                    node.text = CLIENT_PLACEHOLDER
                    after_client = True
                    continue
                if "RELATÓRIO" in value.upper():
                    after_client = False
                    continue
                if after_client:
                    node.text = ""


def _compact_back_cover(document: Document) -> None:
    if not document.tables:
        return
    contact_table = document.tables[-1]
    row_properties = contact_table.rows[0]._tr.get_or_add_trPr()
    if row_properties.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        row_properties.append(header)
    for row in contact_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    if len(contact_table.columns) >= 3:
        paragraph = contact_table.cell(0, 2).add_paragraph()
        run = paragraph.add_run("www.itprotect.com.br")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x59, 0xFC)
    body = document._element.body
    for element in reversed(list(body)):
        text = " ".join("".join(element.itertext()).split()).lower()
        if "www.itprotect.com.br" in text and element.tag == qn("w:p"):
            body.remove(element)
            break
    final_section = body[-1]
    spacer = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "1")
    spacing.set(qn("w:lineRule"), "exact")
    properties.append(spacing)
    spacer.append(properties)
    body.insert(list(body).index(final_section), spacer)

def _set_corporate_image_alt_text(document: Document) -> None:
    roots = [document._element]
    seen: set[int] = set()
    for section in document.sections:
        for container in (section.header, section.footer):
            identity = id(container._element)
            if identity in seen:
                continue
            seen.add(identity)
            roots.append(container._element)
    for root in roots:
        for node in root.iter():
            if node.tag.rsplit("}", 1)[-1] == "docPr":
                node.set("descr", "Elemento gráfico corporativo ITProtect")
                node.set("title", "Elemento gráfico corporativo ITProtect")

def _scrub_properties(document: Document) -> None:
    properties = document.core_properties
    properties.author = ""
    properties.last_modified_by = ""
    properties.title = "Modelo Tenable Cloud Security"
    properties.subject = "Relatório mensal Tenable Cloud Security"
    properties.keywords = ""
    properties.comments = ""
    properties.category = ""


def _owner_part_for_relationships(name: str) -> str | None:
    if name == "_rels/.rels":
        return None
    directory, filename = posixpath.split(name)
    if not directory.endswith("/_rels") or not filename.endswith(".rels"):
        return None
    owner_directory = directory[: -len("/_rels")]
    return posixpath.join(owner_directory, filename[: -len(".rels")])


def _resolve_target(owner_part: str | None, target: str) -> str:
    if target.startswith("/"):
        return target.lstrip("/")
    base = "" if owner_part is None else posixpath.dirname(owner_part)
    return posixpath.normpath(posixpath.join(base, target))


def _used_relationship_ids(xml_bytes: bytes) -> set[str]:
    root = ET.fromstring(xml_bytes)
    return {
        value
        for element in root.iter()
        for key, value in element.attrib.items()
        if key.startswith("{" + _OFFICE_REL_NS + "}")
    }


def _privacy_and_media_scrub(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as package:
        entries = {name: package.read(name) for name in package.namelist()}

    removed_parts = {
        name
        for name in entries
        if name == "docProps/custom.xml"
        or name.startswith("word/comments")
        or name.startswith("word/people")
    }
    retained_image_targets: set[str] = set()
    rewritten: dict[str, bytes] = {}

    for name, payload in entries.items():
        if not name.endswith(".rels"):
            continue
        owner_part = _owner_part_for_relationships(name)
        owner_payload = entries.get(owner_part or "")
        used_ids = _used_relationship_ids(owner_payload) if owner_payload else set()
        root = ET.fromstring(payload)
        changed = False
        for relationship in list(root):
            relationship_id = relationship.attrib.get("Id", "")
            relationship_type = relationship.attrib.get("Type", "")
            target = relationship.attrib.get("Target", "")
            resolved = _resolve_target(owner_part, target)
            remove = resolved in removed_parts
            if relationship_type.endswith("/image"):
                if owner_payload is not None and relationship_id not in used_ids:
                    remove = True
                else:
                    retained_image_targets.add(resolved)
            if remove:
                root.remove(relationship)
                changed = True
        if changed:
            ET.register_namespace("", _REL_NS)
            payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        rewritten[name] = payload

    content_types = entries.get("[Content_Types].xml")
    if content_types:
        root = ET.fromstring(content_types)
        for child in list(root):
            part_name = child.attrib.get("PartName", "").lstrip("/")
            if part_name in removed_parts:
                root.remove(child)
        ET.register_namespace("", _CONTENT_TYPES_NS)
        rewritten["[Content_Types].xml"] = ET.tostring(
            root,
            encoding="utf-8",
            xml_declaration=True,
        )

    with tempfile.NamedTemporaryFile(
        suffix=".docx",
        delete=False,
        dir=path.parent,
    ) as handle:
        temporary = Path(handle.name)
    try:
        with zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as package:
            for name, payload in entries.items():
                if name in removed_parts:
                    continue
                if name.startswith("word/media/") and name not in retained_image_targets:
                    continue
                package.writestr(name, rewritten.get(name, payload))
        temporary.replace(path)
    finally:
        temporary.unlink(missing_ok=True)


def distill_cloud_template(
    *,
    source: str | Path,
    output: str | Path,
    forbidden_term_file: str | Path | None = None,
) -> Path:
    source_path = Path(source).resolve()
    output_path = Path(output).resolve()
    if source_path == output_path:
        raise ValueError("O template de origem nunca pode ser sobrescrito.")
    if not source_path.is_file():
        raise FileNotFoundError(f"Modelo Cloud de origem não encontrado: {source_path}")
    source_hash = _sha256(source_path)
    forbidden_terms = _read_forbidden_terms(forbidden_term_file)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="cloud-template-") as directory:
        working = Path(directory) / "cloud-template.docx"
        shutil.copy2(source_path, working)
        document = Document(working)
        _distill_body(document, forbidden_terms)
        _sanitize_headers_and_footers(document, forbidden_terms)
        _compact_back_cover(document)
        _set_corporate_image_alt_text(document)
        _scrub_properties(document)
        update_fields = document.settings.element.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            document.settings.element.append(update_fields)
        update_fields.set(qn("w:val"), "true")
        document.save(working)
        _privacy_and_media_scrub(working)
        shutil.copy2(working, output_path)

    if _sha256(source_path) != source_hash:
        output_path.unlink(missing_ok=True)
        raise RuntimeError("O arquivo de origem foi alterado durante a destilação.")
    return output_path


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Destila um modelo DOCX Cloud sanitizado sem alterar a origem."
    )
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--forbidden-term-file")
    return parser


def main() -> int:
    args = _parser().parse_args()
    output = distill_cloud_template(
        source=args.source,
        output=args.output,
        forbidden_term_file=args.forbidden_term_file,
    )
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
