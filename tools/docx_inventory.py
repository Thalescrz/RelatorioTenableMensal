#!/usr/bin/env python3
"""Inventaria um DOCX e rasteriza um PDF já exportado pelo Word.

Ferramenta apenas para discovery local. Os artefatos gerados podem conter dados
sensíveis e devem permanecer em diretórios ignorados pelo Git.
"""

from __future__ import annotations

import argparse
import json
import math
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any, Iterator

import pypdfium2 as pdfium
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def twips_to_inches(value: Any) -> float | None:
    if value is None:
        return None
    return round(int(value) / 1440, 3)


def emu_to_inches(value: Any) -> float | None:
    if value is None:
        return None
    return round(int(value) / 914400, 3)


def iter_blocks(parent: DocumentObject) -> Iterator[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def color_value(run: Any) -> str | None:
    color = run.font.color.rgb
    return str(color) if color is not None else None


def run_signature(run: Any) -> tuple[Any, ...]:
    size = round(run.font.size.pt, 2) if run.font.size else None
    return (
        run.font.name,
        size,
        bool(run.bold) if run.bold is not None else None,
        bool(run.italic) if run.italic is not None else None,
        color_value(run),
    )


def paragraph_record(paragraph: Paragraph, index: int) -> dict[str, Any]:
    text = clean_text(paragraph.text)
    runs = []
    for run in paragraph.runs:
        run_text = clean_text(run.text)
        if not run_text:
            continue
        name, size, bold, italic, color = run_signature(run)
        runs.append(
            {
                "text": run_text,
                "font": name,
                "size_pt": size,
                "bold": bold,
                "italic": italic,
                "color": color,
            }
        )
    pPr = paragraph._p.pPr
    return {
        "index": index,
        "text": text,
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
        "page_break_before": bool(
            paragraph.paragraph_format.page_break_before
            or (pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None)
        ),
        "keep_with_next": paragraph.paragraph_format.keep_with_next,
        "runs": runs,
    }


def cell_text(cell: Any) -> str:
    return clean_text("\n".join(p.text for p in cell.paragraphs))


def table_record(table: Table, index: int) -> dict[str, Any]:
    rows = [[cell_text(cell) for cell in row.cells] for row in table.rows]
    tbl_pr = table._tbl.tblPr
    style = table.style.name if table.style else None
    return {
        "index": index,
        "style": style,
        "rows": len(table.rows),
        "columns": max((len(row.cells) for row in table.rows), default=0),
        "matrix": rows,
        "has_repeat_header": bool(
            table.rows
            and table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
        ),
        "autofit": table.autofit,
        "alignment": str(table.alignment) if table.alignment is not None else None,
        "has_borders": bool(tbl_pr.find(qn("w:tblBorders")) is not None),
    }


def section_record(section: Any, index: int) -> dict[str, Any]:
    return {
        "index": index,
        "start_type": str(section.start_type),
        "orientation": str(section.orientation),
        "page_width_in": emu_to_inches(section.page_width),
        "page_height_in": emu_to_inches(section.page_height),
        "margins_in": {
            "top": emu_to_inches(section.top_margin),
            "right": emu_to_inches(section.right_margin),
            "bottom": emu_to_inches(section.bottom_margin),
            "left": emu_to_inches(section.left_margin),
            "header": emu_to_inches(section.header_distance),
            "footer": emu_to_inches(section.footer_distance),
        },
        "header_linked_to_previous": section.header.is_linked_to_previous,
        "footer_linked_to_previous": section.footer.is_linked_to_previous,
        "header_text": clean_text("\n".join(p.text for p in section.header.paragraphs)),
        "first_page_header_text": clean_text(
            "\n".join(p.text for p in section.first_page_header.paragraphs)
        ),
        "footer_text": clean_text("\n".join(p.text for p in section.footer.paragraphs)),
        "different_first_page": bool(section.different_first_page_header_footer),
    }


def ooxml_inventory(docx_path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {}
    with zipfile.ZipFile(docx_path) as archive:
        names = archive.namelist()
        doc_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        settings_xml = (
            archive.read("word/settings.xml").decode("utf-8", errors="replace")
            if "word/settings.xml" in names
            else ""
        )
        result["package_parts"] = {
            "images": sorted(n for n in names if n.startswith("word/media/")),
            "charts": sorted(n for n in names if n.startswith("word/charts/chart")),
            "embeddings": sorted(n for n in names if n.startswith("word/embeddings/")),
            "headers": sorted(n for n in names if re.match(r"word/header\d+\.xml$", n)),
            "footers": sorted(n for n in names if re.match(r"word/footer\d+\.xml$", n)),
            "comments": "word/comments.xml" in names,
            "footnotes": "word/footnotes.xml" in names,
            "endnotes": "word/endnotes.xml" in names,
            "theme": "word/theme/theme1.xml" in names,
        }
        result["xml_counts"] = {
            "drawings": doc_xml.count("<w:drawing"),
            "pict_vml": doc_xml.count("<w:pict"),
            "textboxes": doc_xml.count("<w:txbxContent"),
            "manual_page_breaks": len(re.findall(r'<w:br[^>]+w:type="page"', doc_xml)),
            "rendered_page_breaks": doc_xml.count("<w:lastRenderedPageBreak"),
            "section_properties": doc_xml.count("<w:sectPr"),
            "bookmarks": doc_xml.count("<w:bookmarkStart"),
            "fields": doc_xml.count("<w:fldChar"),
            "tracked_insertions": doc_xml.count("<w:ins"),
            "tracked_deletions": doc_xml.count("<w:del"),
        }
        result["field_instructions"] = [
            clean_text(re.sub(r"<[^>]+>", "", value))
            for value in re.findall(r"<w:instrText[^>]*>(.*?)</w:instrText>", doc_xml, re.S)
        ]
        result["settings"] = {
            "track_revisions": "<w:trackRevisions" in settings_xml,
            "document_protection": "<w:documentProtection" in settings_xml,
            "update_fields_on_open": "<w:updateFields" in settings_xml,
        }
    return result


def inventory_docx(docx_path: Path) -> dict[str, Any]:
    doc = Document(docx_path)
    paragraphs = []
    tables = []
    order = []
    style_counter: Counter[str] = Counter()
    font_counter: Counter[str] = Counter()
    size_counter: Counter[str] = Counter()
    color_counter: Counter[str] = Counter()

    p_index = 0
    t_index = 0
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            record = paragraph_record(block, p_index)
            paragraphs.append(record)
            order.append({"type": "paragraph", "index": p_index})
            if record["text"]:
                style_counter[record["style"] or "[none]"] += 1
            for run in record["runs"]:
                if run["font"]:
                    font_counter[run["font"]] += 1
                if run["size_pt"]:
                    size_counter[str(run["size_pt"])] += 1
                if run["color"]:
                    color_counter[run["color"]] += 1
            p_index += 1
        else:
            record = table_record(block, t_index)
            tables.append(record)
            order.append({"type": "table", "index": t_index})
            t_index += 1

    core = doc.core_properties
    return {
        "file": str(docx_path),
        "core_properties": {
            "title": core.title,
            "subject": core.subject,
            "author": core.author,
            "last_modified_by": core.last_modified_by,
            "created": core.created.isoformat() if core.created else None,
            "modified": core.modified.isoformat() if core.modified else None,
            "revision": core.revision,
        },
        "counts": {
            "paragraphs": len(doc.paragraphs),
            "nonempty_paragraphs": sum(1 for p in paragraphs if p["text"]),
            "tables": len(doc.tables),
            "inline_shapes": len(doc.inline_shapes),
            "sections": len(doc.sections),
        },
        "styles_used": style_counter.most_common(),
        "direct_fonts": font_counter.most_common(),
        "direct_sizes_pt": size_counter.most_common(),
        "direct_colors": color_counter.most_common(),
        "sections": [section_record(s, i) for i, s in enumerate(doc.sections)],
        "paragraphs": paragraphs,
        "tables": tables,
        "document_order": order,
        "ooxml": ooxml_inventory(docx_path),
    }


def render_pdf(pdf_path: Path, output_dir: Path, dpi: int) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    document = pdfium.PdfDocument(pdf_path)
    scale = dpi / 72.0
    paths: list[Path] = []
    for index in range(len(document)):
        page = document[index]
        bitmap = page.render(scale=scale, rev_byteorder=True)
        image = bitmap.to_pil().convert("RGB")
        path = output_dir / f"page-{index + 1:03d}.png"
        image.save(path, format="PNG", optimize=True)
        paths.append(path)
        page.close()
    document.close()
    return paths


def make_contact_sheets(page_paths: list[Path], output_dir: Path, columns: int = 3) -> list[Path]:
    if not page_paths:
        return []
    thumb_width = 360
    label_height = 34
    margin = 18
    per_sheet = 12
    font = ImageFont.load_default()
    results: list[Path] = []

    for sheet_index in range(math.ceil(len(page_paths) / per_sheet)):
        batch = page_paths[sheet_index * per_sheet : (sheet_index + 1) * per_sheet]
        thumbs: list[tuple[Image.Image, str]] = []
        for path in batch:
            with Image.open(path) as source:
                image = source.convert("RGB")
                ratio = thumb_width / image.width
                thumb = image.resize((thumb_width, int(image.height * ratio)), Image.Resampling.LANCZOS)
                thumbs.append((thumb, path.stem))
        rows = math.ceil(len(thumbs) / columns)
        cell_height = max(img.height for img, _ in thumbs) + label_height
        sheet = Image.new(
            "RGB",
            (
                columns * thumb_width + (columns + 1) * margin,
                rows * cell_height + (rows + 1) * margin,
            ),
            "#d9dde3",
        )
        draw = ImageDraw.Draw(sheet)
        for i, (thumb, label) in enumerate(thumbs):
            row, column = divmod(i, columns)
            x = margin + column * (thumb_width + margin)
            y = margin + row * (cell_height + margin)
            sheet.paste(thumb, (x, y + label_height))
            draw.text((x + 4, y + 8), label, fill="black", font=font)
        out_path = output_dir / f"contact-sheet-{sheet_index + 1:02d}.png"
        sheet.save(out_path, format="PNG", optimize=True)
        results.append(out_path)
    return results


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--pdf", type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--dpi", type=int, default=144)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    inventory = inventory_docx(args.docx)
    structure_path = args.output_dir / "structure.json"
    structure_path.write_text(json.dumps(inventory, ensure_ascii=False, indent=2), encoding="utf-8")
    pages_dir = args.output_dir / "pages"
    pages = render_pdf(args.pdf, pages_dir, args.dpi) if args.pdf else []
    sheets = make_contact_sheets(pages, args.output_dir)
    print(
        json.dumps(
            {
                "structure": str(structure_path),
                "page_count": len(pages),
                "pages_dir": str(pages_dir),
                "contact_sheets": [str(path) for path in sheets],
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
